"""
Generatore unico delle 5 schede PG (anteprime).
Replica esattamente il layout di ANTEPRIMA_Yamamoto_Kenji.docx originale,
applicando i dati specifici di ciascun personaggio.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Colori ──
GOLD = RGBColor(0xB4, 0x96, 0x50)
NAVY = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x78, 0x78, 0x78)
DARK_RED = RGBColor(0x8B, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CREAM_HEX = "F5F3EB"
INDENT = Emu(107950)
GOU_INDENT = Emu(215900)

# ── Descrizioni attributi (uguali per tutti i PG, da Regolamento v1.3) ──
ATTR_DESC = {
    "Distacco": "Non farti coinvolgere emotivamente. Restare lucido davanti all'orrore. Mantenere la distanza professionale quando tutto ti spinge a reagire.",
    "Pazienza": "Resistere alla fretta. Aspettare il momento giusto. Sopportare lunghe attese senza cedere alla pressione del tempo.",
    "Silenzio": "Calma interiore. Non reagire d'impulso. Incassare provocazioni senza esplodere. Non tradire le tue intenzioni con un gesto o una parola di troppo.",
    "Lucidità": "Vedere chiaro. Collegare i pezzi. Capire cosa è importante e cosa no. Ricostruire cosa è successo da quello che resta.",
    "Ascolto": "Far parlare gli altri. Cogliere il non detto. Leggere tra le righe. Sentire quando qualcosa non torna in quello che qualcuno ti sta dicendo.",
    "Presenza": "L'autorità che ispiri. Farti prendere sul serio. Intimidire, comandare attenzione, far pesare la tua posizione in una stanza.",
}
ATTR_KANJI = {
    "Distacco": "冷静",
    "Pazienza": "忍耐",
    "Silenzio": "沈黙",
    "Lucidità": "明晰",
    "Ascolto": "傾聴",
    "Presenza": "存在",
}


# ── Helpers ──
def add_run(para, text, size=10, bold=False, italic=False, color=NAVY):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def remove_all_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/><w:left w:val="none"/>'
        f'<w:right w:val="none"/><w:bottom w:val="none"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_valign(cell, val="center"):
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    cell._tc.get_or_add_tcPr().append(vAlign)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcW')):
        tcPr.remove(ex)
    tcPr.append(parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>'))


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "━" * 70, size=6, color=GOLD)


def add_section_header(doc, kanji, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, kanji, size=13, color=GOLD)
    add_run(p, "  ", size=13)
    add_run(p, title, size=11, bold=True, color=NAVY)


def add_text(doc, text, size=10, color=NAVY, italic=False, bold=False, indent=True, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = INDENT
    add_run(p, text, size=size, color=color, italic=italic, bold=bold)
    return p


def add_field(doc, label, value, label_color=GRAY, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = INDENT
    add_run(p, f"{label}: ", size=size, bold=True, color=label_color)
    add_run(p, value, size=size, color=NAVY)


def add_gou(doc, name, desc, attributo, costo, successo, fallimento):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = INDENT
    add_run(p, "☐ ", size=11, color=NAVY)
    add_run(p, name, size=10, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, desc, size=9, italic=True, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, "Attributo: ", size=9, bold=True, color=GRAY)
    add_run(p, attributo, size=9, color=GRAY)
    add_run(p, "    ", size=9)
    add_run(p, "Costo: ", size=9, bold=True, color=GRAY)
    add_run(p, costo, size=9, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, "Successo: ", size=9, bold=True, color=GRAY)
    add_run(p, successo, size=9, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, "Fallimento: ", size=9, bold=True, color=GRAY)
    add_run(p, fallimento, size=9, color=NAVY)


# ── Costruzione scheda ──
def genera_scheda(d, output_filename, image_path=None):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # Header riga GENKAI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "限界  ·  GENKAI v1.3  ·  限界", size=8, color=GOLD)

    # ── TABLE 0 (1x3): nome | info | foto ──
    t0 = doc.add_table(rows=1, cols=3)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER

    c_name = t0.rows[0].cells[0]
    set_cell_width(c_name, 8)
    set_cell_valign(c_name, "center")
    remove_all_borders(c_name)
    p = c_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, d["nome_kanji"], size=24, color=NAVY)
    p = c_name.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, d["nome_romanizzato"], size=13, bold=True, color=NAVY)

    c_info = t0.rows[0].cells[1]
    set_cell_width(c_info, 9)
    set_cell_valign(c_info, "center")
    remove_all_borders(c_info)
    p = c_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Età ", size=10, bold=True, color=GRAY)
    add_run(p, d["eta"], size=10, color=NAVY)
    p = c_info.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Ruolo ", size=10, bold=True, color=GRAY)
    add_run(p, d["ruolo"], size=10, color=NAVY)
    p = c_info.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Grado ", size=10, bold=True, color=GRAY)
    add_run(p, d["grado"], size=10, color=NAVY)
    p = c_info.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Servizio ", size=10, bold=True, color=GRAY)
    add_run(p, d["servizio"], size=10, color=NAVY)

    c_foto = t0.rows[0].cells[2]
    set_cell_width(c_foto, 4)
    set_cell_valign(c_foto, "center")
    remove_all_borders(c_foto)
    p_foto = c_foto.paragraphs[0]
    p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path and os.path.exists(image_path):
        run = p_foto.add_run()
        run.add_picture(image_path, width=Cm(3))
    else:
        add_run(p_foto, "[FOTO]", size=9, italic=True, color=GRAY)

    add_separator(doc)

    # ── ATTRIBUTI ──
    add_section_header(doc, "能力", "ATTRIBUTI")

    # Table 1 (7 righe x 6 colonne): Attributo | Base | Finale | (sep) | Valori | Note
    t1 = doc.add_table(rows=7, cols=6)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Attributo", "Base", "Finale", "", "Valori", "Note"]
    widths = [3.0, 1.2, 1.2, 0.3, 2.0, 11.3]
    for i, h in enumerate(headers):
        cell = t1.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        if i in (1, 2, 4):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if h:
            add_run(p, h, size=9, bold=True, color=WHITE)

    attrs = [
        ("Distacco", d["distacco_base"]),
        ("Pazienza", d["pazienza_base"]),
        ("Silenzio", d["silenzio_base"]),
        ("Lucidità", d["lucidita_base"]),
        ("Ascolto", d["ascolto_base"]),
        ("Presenza", d["presenza_base"]),
    ]
    for i, (name, base) in enumerate(attrs):
        row = t1.rows[i + 1]
        for ci in range(6):
            set_cell_width(row.cells[ci], widths[ci])
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, CREAM_HEX)

        # Col 0: nome + kanji
        p = row.cells[0].paragraphs[0]
        add_run(p, name + " ", size=10, bold=True, color=NAVY)
        add_run(p, ATTR_KANJI[name], size=10, color=NAVY)

        # Col 1: base
        p = row.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, base, size=12, bold=True, color=NAVY)

        # Col 2: finale (vuoto da riempire a mano)
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Col 3: separatore (vuoto)
        # Col 4: valori (vuoto da riempire a mano in sessione)
        p = row.cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Col 5: descrizione
        p = row.cells[5].paragraphs[0]
        add_run(p, ATTR_DESC[name], size=8, color=NAVY)

    # Bordi sottili sotto ogni riga
    for row in t1.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/>'
                f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="D0C8B0"/>'
                f'</w:tcBorders>'))

    # Nota sotto tabella
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = INDENT
    add_run(p, "3 punti da distribuire ", size=9, bold=True, color=GRAY)
    add_run(p, "(max 2 per attributo, max 8)", size=9, color=GRAY)

    add_separator(doc)

    # ── KI / SATORI (table 1x2) ──
    t2 = doc.add_table(rows=1, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER

    c_ki = t2.rows[0].cells[0]
    set_cell_width(c_ki, 12)
    set_cell_valign(c_ki, "top")
    remove_all_borders(c_ki)
    p = c_ki.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "気", size=12, color=GOLD)
    add_run(p, "  KI  — Energia Vitale", size=10, bold=True, color=NAVY)
    p = c_ki.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Ki Max: ________    ", size=11, bold=True, color=NAVY)
    add_run(p, "attributo più basso + 2d6 prendi il più alto; l'1 si ritira una sola volta. Il dado si ritira a ogni nuovo caso; i comprati si sommano sempre", size=8, italic=True, color=DARK_RED)
    p = c_ki.add_paragraph()
    add_run(p, "Ki attuale: ________________ ", size=11, bold=True, color=NAVY)
    add_run(p, "≤3 Genkai", size=9, italic=True, color=DARK_RED)
    p = c_ki.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    add_run(p, "Soroban: ________ ", size=11, bold=True, color=NAVY)
    add_run(p, "parte da 5 (0-9); +1 su 3, +2 su 2, -1 su 11, -2 su 12; decide il recupero notturno; domani riparti dalla chiusura, avvicinata di 1 al riposo", size=8, italic=True, color=GRAY)

    c_sat = t2.rows[0].cells[1]
    set_cell_width(c_sat, 7)
    set_cell_valign(c_sat, "top")
    remove_all_borders(c_sat)
    p = c_sat.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "悟り", size=12, color=GOLD)
    add_run(p, "  SATORI", size=10, bold=True, color=NAVY)
    p = c_sat.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "1/sessione — dichiarato prima del tiro", size=9, color=GRAY)
    p = c_sat.add_paragraph()
    add_run(p, "il dado vale 2, niente Kiwami ☐", size=9, color=NAVY)

    add_separator(doc)

    # ── NASAKE / SOROBAN (table 1x2) ──
    t3 = doc.add_table(rows=1, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER

    c_nas = t3.rows[0].cells[0]
    set_cell_width(c_nas, 9.5)
    set_cell_valign(c_nas, "top")
    remove_all_borders(c_nas)
    p = c_nas.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "情け", size=12, color=GOLD)
    add_run(p, "  NASAKE — Compassione", size=10, bold=True, color=NAVY)
    p = c_nas.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Solo donabile a un altro PG giocando la scena.\nSi perde a fine sessione.", size=9, color=GRAY)
    p = c_nas.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    add_run(p, "☐ Vuoto   ☐ Pieno", size=10, bold=True, color=NAVY)

    c_sor = t3.rows[0].cells[1]
    set_cell_width(c_sor, 9.5)
    set_cell_valign(c_sor, "top")
    remove_all_borders(c_sor)
    p = c_sor.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "算盤", size=12, color=GOLD)
    add_run(p, "  SOROBAN — Giornata", size=10, bold=True, color=NAVY)
    p = c_sor.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Nami+ → +1    Kiwami+ → +2\nNami- → -1     Kiwami- → -2", size=9, color=GRAY)
    p = c_sor.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    add_run(p, "0  1  2  3  4  ", size=10, color=GRAY)
    add_run(p, "⑤", size=12, bold=True, color=GOLD)
    add_run(p, "  6  7  8  9", size=10, color=GRAY)

    add_separator(doc)

    # ── GOU ──
    add_section_header(doc, "業", "GOU — IL DEBITO  (scegli uno)")
    add_text(doc, "Il Gou funziona SEMPRE. Successo = preciso, Fallimento = vago. Ogni uso raddoppia il costo del successivo; una notte di sonno lo riabbassa di un grado. Il costo si paga per intero: se ti porterebbe a 0 o sotto, il Gou non si attiva.",
             size=9, color=GRAY, italic=True, after=2)

    add_gou(doc, d["gou_1_nome"], d["gou_1_desc"], d["gou_1_attributo"],
            d["gou_1_costo"], d["gou_1_successo"], d["gou_1_fallimento"])
    add_gou(doc, d["gou_2_nome"], d["gou_2_desc"], d["gou_2_attributo"],
            d["gou_2_costo"], d["gou_2_successo"], d["gou_2_fallimento"])
    add_gou(doc, d["gou_3_nome"], d["gou_3_desc"], d["gou_3_attributo"],
            d["gou_3_costo"], d["gou_3_successo"], d["gou_3_fallimento"])

    add_separator(doc)

    # ── SENMON ──
    add_section_header(doc, "専門", "SENMON — SPECIALIZZAZIONI  (scegli una, grado 1)")
    add_text(doc, "Parti da Praticante: +1 ai tiri pertinenti, conosci le cose comuni del campo senza tiro. Cresce con gli usi e i punti Shugyō — vedi GENKAI_Specializzazioni.md.",
             size=9, color=GRAY, italic=True, after=2)
    for i in (1, 2, 3):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = INDENT
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "☐ ", size=10, color=NAVY)
        add_run(p, d[f"senmon_{i}_nome"], size=10, bold=True, color=NAVY)
        add_run(p, f"  ({d[f'senmon_{i}_chiave']})  —  ", size=9, color=GOLD)
        add_run(p, d[f"senmon_{i}_desc"], size=9, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INDENT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Grado: ______    Usi: ", size=10, bold=True, color=NAVY)
    add_run(p, "☐☐☐☐☐☐☐☐☐☐", size=11, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INDENT
    add_run(p, "10 usi + 19 Shugyō (attr. chiave di scheda 6+) → Esperto (+2)  ·  25 usi + 39 Shugyō (paletti) → Maestro", size=8, italic=True, color=GRAY)

    add_separator(doc)

    # ── EQUIPAGGIAMENTO ──
    add_section_header(doc, "装", "EQUIPAGGIAMENTO DI SERVIZIO")
    add_text(doc, "In borghese non porti l'arma: resta nell'armadietto in centrale e si preleva, firmando, solo per le operazioni. (Regole di scontro: GENKAI_Combattimento.md)",
             size=9, color=GRAY, italic=True, after=2)
    for line in (
        "Armadietto — Revolver New Nambu M60 (.38, 5 colpi): Lucidità · vel. 3/1 · ricarica 4 · danno 3. Addestramento base: sai usarla, non sei un tiratore scelto",
        "Armadietto — Giubbotto antiproiettile: Assorbe 3 (fisso, contro ogni colpo) · indossare 4",
        "Operazioni — Keibō (manganello): Pazienza o Silenzio · vel. 2/1 · danno 2",
        "Sempre con te — Keisatsu techō (tesserino), manette, taccuino",
        "A mani nude — Lotta 1 d'accademia: Presenza · vel. 1/1 · danno 1 (prese e immobilizzazioni; il grado 1 ce l'hanno tutti gli investigatori)",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = INDENT
        p.paragraph_format.space_after = Pt(1)
        add_run(p, "• " + line, size=9, color=GRAY)

    add_separator(doc)

    # ── SHUGYO ──
    add_section_header(doc, "修行", "SHUGYŌ — CRESCITA")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INDENT
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Punti Shugyō: ____________________", size=11, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INDENT
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Guadagno: 1 a sessione · 4-6 a caso chiuso · +1 scena personale ben gestita · +1 momento eccezionale (max 1/sessione)", size=8, color=GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INDENT
    add_run(p, "Spesa tra i casi: attributo = arrivo ×3 · Ki max = arrivo ×4 (tetto 12) · Senmon 9/19/39 · Enja extra 12 · affinare il Gou = base ×11", size=8, color=GRAY)

    add_separator(doc)

    # ── PAGE BREAK ──
    doc.add_page_break()

    # ── CHI SEI ──
    add_section_header(doc, "人", "CHI SEI")
    add_text(doc, d["chi_sei"])

    add_separator(doc)

    # ── PROBLEMA ──
    add_section_header(doc, "影", f"IL TUO PROBLEMA — {d['problema_titolo']}")
    add_text(doc, d["problema_testo"])

    add_separator(doc)

    # ── PNG DEL PROBLEMA ──
    add_section_header(doc, "縁", d["png_nome"])
    t4 = doc.add_table(rows=4, cols=2)
    t4.autofit = True
    for i, (label, val) in enumerate([
        ("Età", d["png_eta"]),
        ("Occupazione", d["png_occupazione"]),
        ("Relazione", d["png_relazione"]),
        ("Cosa vuole", d["png_vuole"]),
    ]):
        p = t4.rows[i].cells[0].paragraphs[0]
        add_run(p, label, size=9, bold=True, color=GRAY)
        p = t4.rows[i].cells[1].paragraphs[0]
        add_run(p, val, size=10, color=NAVY)
        for cell in t4.rows[i].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/>'
                f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="D0C8B0"/>'
                f'</w:tcBorders>'))

    add_text(doc, d["png_desc"], italic=True, after=2)

    add_separator(doc)

    # ── CONOSCENZA ──
    add_section_header(doc, "縁者", f"CONOSCENZA — {d['conoscenza_nome']}")
    add_text(doc, d["conoscenza_ruolo"], italic=True, color=GRAY, after=2)
    add_text(doc, d["conoscenza_desc"], after=2)
    add_field(doc, "Contatto", d["conoscenza_contatto"], size=9)
    add_field(doc, "Costo", d["conoscenza_costo"], size=9)
    add_text(doc, d["conoscenza_limite"], size=8, italic=True, color=GRAY, after=2)

    add_separator(doc)

    # ── COME TI COMPORTI ──
    add_section_header(doc, "面", "COME TI COMPORTI")
    for kanji, label, field in [
        ("建前", "In pubblico", d["tatemae"]),
        ("本音", "In privato", d["honne"]),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = INDENT
        add_run(p, f"{kanji}  {label}: ", size=10, bold=True, color=GOLD)
        add_run(p, field, size=10, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = INDENT
    add_run(p, "Frase tipica: ", size=10, bold=True, color=GRAY)
    add_run(p, d["frase"], size=10, italic=True, color=NAVY)

    for label, field in [("Sotto pressione", d["pressione"]), ("Debolezza", d["debolezza"])]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = INDENT
        add_run(p, f"{label}: ", size=10, bold=True, color=DARK_RED)
        add_run(p, field, size=10, color=NAVY)

    add_separator(doc)

    # ── TRATTI PERSONALI ──
    add_section_header(doc, "癖", "TRATTI PERSONALI")
    for label, val in [
        ("Vizio", d["vizio"]),
        ("Tic", d["tic"]),
        ("Oggetto", d["oggetto"]),
        ("Gusto", d["gusto"]),
        ("Rituale", d["rituale"]),
    ]:
        add_field(doc, label, val, size=10)

    add_separator(doc)

    # ── LA SQUADRA ──
    add_section_header(doc, "仲", "LA SQUADRA")
    for i in range(1, 5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = INDENT
        add_run(p, f"{d[f'rapporto_{i}_nome']}: ", size=10, bold=True, color=GRAY)
        add_run(p, d[f"rapporto_{i}_testo"], size=10, italic=True, color=NAVY)

    out_path = os.path.join(BASE_DIR, output_filename)
    doc.save(out_path)
    print(f"Salvata: {output_filename}")


# =============================================================
# DATI DEI 5 PG
# =============================================================

PG_01 = {
    "nome_kanji": "山本 健二",
    "nome_romanizzato": "YAMAMOTO Kenji",
    "eta": "42",
    "ruolo": "Capo Squadra Investigativa",
    "grado": "Ispettore Capo (Keibu)",
    "servizio": "18 anni",
    "distacco_base": "6", "pazienza_base": "4", "silenzio_base": "5",
    "lucidita_base": "6", "ascolto_base": "4", "presenza_base": "5",
    "gou_1_nome": "Pugno di Ferro 鉄拳",
    "gou_1_desc": "Pressione psicologica costante e implacabile. Non un momento — un peso che non si toglie.",
    "gou_1_attributo": "Presenza",
    "gou_1_costo": "3 Ki",
    "gou_1_successo": "Il PNG cede su un punto importante — ammette, collabora, si ritira",
    "gou_1_fallimento": "Il PNG è scosso e sotto pressione, ma non cede",
    "gou_2_nome": "Teatro delle Ombre 影芝居",
    "gou_2_desc": "Guardando una scena, vedi le sagome sbiadite di ciò che è accaduto.",
    "gou_2_attributo": "Lucidità (requisito: ≥ 7)",
    "gou_2_costo": "4 Ki",
    "gou_2_successo": "2-3 momenti chiave come sagome sbiadite, nell'ordine giusto, con vuoti. Non volti, non dettagli",
    "gou_2_fallimento": "Un singolo momento congelato — una sagoma, una posizione, senza prima né dopo",
    "gou_3_nome": "Lo Spirito che Resta 残心",
    "gou_3_desc": "Lo spirito del samurai che non cede. Attivabile solo a Ki 4-5; il costo non innesca il Genkai: si verifica a fine scena.",
    "gou_3_attributo": "Distacco o Pazienza (dual)",
    "gou_3_costo": "2 Ki",
    "gou_3_successo": "Per il resto della scena, 11 e 12 contano come fallimento leggero — non crolli",
    "gou_3_fallimento": "Solo il 12 è protetto. L'11 resta Nami negativo",
    "senmon_1_nome": "Ambienti yakuza",
    "senmon_1_chiave": "Silenzio",
    "senmon_1_desc": "Diciott'anni a Kyoto: sai chi comanda in quali strade e come si parla a un kumichō senza perdere la faccia.",
    "senmon_2_nome": "Stampa e media",
    "senmon_2_chiave": "Presenza",
    "senmon_2_desc": "Sai come ragionano le redazioni, cosa fa notizia e cosa può aspettare. E hai un amico al Kyoto Shimbun.",
    "senmon_3_nome": "Interrogatorio",
    "senmon_3_chiave": "Pazienza o Ascolto",
    "senmon_3_desc": "Diciott'anni di stanze chiuse: sblocca le tecniche di apertura — il terreno emotivo che si prepara prima delle domande vere.",
    "chi_sei": "Sei un veterano rispettato. Hai risolto casi importanti, hai costruito una reputazione solida. Ma la tua vita privata è un disastro. Sei mesi fa tua moglie Yuko ti ha lasciato. Non ci sei mai, ha detto. Aveva ragione. Ora vivi in un piccolo appartamento a Fushimi, con le scatole ancora da disfare. Tuo figlio Takeshi (8 anni) vive con la madre. Il martedì e il giovedì è con te. O dovrebbe esserlo. Il problema: il lavoro viene sempre prima.",
    "problema_titolo": "LA FAMIGLIA",
    "problema_testo": "Tua sorella Yamamoto Noriko è stata il tuo supporto durante il divorzio. Ha coperto le tue assenze, ha badato a Takeshi quando tu non potevi. Ma la pazienza di Noriko ha un limite. L'ultima volta che le hai chiesto aiuto, te lo ha detto chiaramente: Non sono la tua baby-sitter di riserva. Devi scegliere cosa conta davvero. Sai che Noriko è stanca. Sai che il rapporto con Takeshi è fragile.",
    "png_nome": "YAMAMOTO Noriko",
    "png_eta": "38 anni",
    "png_occupazione": "Insegnante di ikebana",
    "png_relazione": "Sorella minore, una volta eravate molto uniti",
    "png_vuole": "Che tu ammetta di avere un problema. Non singoli episodi — il pattern.",
    "png_desc": "Non urla. È peggio. Parla con calma, con quella delusione quieta che ferisce più della rabbia. Fa domande scomode.",
    "conoscenza_nome": "TANAKA Shuichi",
    "conoscenza_ruolo": "Giornalista al Kyoto Shimbun",
    "conoscenza_desc": "Vecchio compagno di università. Vi vedete per bere ogni tanto. Accesso agli archivi del giornale, voci di corridoio, contatti nel mondo della stampa.",
    "conoscenza_contatto": "Cellulare personale. Risponde sempre, anche di notte.",
    "conoscenza_costo": "Ogni tanto una soffiata, niente di compromettente.",
    "conoscenza_limite": "Una volta per sessione senza conseguenze. La seconda volta, chiederà qualcosa in cambio.",
    "tatemae": "Sei il capo. Calmo, competente, rispettato. Non mostri mai debolezza davanti alla squadra.",
    "honne": "Sei esausto. Ti senti in colpa per Takeshi. A volte ti chiedi se ne vale la pena.",
    "frase": "Abbiamo un lavoro da fare. Concentriamoci su quello.",
    "pressione": "Ti chiudi. Parli meno. Diventi più brusco. Non deleghi — fai tutto tu.",
    "debolezza": "Quando qualcuno mette in dubbio le sue priorità tra lavoro e famiglia.",
    "vizio": "Caffè nero, almeno sei tazze al giorno.",
    "tic": "Si massaggia il ponte del naso quando è stanco o frustrato.",
    "oggetto": "Un orologio Seiko automatico, regalo di suo padre per la promozione.",
    "gusto": "Udon caldo da un chiosco vicino alla centrale. Sempre lo stesso, sempre in piedi.",
    "rituale": "Ogni mattina, caffè Boss in lattina al distributore.",
    "rapporto_1_nome": "HONDA Ryota",
    "rapporto_1_testo": "Brillante, ma gioca col fuoco. Un giorno si brucerà.",
    "rapporto_2_nome": "NAKAMURA Shota",
    "rapporto_2_testo": "Il migliore di noi negli interrogatori. Mi fido di lui più di chiunque altro.",
    "rapporto_3_nome": "SATO Yuki",
    "rapporto_3_testo": "Mi ricorda me a quell'età. Spero che non faccia i miei stessi errori.",
    "rapporto_4_nome": "FUJITA Emi",
    "rapporto_4_testo": "Competente, riservata, non chiede mai nulla. Questo mi preoccupa.",
}

PG_02 = {
    "nome_kanji": "本田 涼太",
    "nome_romanizzato": "HONDA Ryota",
    "eta": "35",
    "ruolo": "Analista Scene del Crimine",
    "grado": "Sergente (Junsa-bucho)",
    "servizio": "10 anni",
    "distacco_base": "6", "pazienza_base": "4", "silenzio_base": "4",
    "lucidita_base": "7", "ascolto_base": "4", "presenza_base": "5",
    "gou_1_nome": "Occhio della Gru 鶴の目",
    "gou_1_desc": "Vedi ciò che altri non vedono — il dettaglio che cambia tutto.",
    "gou_1_attributo": "Lucidità",
    "gou_1_costo": "3 Ki",
    "gou_1_successo": "Trovi il dettaglio nascosto e capisci in che direzione punta",
    "gou_1_fallimento": "Percepisci quanti elementi fuori posto ci sono (uno, più di uno, molti) — non sai cosa sono",
    "gou_2_nome": "Cuore di Ghiaccio 氷の心",
    "gou_2_desc": "Ti distacchi completamente dalla situazione. Niente ti tocca.",
    "gou_2_attributo": "Distacco",
    "gou_2_costo": "3 Ki",
    "gou_2_successo": "Il primo tiro emotivo della scena è successo automatico",
    "gou_2_fallimento": "+2 al prossimo tiro emotivo",
    "gou_3_nome": "La Brace che Resta 残り火",
    "gou_3_desc": "Le emozioni lasciano calore. Il luogo è ancora tiepido di ciò che è stato provato.",
    "gou_3_attributo": "Ascolto o Distacco (dual)",
    "gou_3_costo": "3 Ki",
    "gou_3_successo": "Senti le emozioni dominanti e quante presenze diverse c'erano — distingui paura da rabbia da determinazione",
    "gou_3_fallimento": "Un'impressione emotiva generale, senza distinzione tra le persone",
    "senmon_1_nome": "Rilievi e fotografia",
    "senmon_1_chiave": "Lucidità",
    "senmon_1_desc": "Documentare una scena è il tuo mestiere: luci, angoli, impronte, catena di custodia.",
    "senmon_2_nome": "Pedinamento",
    "senmon_2_chiave": "Lucidità",
    "senmon_2_desc": "Distanze, riflessi nelle vetrine, cambi di marciapiede. Vedi ogni dettaglio — anche quando sei tu a seguire.",
    "senmon_3_nome": "Sport e scommesse",
    "senmon_3_chiave": "Ascolto",
    "senmon_3_desc": "Baseball, sumo, corse. Quote, allibratori, risultati pilotati. Conosci quel mondo fin troppo bene.",
    "chi_sei": "Sei brillante. Lo sanno tutti, tu per primo. Vedi cose che altri non vedono, ricostruisci scene del crimine come se avessi assistito in persona. Il problema è che la tua mente non si ferma mai — cerchi stimoli, adrenalina, sfide. Hai iniziato con le scommesse sportive tre anni fa. Solo per divertimento. Poi le cifre sono cresciute. Ora scommetti su tutto: baseball, sumo, corse di cavalli. Vinci spesso — sei intelligente. Ma ultimamente hai perso più di quanto hai vinto.",
    "problema_titolo": "IL GIOCO D'AZZARDO",
    "problema_testo": "Devi soldi a Murakami, un allibratore indipendente. Non è yakuza — è quasi rispettabile. Ma i debiti sono debiti. Il suo uomo per le riscossioni è Goto Masaru: non minaccia mai, è educato, ragionevole, professionale — il che lo rende più inquietante. Sai che se qualcuno in centrale scopre che scommetti illegalmente, la tua carriera è finita.",
    "png_nome": "GOTO Masaru",
    "png_eta": "45 anni",
    "png_occupazione": "Consulente finanziario (riscossore per Murakami)",
    "png_relazione": "Riscossore — ti tiene sotto pressione",
    "png_vuole": "I soldi, preferibilmente. Ma è pragmatico — potrebbe accettare alternative.",
    "png_desc": "Parla di soluzioni, accordi, comprensione reciproca. Sorride sempre. Le minacce sono solo sottintese.",
    "conoscenza_nome": "ODA Takumi",
    "conoscenza_ruolo": "Informatore / piccolo ricettatore",
    "conoscenza_desc": "Ex testimone di un vecchio caso. Gli hai evitato il carcere, ti deve un favore. Conosce il sottobosco di Kyoto, sente voci, sa chi vende cosa a chi.",
    "conoscenza_contatto": "Bazzica al Bar Shinjuku a Gion, tutte le sere dopo le 21.",
    "conoscenza_costo": "Che tu continui a chiudere un occhio sui suoi affari minori.",
    "conoscenza_limite": "Una volta per sessione senza conseguenze. La seconda volta, il favore diventa più grande.",
    "tatemae": "Sei il tecnico sicuro di sé. Fai battute, smorzi la tensione. Sembri sempre sotto controllo. Sei quello che trova l'indizio che cambia tutto.",
    "honne": "L'adrenalina è una droga. Quando non c'è, ti senti vuoto. Le scommesse riempiono quel vuoto — e adesso il vuoto sta per inghiottirti.",
    "frase": "Aspetta. Guarda qui. Vedi questo? Nessuno l'aveva notato.",
    "pressione": "Diventi più arrogante. Prendi rischi. Scommetti — letteralmente e metaforicamente.",
    "debolezza": "Quando qualcuno lo chiama fortunato invece di riconoscere il suo talento. O quando i debiti vengono sfiorati.",
    "vizio": "Mild Seven. Fuma solo sulla scena del crimine — dice che lo aiuta a pensare.",
    "tic": "Fa roteare una moneta da 500 yen tra le dita quando ragiona.",
    "oggetto": "Una lente d'ingrandimento pieghevole nel taschino — vecchia, graffiata, non la cambierebbe mai.",
    "gusto": "Yakitori e birra Asahi al bancone. Mai al tavolo, sempre al bancone.",
    "rituale": "Non inizia mai un sopralluogo dal lato destro della scena. Dice che porta sfortuna.",
    "rapporto_1_nome": "YAMAMOTO Kenji",
    "rapporto_1_testo": "Il capo. Duro, giusto, ma ultimamente lo vedo stanco. Non glielo dirò mai.",
    "rapporto_2_nome": "NAKAMURA Shota",
    "rapporto_2_testo": "Troppo lento per i miei gusti, ma quando trova qualcosa è sempre la cosa giusta.",
    "rapporto_3_nome": "SATO Yuki",
    "rapporto_3_testo": "Il ragazzino del laboratorio. Bravo, ma deve smettere di chiedere il permesso per tutto.",
    "rapporto_4_nome": "FUJITA Emi",
    "rapporto_4_testo": "Mi legge come un libro aperto. Questo mi innervosisce.",
}

PG_03 = {
    "nome_kanji": "中村 翔太",
    "nome_romanizzato": "NAKAMURA Shota",
    "eta": "38",
    "ruolo": "Specialista Interrogatori",
    "grado": "Ispettore (Keibu-ho)",
    "servizio": "14 anni",
    "distacco_base": "4", "pazienza_base": "7", "silenzio_base": "4",
    "lucidita_base": "4", "ascolto_base": "7", "presenza_base": "4",
    "gou_1_nome": "Ombra della Verità 影の真実",
    "gou_1_desc": "Senti quando qualcuno mente. Non sai come, ma lo senti.",
    "gou_1_attributo": "Ascolto",
    "gou_1_costo": "3 Ki",
    "gou_1_successo": "Certezza della bugia + l'area generale (orario, persona, luogo) — non il contenuto specifico",
    "gou_1_fallimento": "Senti disonestà, ma non distingui se è bugia, nervosismo, o qualcosa di personale",
    "gou_2_nome": "Porta Socchiusa 開きかけの扉",
    "gou_2_desc": "Le persone si aprono con te. Dicono più di quanto vorrebbero.",
    "gou_2_attributo": "Ascolto",
    "gou_2_costo": "3 Ki",
    "gou_2_successo": "Il PNG rivela qualcosa che non voleva assolutamente dire",
    "gou_2_fallimento": "Il PNG lascia trapelare qualcosa, ma si ferma prima di dire troppo",
    "gou_3_nome": "La Risalita della Carpa 鯉の滝登り",
    "gou_3_desc": "Non molli. Mai. La tua pazienza è una forza della natura.",
    "gou_3_attributo": "Pazienza o Presenza (dual)",
    "gou_3_costo": "3 Ki",
    "gou_3_successo": "Il PNG si contraddice su un punto specifico — la sua storia cede sotto il peso della ripetizione",
    "gou_3_fallimento": "Il PNG è stanco, la coerenza vacilla, ma la storia regge",
    "senmon_1_nome": "Interrogatorio",
    "senmon_1_chiave": "Pazienza o Ascolto",
    "senmon_1_desc": "Il tuo mestiere. Sblocca le tecniche di apertura — il terreno emotivo che prepari prima delle domande vere.",
    "senmon_2_nome": "Conoscere il quartiere",
    "senmon_2_chiave": "Ascolto — quartiere a scelta",
    "senmon_2_desc": "Le strade dove sei cresciuto o dove lavori da anni: le facce, le abitudini, chi vede tutto e chi non parla.",
    "senmon_3_nome": "Ricerca d'archivio",
    "senmon_3_chiave": "Pazienza",
    "senmon_3_desc": "Registri, protocolli, fascicoli. Dove gli altri si arrendono dopo un'ora, tu stai ancora leggendo.",
    "chi_sei": "Sei quello che fa parlare le persone. Non con la forza — con la pazienza. Ti siedi, aspetti, ascolti. Prima o poi tutti parlano. Figlio maggiore di una famiglia modesta di Osaka, hai lavorato duro per arrivare dove sei. Tuo padre è morto quando avevi 15 anni. Tuo fratello minore Kazuo (34 anni) è l'opposto — vive di espedienti, piccoli imbrogli, soldi prestati e mai restituiti.",
    "problema_titolo": "IL FRATELLO",
    "problema_testo": "Kazuo è un peso che ti porti dietro da sempre. Gli hai già prestato soldi tre volte. Ogni volta l'ultima. Ogni volta una bugia. Tua moglie non sa di questi prestiti: se lo scoprisse, ci sarebbero problemi anche a casa. Ma Kazuo è sangue del tuo sangue, e quando ha bisogno sa sempre dove trovare suo fratello.",
    "png_nome": "NAKAMURA Kazuo",
    "png_eta": "34 anni",
    "png_occupazione": "Nessuno fisso — imprenditore",
    "png_relazione": "Fratello minore",
    "png_vuole": "Dipende dalla situazione. Ma di solito sono soldi, e di solito sono urgenti.",
    "png_desc": "Prima la simpatia. Poi il senso di colpa. Poi la minaccia velata. Conosce i tuoi punti deboli.",
    "conoscenza_nome": "NAKAMURA Hideki",
    "conoscenza_ruolo": "Avvocato penalista (cugino)",
    "conoscenza_desc": "Cugino di primo grado, siete cresciuti insieme. Consulenza legale, accesso a fascicoli pubblici, contatti in tribunale.",
    "conoscenza_contatto": "Studio legale Nakamura & Associati, Via Kawaramachi. O al cellulare.",
    "conoscenza_costo": "Niente di specifico — famiglia è famiglia.",
    "conoscenza_limite": "Una volta per sessione senza conseguenze. Hideki non fa mai nulla di illegale, ma conosce le zone grigie.",
    "tatemae": "Sei quello calmo. Non ti agiti mai. Parli poco, ascolti molto. Quando parli, le persone tendono ad ascoltare.",
    "honne": "Sei stanco di essere quello responsabile. Di dover tenere insieme tutto. A volte vorresti mandare tutto al diavolo. Ma non lo farai mai.",
    "frase": "Prendiamoci un momento. Raccontami tutto dall'inizio.",
    "pressione": "Diventi ancora più silenzioso. Ti ritiri in te stesso. Smetti di chiedere e inizi a osservare.",
    "debolezza": "Quando qualcuno usa i legami di sangue come arma. Tocca troppo vicino a casa.",
    "vizio": "Tè verde. Ne beve litri durante gli interrogatori. Offre sempre una tazza all'interrogato.",
    "tic": "Annuisce lentamente anche quando non è d'accordo. Le persone lo trovano rassicurante.",
    "oggetto": "Un taccuino Moleskine nero, pieno di appunti in calligrafia minuscola. Non usa registratori.",
    "gusto": "Ramen miso con extra chashu. Conosce tutti i posti migliori di Kyoto a memoria.",
    "rituale": "Colleziona manga shoujo. Li tiene nascosti in un cassetto della scrivania.",
    "rapporto_1_nome": "YAMAMOTO Kenji",
    "rapporto_1_testo": "Un capo che si porta il lavoro a casa e la casa al lavoro. Lo capisco fin troppo bene.",
    "rapporto_2_nome": "HONDA Ryota",
    "rapporto_2_testo": "Geniale e imprudente. Lo tengo d'occhio, anche se lui non se ne accorge.",
    "rapporto_3_nome": "SATO Yuki",
    "rapporto_3_testo": "Ha talento vero. Deve solo imparare che non tutti i problemi si risolvono col microscopio.",
    "rapporto_4_nome": "FUJITA Emi",
    "rapporto_4_testo": "L'unica che capisce il silenzio come lo capisco io. Ci rispettiamo senza bisogno di parole.",
}

PG_04 = {
    "nome_kanji": "佐藤 勇気",
    "nome_romanizzato": "SATO Yuki",
    "eta": "27",
    "ruolo": "Tecnico Analisi Tracce",
    "grado": "Agente Scelto (Junsa-cho)",
    "servizio": "3 anni",
    "distacco_base": "4", "pazienza_base": "6", "silenzio_base": "4",
    "lucidita_base": "7", "ascolto_base": "5", "presenza_base": "4",
    "gou_1_nome": "Palazzo della Memoria 記憶の宮殿",
    "gou_1_desc": "Puoi richiamare con precisione fotografica qualcosa che hai visto o sentito.",
    "gou_1_attributo": "Lucidità o Pazienza (dual)",
    "gou_1_costo": "2 Ki",
    "gou_1_successo": "Ricordi il dettaglio e anche elementi periferici che non avevi notato consciamente",
    "gou_1_fallimento": "Ricordi il dettaglio principale, ma sfocato o incompleto",
    "gou_2_nome": "L'Ora Giusta 正しい時",
    "gou_2_desc": "Sai quando è il momento perfetto per parlare, agire, colpire. Il bonus decade se non usato entro la scena.",
    "gou_2_attributo": "Pazienza",
    "gou_2_costo": "3 Ki",
    "gou_2_successo": "+3 al prossimo tiro",
    "gou_2_fallimento": "+1 al prossimo tiro",
    "gou_3_nome": "L'Istante della Caduta 散り際",
    "gou_3_desc": "Il fiore di ciliegio nel momento in cui cade. Cogliere ciò che sta per svanire.",
    "gou_3_attributo": "Pazienza o Ascolto (dual)",
    "gou_3_costo": "2 Ki",
    "gou_3_successo": "Sai cosa sta per svanire e hai un istante per agire",
    "gou_3_fallimento": "Senti urgenza, sai che qualcosa sta sfuggendo, ma non cosa",
    "senmon_1_nome": "Medicinali e veleni",
    "senmon_1_chiave": "Lucidità",
    "senmon_1_desc": "La tua laurea: farmaci, dosaggi, interazioni. Riconosci un'anomalia chimica prima ancora del laboratorio.",
    "senmon_2_nome": "Computer e reti",
    "senmon_2_chiave": "Lucidità",
    "senmon_2_desc": "Sistemi, archivi digitali, recupero dati. È il 1997: sei uno dei pochi in centrale che ci capisce qualcosa.",
    "senmon_3_nome": "Medicina",
    "senmon_3_chiave": "Lucidità",
    "senmon_3_desc": "Anatomia, traumi, referti. Il linguaggio del medico legale non ha segreti per te.",
    "chi_sei": "Sei il più giovane della squadra. Brillante, laureato con lode in chimica forense, pieno di entusiasmo. Sei quello che lavora fino a tardi, che controlla tre volte, che non si arrende finché non trova la risposta. Il problema è tua madre: Sato Michiko non ha mai accettato che tu facessi il poliziotto. Uno spreco, dice. Tuo zio Tanaka Jiro ha un'azienda di import-export e ti aspetta.",
    "problema_titolo": "LA MADRE",
    "problema_testo": "Tua madre chiama. Spesso. A volte in centrale. È urgente, dice alla segretaria. Non è mai urgente. Vuole sapere se hai ripensato al lavoro dello zio, quando ti sposi, perché sprechi la tua vita. Non urla, non minaccia — usa il senso di colpa come un'arma di precisione. E non puoi riattaccare in faccia a tua madre — non in Giappone.",
    "png_nome": "SATO Michiko",
    "png_eta": "58 anni",
    "png_occupazione": "Casalinga, vedova da 10 anni",
    "png_relazione": "Madre",
    "png_vuole": "Che tu sistemi la tua vita. Matrimonio, lavoro stabile, nipotini.",
    "png_desc": "Mai aggressiva direttamente. Io voglio solo il tuo bene. Tuo padre sarebbe così deluso. Il senso di colpa è la sua arma principale.",
    "conoscenza_nome": "KATO Hiroshi",
    "conoscenza_ruolo": "Tecnico informatico (amico d'infanzia)",
    "conoscenza_desc": "Migliori amici dalle elementari, vi vedete ancora per giocare ai videogiochi. Recupero dati, analisi file, hacking leggero.",
    "conoscenza_contatto": "Telefono di casa. Risponde anche a tarda sera.",
    "conoscenza_costo": "Storie dal lavoro, cene, compagnia.",
    "conoscenza_limite": "Una volta per sessione senza conseguenze. Kato è l'unica persona con cui puoi essere completamente te stesso.",
    "tatemae": "Sei il giovane promettente. Entusiasta, competente, sempre pronto ad aiutare. Vuoi dimostrare di meritare il tuo posto.",
    "honne": "Sei esausto dal peso delle aspettative. Di tua madre, del lavoro, di te stesso. A volte ti chiedi se non avresti fatto meglio ad ascoltarla.",
    "frase": "Datemi ancora un'ora, ho quasi finito l'analisi.",
    "pressione": "Parli troppo. Riempi i silenzi. Ti giustifichi anche quando non devi. Cerchi approvazione.",
    "debolezza": "Quando qualcuno lo tratta da ragazzino o mette in dubbio la sua esperienza.",
    "vizio": "Gomme da masticare alla menta. Ne consuma un pacchetto al giorno in laboratorio.",
    "tic": "Si sistema gli occhiali spingendoli sul naso con l'indice, anche quando non scivolano.",
    "oggetto": "Un portachiavi a forma di struttura molecolare, regalo di laurea dei compagni.",
    "gusto": "Curry rice della mensa della centrale. Lo mangia quasi ogni giorno, senza vergogna.",
    "rituale": "Conta i passi quando è nervoso. Non se ne accorge, ma i colleghi sì.",
    "rapporto_1_nome": "YAMAMOTO Kenji",
    "rapporto_1_testo": "L'ispettore. Lo ammiro, ma a volte ho paura di deluderlo.",
    "rapporto_2_nome": "HONDA Ryota",
    "rapporto_2_testo": "Il più figo della squadra. Vorrei avere la sua sicurezza. O almeno fingerla così bene.",
    "rapporto_3_nome": "NAKAMURA Shota",
    "rapporto_3_testo": "Calmo come un lago. Quando parla con me, mi sento meno ansioso.",
    "rapporto_4_nome": "FUJITA Emi",
    "rapporto_4_testo": "A volte mi guarda come se sapesse esattamente cosa sto pensando. Probabilmente lo sa.",
}

PG_05 = {
    "nome_kanji": "藤田 恵美",
    "nome_romanizzato": "FUJITA Emi",
    "eta": "36",
    "ruolo": "Profiler / Psicologa Investigativa",
    "grado": "Ispettore (Keibu-ho)",
    "servizio": "8 anni",
    "distacco_base": "4", "pazienza_base": "4", "silenzio_base": "7",
    "lucidita_base": "5", "ascolto_base": "6", "presenza_base": "4",
    "gou_1_nome": "Specchio dell'Anima 魂の鏡",
    "gou_1_desc": "Vedi oltre la maschera. Senti le emozioni vere di una persona.",
    "gou_1_attributo": "Silenzio",
    "gou_1_costo": "2 Ki",
    "gou_1_successo": "Senti l'emozione dominante e capisci perché la prova",
    "gou_1_fallimento": "Senti l'emozione, ma non la sua origine",
    "gou_2_nome": "L'Eco della Montagna 山彦",
    "gou_2_desc": "Le parole non scompaiono. Nel tuo silenzio, le senti.",
    "gou_2_attributo": "Silenzio (requisito: ≥ 7)",
    "gou_2_costo": "4 Ki",
    "gou_2_successo": "Frammenti di frasi riconoscibili — parole chiave, toni, numero di voci presenti",
    "gou_2_fallimento": "Suoni indistinti — un tono emotivo generale senza parole",
    "gou_3_nome": "Tocco del Medico 医者の手",
    "gou_3_desc": "Leggi il corpo e la mente. Vedi i segni che altri ignorano.",
    "gou_3_attributo": "Lucidità",
    "gou_3_costo": "2 Ki",
    "gou_3_successo": "Dettagli precisi — tipo di trauma, natura della paura, segni di stress cronico",
    "gou_3_fallimento": "Capisci che qualcosa non va, ma non riesci a definirlo con precisione",
    "senmon_1_nome": "Copertura e travestimento",
    "senmon_1_chiave": "Silenzio",
    "senmon_1_desc": "Reggere un ruolo è psicologia applicata: postura, voce, storia di copertura.",
    "senmon_2_nome": "Mondo della notte",
    "senmon_2_chiave": "Presenza",
    "senmon_2_desc": "Bar, hostess club, mizu shōbai: i posti dove le maschere cadono. Sai chi tiene i conti e chi i segreti.",
    "senmon_3_nome": "Arte e antiquariato",
    "senmon_3_chiave": "Lucidità",
    "senmon_3_desc": "Epoche, autori, falsi. La tua estetica coltivata — e il mercato, legale e non, che ci gira attorno.",
    "chi_sei": "Sei arrivata in polizia dopo una laurea in psicologia e anni di pratica clinica. Capisci le persone, vedi i pattern, sai perché la gente fa quello che fa. Sei anche una donna single in un ambiente dominato dagli uomini — hai dovuto lavorare il doppio per guadagnarti il rispetto. Cinque anni fa hai chiesto un prestito a Iwamoto, un conoscente di famiglia, per le cure di tuo padre. Tuo padre è morto comunque. Il debito è quasi saldato. Ma Iwamoto non sembra volerlo chiudere.",
    "problema_titolo": "IL CREDITORE",
    "problema_testo": "Iwamoto Koji continua a presentarsi. Non c'è fretta. Siamo amici. Quando sarà il momento, ne parleremo. Non chiede mai nulla di specifico — questo è il punto. L'ambiguità è l'arma. Vuole tenerti in debito non per i soldi, ma per il potere di poter chiedere qualcosa un giorno. E se qualcuno in centrale lo vede parlare con te, ci saranno domande.",
    "png_nome": "IWAMOTO Koji",
    "png_eta": "55 anni",
    "png_occupazione": "Uomo d'affari — prestiti privati, investimenti grigi",
    "png_relazione": "Creditore",
    "png_vuole": "Controllo. Il potere di poter chiedere un favore quando vuole.",
    "png_desc": "Gentilissimo. Premuroso. Mai una parola su cosa potrebbe chiedere. L'ambiguità è l'arma.",
    "conoscenza_nome": "MORITA Akiko",
    "conoscenza_ruolo": "Medico legale (ex collega)",
    "conoscenza_desc": "Amiche dai tempi dell'università, vi rispettate professionalmente. Autopsie, consulenze mediche, accesso ai dati sanitari (con cautela).",
    "conoscenza_contatto": "Ospedale Universitario, reparto Medicina Legale. O al cellulare personale.",
    "conoscenza_costo": "Caffè e chiacchierate, ogni tanto una consulenza psicologica informale.",
    "conoscenza_limite": "Una volta per sessione senza conseguenze. Siete le uniche due donne in posizioni di responsabilità che conoscete nel vostro campo.",
    "tatemae": "Sei professionale, competente, riservata. Non parli della tua vita privata. Fai il tuo lavoro e lo fai bene.",
    "honne": "Sei stufa di dover dimostrare il doppio. Sei stufa di Iwamoto. Sei stufa di essere sempre all'erta. Ma non lo mostrerai mai.",
    "frase": "Non guardare cosa ha fatto. Guarda perché l'ha fatto.",
    "pressione": "Ti chiudi ancora di più. Diventi fredda, quasi clinica. Analizzi tutto, anche le tue emozioni — come se fossero di qualcun altro.",
    "debolezza": "Quando un uomo la sminuisce o le dice che esagera. Anni di lotta per il rispetto condensati in un istante.",
    "vizio": "Whisky Nikka. Un dito, da sola, la sera dopo i casi pesanti. Non in compagnia.",
    "tic": "Inclina leggermente la testa a sinistra quando ascolta qualcuno mentire.",
    "oggetto": "Una penna stilografica Pilot nera. Scrive i profili sempre a mano, mai al computer.",
    "gusto": "Wagashi e tè matcha. Ha un debole per la pasticceria tradizionale vicino a Kiyomizu.",
    "rituale": "Il tempio di Nanzen-ji, la sera. Si siede nel giardino di pietra e non pensa a nulla.",
    "rapporto_1_nome": "YAMAMOTO Kenji",
    "rapporto_1_testo": "Un buon capo, ma si sta consumando. Non è il mio ruolo dirglielo, ma qualcuno dovrebbe.",
    "rapporto_2_nome": "HONDA Ryota",
    "rapporto_2_testo": "Nasconde qualcosa. Non so cosa, ma il linguaggio del corpo non mente.",
    "rapporto_3_nome": "NAKAMURA Shota",
    "rapporto_3_testo": "L'unico che non mi ha mai chiesto perché non sono sposata. Per questo lo rispetto.",
    "rapporto_4_nome": "SATO Yuki",
    "rapporto_4_testo": "Giovane, entusiasta, fragile. Mi ricorda una versione di me che non esiste più.",
}


PG_LIST = [
    ("PG_01", "Yamamoto_Kenji", PG_01),
    ("PG_02", "Honda_Ryota", PG_02),
    ("PG_03", "Nakamura_Shota", PG_03),
    ("PG_04", "Sato_Yuki", PG_04),
    ("PG_05", "Fujita_Emi", PG_05),
]

if __name__ == "__main__":
    for code, name, data in PG_LIST:
        img = os.path.join(BASE_DIR, "Immagini", f"{code}_{name}_foto.png")
        out = f"ANTEPRIMA_{name}.docx"
        genera_scheda(data, out, img)
    print("Fatto.")
