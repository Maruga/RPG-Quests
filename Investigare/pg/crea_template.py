"""
Genera il template DOCX per le schede PG di GENKAI v1.3
Parametro opzionale: percorso immagine per la foto del PG.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, sys

# Colori
GOLD = RGBColor(0xB4, 0x96, 0x50)
NAVY = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x78, 0x78, 0x78)
DARK_RED = RGBColor(0x8B, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CREAM_HEX = "F5F3EB"
INDENT = Emu(107950)
GOU_INDENT = Emu(215900)

# Parametro immagine opzionale
IMAGE_PATH = sys.argv[1] if len(sys.argv) > 1 else None

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(0.5)
section.bottom_margin = Cm(0.5)
section.left_margin = Cm(1.0)
section.right_margin = Cm(1.0)


def add_run(para, text, size=10, bold=False, italic=False, color=NAVY):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def remove_all_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/><w:left w:val="none"/>'
        f'  <w:right w:val="none"/><w:bottom w:val="none"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_valign(cell, val="center"):
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    cell._tc.get_or_add_tcPr().append(vAlign)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcW')):
        tcPr.remove(ex)
    tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>'))


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "━" * 70, size=6, color=GOLD)


def add_section_header(doc, kanji, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, kanji, size=13, color=GOLD)
    add_run(p, "  ", size=13)
    add_run(p, title, size=11, bold=True, color=NAVY)


def add_text(doc, text, size=10, color=NAVY, italic=False, bold=False, indent=True, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = INDENT
    add_run(p, text, size=size, color=color, italic=italic, bold=bold)
    return p


def add_field(doc, label, value, label_color=GRAY, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = INDENT
    add_run(p, f"{label}: ", size=size, bold=True, color=label_color)
    add_run(p, value, size=size, color=NAVY)


def add_gou(doc, name, desc, attributo, costo, successo, fallimento):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = INDENT
    add_run(p, "▸ ", size=10, color=GOLD)
    add_run(p, name, size=10, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, desc, size=9, italic=True, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, "Attributo: ", size=9, bold=True, color=GRAY)
    add_run(p, attributo, size=9, color=GRAY)
    add_run(p, "    ", size=9)
    add_run(p, "Costo: ", size=9, bold=True, color=GRAY)
    add_run(p, costo, size=9, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, "Successo: ", size=9, bold=True, color=GRAY)
    add_run(p, successo, size=9, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = GOU_INDENT
    add_run(p, "Fallimento: ", size=9, bold=True, color=GRAY)
    add_run(p, fallimento, size=9, color=NAVY)


# ============================================================
# PAGINA 1
# ============================================================

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
add_run(p, "限界  ·  GENKAI v1.3  ·  限界", size=8, color=GOLD)

# NOME + FOTO
name_tbl = doc.add_table(rows=1, cols=2)
name_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

c_name = name_tbl.rows[0].cells[0]
set_cell_width(c_name, 15)
set_cell_valign(c_name, "center")
remove_all_borders(c_name)
p = c_name.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
add_run(p, "{{nome_kanji}}", size=26, color=NAVY)
p = c_name.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "{{nome_romanizzato}}", size=14, bold=True, color=NAVY)

c_foto = name_tbl.rows[0].cells[1]
set_cell_width(c_foto, 4)
set_cell_valign(c_foto, "center")
remove_all_borders(c_foto)
p_foto = c_foto.paragraphs[0]
p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Inserisci foto se disponibile
if IMAGE_PATH and os.path.exists(IMAGE_PATH):
    run = p_foto.add_run()
    run.add_picture(IMAGE_PATH, width=Cm(3))
else:
    add_run(p_foto, "{{foto}}", size=9, italic=True, color=GRAY)

# Info
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
add_run(p, "Età ", size=10, bold=True, color=GRAY)
add_run(p, "{{eta}}", size=10, color=NAVY)
add_run(p, "  ·  ", size=10, color=GOLD)
add_run(p, "Ruolo ", size=10, bold=True, color=GRAY)
add_run(p, "{{ruolo}}", size=10, color=NAVY)
add_run(p, "  ·  ", size=10, color=GOLD)
add_run(p, "Grado ", size=10, bold=True, color=GRAY)
add_run(p, "{{grado}}", size=10, color=NAVY)
add_run(p, "  ·  ", size=10, color=GOLD)
add_run(p, "Servizio ", size=10, bold=True, color=GRAY)
add_run(p, "{{servizio}}", size=10, color=NAVY)

add_separator(doc)

# ATTRIBUTI
add_section_header(doc, "能力", "ATTRIBUTI")

# Tabella compatta: 3 colonne
table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

for i, h in enumerate(["Attributo", "Base", "Finale"]):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A1A2E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, h, size=9, bold=True, color=WHITE)

attrs = [
    ("Distacco 冷静", "{{distacco_base}}"),
    ("Pazienza 忍耐", "{{pazienza_base}}"),
    ("Silenzio 沈黙", "{{silenzio_base}}"),
    ("Lucidità 明晰", "{{lucidita_base}}"),
    ("Ascolto 傾聴", "{{ascolto_base}}"),
    ("Presenza 存在", "{{presenza_base}}"),
]
for i, (name, base) in enumerate(attrs):
    row = table.rows[i + 1]
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, CREAM_HEX)
    p = row.cells[0].paragraphs[0]
    add_run(p, name, size=10, bold=True, color=NAVY)
    p = row.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, base, size=12, bold=True, color=NAVY)
    p = row.cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "", size=12, color=GRAY)

for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/>'
            f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="D0C8B0"/>'
            f'</w:tcBorders>'))

# Note uso sotto tabella
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = INDENT
add_run(p, "5 punti da distribuire ", size=9, bold=True, color=GRAY)
add_run(p, "(max 2 per attributo, max 8)", size=9, color=GRAY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = INDENT
add_run(p, "Distacco: ", size=8, bold=True, color=GRAY)
add_run(p, "scene brutali, non farti coinvolgere", size=8, italic=True, color=GRAY)
add_run(p, "  ·  ", size=8, color=GOLD)
add_run(p, "Pazienza: ", size=8, bold=True, color=GRAY)
add_run(p, "attese, interrogatori lunghi", size=8, italic=True, color=GRAY)
add_run(p, "  ·  ", size=8, color=GOLD)
add_run(p, "Silenzio: ", size=8, bold=True, color=GRAY)
add_run(p, "incassare, non reagire", size=8, italic=True, color=GRAY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = INDENT
add_run(p, "Lucidità: ", size=8, bold=True, color=GRAY)
add_run(p, "analizzare, ricostruire", size=8, italic=True, color=GRAY)
add_run(p, "  ·  ", size=8, color=GOLD)
add_run(p, "Ascolto: ", size=8, bold=True, color=GRAY)
add_run(p, "cogliere bugie, far parlare", size=8, italic=True, color=GRAY)
add_run(p, "  ·  ", size=8, color=GOLD)
add_run(p, "Presenza: ", size=8, bold=True, color=GRAY)
add_run(p, "fare pressione, autorità", size=8, italic=True, color=GRAY)

add_separator(doc)

# KI
add_section_header(doc, "気", "KI — TENUTA")

add_text(doc, "Attributo Finale più basso + 2d6 (prendi il dado più alto). Se esce 1, ritira.", size=9, color=GRAY, italic=True, after=4)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = INDENT
add_run(p, "Ki Max: ________", size=12, bold=True, color=NAVY)
add_run(p, "     ≤3 Genkai", size=9, italic=True, color=DARK_RED)

# Ki attuale + Satori sulla stessa riga
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = INDENT
add_run(p, "Ki attuale: ________", size=12, bold=True, color=NAVY)
add_run(p, "          ", size=10)
add_run(p, "悟り", size=12, color=GOLD)
add_run(p, "  SATORI — 1/sessione, successo automatico  ", size=9, color=NAVY)
add_run(p, "☐", size=11, color=NAVY)

add_separator(doc)

# NASAKE + SOROBAN AFFIANCATI
ns_tbl = doc.add_table(rows=1, cols=2)
ns_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# NASAKE
c_nas = ns_tbl.rows[0].cells[0]
set_cell_width(c_nas, 9.5)
set_cell_valign(c_nas, "top")
remove_all_borders(c_nas)

p = c_nas.paragraphs[0]
p.paragraph_format.space_after = Pt(3)
add_run(p, "情け", size=12, color=GOLD)
add_run(p, "  NASAKE — Compassione", size=10, bold=True, color=NAVY)

p = c_nas.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p, "情けは人の為ならず", size=9, italic=True, color=GOLD)

p = c_nas.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p, "Bonus Kiwami+/Nami+ perso → conserva 1 Ki.\nSolo donabile a un altro PG giocando la scena.\nSi perde a fine sessione.", size=9, color=GRAY)

p = c_nas.add_paragraph()
p.paragraph_format.space_before = Pt(2)
add_run(p, "Nasake:  ☐ Vuoto   ☐ Pieno", size=10, bold=True, color=NAVY)

# SOROBAN
c_sor = ns_tbl.rows[0].cells[1]
set_cell_width(c_sor, 9.5)
set_cell_valign(c_sor, "top")
remove_all_borders(c_sor)

p = c_sor.paragraphs[0]
p.paragraph_format.space_after = Pt(3)
add_run(p, "算盤", size=12, color=GOLD)
add_run(p, "  SOROBAN — Giornata", size=10, bold=True, color=NAVY)

p = c_sor.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p, "Nami+ → +1    Kiwami+ → +2\nNami- → -1     Kiwami- → -2", size=9, color=GRAY)

p = c_sor.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p, "≥ partenza → 2d6 migliore (reroll 1)\n< partenza → 2d6 peggiore (reroll 1)", size=9, color=GRAY)

p = c_sor.add_paragraph()
p.paragraph_format.space_before = Pt(2)
add_run(p, "0  1  2  3  4  ", size=10, color=GRAY)
add_run(p, "⑤", size=12, bold=True, color=GOLD)
add_run(p, "  6  7  8  9", size=10, color=GRAY)

add_separator(doc)

# GOU
add_section_header(doc, "業", "GOU — IL DEBITO  (scegli uno)")

add_text(doc, "Il Gou funziona SEMPRE. Successo = preciso, Fallimento = vago. Ogni uso raddoppia il costo del successivo; una notte di sonno lo riabbassa di un grado.", size=9, color=GRAY, italic=True, after=2)

add_gou(doc, "{{gou_1_nome}}", "{{gou_1_desc}}",
        "{{gou_1_attributo}}", "{{gou_1_costo}}",
        "{{gou_1_successo}}", "{{gou_1_fallimento}}")

add_gou(doc, "{{gou_2_nome}}", "{{gou_2_desc}}",
        "{{gou_2_attributo}}", "{{gou_2_costo}}",
        "{{gou_2_successo}}", "{{gou_2_fallimento}}")

add_gou(doc, "{{gou_3_nome}}", "{{gou_3_desc}}",
        "{{gou_3_attributo}}", "{{gou_3_costo}}",
        "{{gou_3_successo}}", "{{gou_3_fallimento}}")

add_separator(doc)

# Gou scelto
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Gou scelto: ________________________________", size=10, color=GRAY)


# ============================================================
# PAGE BREAK
# ============================================================
doc.add_page_break()


# ============================================================
# PAGINA 2
# ============================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
add_run(p, "{{nome_kanji}}", size=14, color=GOLD)
add_run(p, "  ·  ", size=12, color=GOLD)
add_run(p, "{{nome_romanizzato}}", size=12, bold=True, color=NAVY)

add_separator(doc)

# CHI SEI
add_section_header(doc, "人", "CHI SEI")
add_text(doc, "{{chi_sei}}")

add_separator(doc)

# IL TUO PROBLEMA
add_section_header(doc, "影", "IL TUO PROBLEMA — {{problema_titolo}}")
add_text(doc, "{{problema_testo}}")

add_separator(doc)

# PNG DEL PROBLEMA
add_section_header(doc, "縁", "{{png_nome}}")

table2 = doc.add_table(rows=4, cols=2)
table2.autofit = True
for i, (label, val) in enumerate([
    ("Età", "{{png_eta}}"), ("Occupazione", "{{png_occupazione}}"),
    ("Relazione", "{{png_relazione}}"), ("Cosa vuole", "{{png_vuole}}"),
]):
    p = table2.rows[i].cells[0].paragraphs[0]
    add_run(p, label, size=9, bold=True, color=GRAY)
    p = table2.rows[i].cells[1].paragraphs[0]
    add_run(p, val, size=10, color=NAVY)
    for cell in table2.rows[i].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/>'
            f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="D0C8B0"/>'
            f'</w:tcBorders>'))

add_text(doc, "{{png_desc}}", italic=True, after=2)

add_separator(doc)

# CONOSCENZA
add_section_header(doc, "縁者", "CONOSCENZA — {{conoscenza_nome}}")

add_text(doc, "{{conoscenza_ruolo}}", italic=True, color=GRAY, after=2)
add_text(doc, "{{conoscenza_desc}}", after=2)
add_field(doc, "Contatto", "{{conoscenza_contatto}}", size=9)
add_field(doc, "Costo", "{{conoscenza_costo}}", size=9)
add_text(doc, "{{conoscenza_limite}}", size=8, italic=True, color=GRAY, after=2)

add_separator(doc)

# COME TI COMPORTI
add_section_header(doc, "面", "COME TI COMPORTI")

for kanji, label, field, lbl_color in [
    ("建前", "In pubblico", "{{tatemae}}", GOLD),
    ("本音", "In privato", "{{honne}}", GOLD),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = INDENT
    add_run(p, f"{kanji}  {label}: ", size=10, bold=True, color=lbl_color)
    add_run(p, field, size=10, color=NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
p.paragraph_format.left_indent = INDENT
add_run(p, "Frase tipica: ", size=10, bold=True, color=GRAY)
add_run(p, "{{frase}}", size=10, italic=True, color=NAVY)

for label, field in [("Sotto pressione", "{{pressione}}"), ("Debolezza", "{{debolezza}}")]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = INDENT
    add_run(p, f"{label}: ", size=10, bold=True, color=DARK_RED)
    add_run(p, field, size=10, color=NAVY)

add_separator(doc)

# TRATTI PERSONALI
add_section_header(doc, "癖", "TRATTI PERSONALI")

for label, val in [("Vizio", "{{vizio}}"), ("Tic", "{{tic}}"), ("Oggetto", "{{oggetto}}"),
                   ("Gusto", "{{gusto}}"), ("Rituale", "{{rituale}}")]:
    add_field(doc, label, val, size=10)

add_separator(doc)

# LA SQUADRA
add_section_header(doc, "仲", "LA SQUADRA")

for i in range(1, 5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = INDENT
    add_run(p, f"{{{{rapporto_{i}_nome}}}}: ", size=10, bold=True, color=GRAY)
    add_run(p, f"{{{{rapporto_{i}_testo}}}}", size=10, italic=True, color=NAVY)


# Salva
output_path = os.path.join(os.path.dirname(__file__), "TEMPLATE_Scheda_PG.docx")
doc.save(output_path)
print(f"Template salvato: {output_path}")
