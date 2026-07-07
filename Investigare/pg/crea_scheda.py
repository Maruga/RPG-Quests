"""
Genera schede PG in stile giapponese elegante per GENKAI 限界.
Pagina 1: Statistiche, Ki, Gou
Pagina 2: Background, problema, PNG, conoscenza, comportamento
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

OUT_DIR = Path(__file__).parent

# ── Colors ──
DARK = RGBColor(26, 26, 46)       # #1A1A2E
GOLD = RGBColor(180, 150, 80)     # #B49650
MUTED = RGBColor(120, 120, 120)   # #787878
RED = RGBColor(139, 0, 0)         # #8B0000
WHITE = RGBColor(255, 255, 255)
CREAM_HEX = "FAF8F0"
DARK_HEX = "1A1A2E"
GOLD_HEX = "B49650"
LIGHT_HEX = "F5F3EB"

# ── Helpers ──

def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each param is (size_eighths, color_hex) or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            sz, color = val
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{sz}" '
                f'w:space="0" w:color="{color}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_borders(table, color="1A1A2E", size=4):
    """Set outer borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_run(paragraph, text, bold=False, italic=False, size=10, color=DARK, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run

def add_para(doc, text="", bold=False, italic=False, size=10, color=DARK,
             align=None, space_after=4, space_before=0, font="Calibri"):
    p = doc.add_paragraph()
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color, font=font)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_separator(doc, char="━", length=50, color=GOLD):
    add_para(doc, char * length, size=7, color=color,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, space_before=6)

def add_thin_separator(doc, color=GOLD):
    add_para(doc, "─" * 65, size=5, color=color,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=4)

def add_section_header(doc, text, kanji=""):
    """Elegant section header with optional kanji."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(2)
    if kanji:
        add_run(p, f"{kanji}  ", size=11, color=GOLD, font="Calibri")
    add_run(p, text.upper(), bold=True, size=10, color=DARK, font="Calibri")
    # underline
    add_para(doc, "─" * 55, size=5, color=GOLD,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, space_before=0)


def add_gou_box(doc, name, kanji_desc, attribute, cost, success, failure):
    """Add a styled Gou option box."""
    # Title line
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(7)
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.3)
    add_run(p, "▸ ", size=9, color=GOLD)
    add_run(p, name, bold=True, size=10, color=DARK)

    # Description
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.left_indent = Cm(0.6)
    add_run(p2, kanji_desc, italic=True, size=9, color=MUTED)

    # Mechanics in compact form
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(1)
    p3.paragraph_format.left_indent = Cm(0.6)
    add_run(p3, "Attributo: ", bold=True, size=8, color=MUTED)
    add_run(p3, attribute, size=8, color=DARK)
    add_run(p3, "  ·  ", size=8, color=GOLD)
    add_run(p3, "Costo: ", bold=True, size=8, color=MUTED)
    add_run(p3, cost, size=8, color=RED)

    # Success
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(1)
    p4.paragraph_format.left_indent = Cm(0.6)
    add_run(p4, "Successo: ", bold=True, size=8, color=MUTED)
    add_run(p4, success, size=8, color=DARK)

    # Failure
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(4)
    p5.paragraph_format.left_indent = Cm(0.6)
    add_run(p5, "Fallimento: ", bold=True, size=8, color=MUTED)
    add_run(p5, failure, size=8, color=DARK)


# ── PG Data ──

PG_DATA = [
    {
        "filename": "Scheda_PG_01_Yamamoto_Kenji.docx",
        "photo": "PG_01_Yamamoto_Kenji_foto.png",
        "name_jp": "山本 健二",
        "name": "YAMAMOTO Kenji",
        "age": "42",
        "role": "Capo Squadra Investigativa",
        "rank": "Ispettore Capo (Keibu 警部)",
        "years": "18",
        "attrs": [
            ("Distacco", "6", "Veterano — sai distaccarti"),
            ("Pazienza", "4", "Normale"),
            ("Silenzio", "6", "Incassi e chiudi tutto dentro"),
            ("Lucidità", "6", "Punto di forza — vedi i collegamenti"),
            ("Ascolto", "4", "Normale"),
            ("Presenza", "5", "Capo rispettato"),
        ],
        "gou": [
            {
                "name": "Teatro delle Ombre",
                "desc": "Guardando una scena, visualizzi come si sono svolti gli eventi.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Vedi la sequenza completa — chi ha fatto cosa, in che ordine, con quali movimenti",
                "failure": "Vedi frammenti disconnessi — immagini, gesti, ma non la sequenza chiara",
            },
            {
                "name": "Pietra Fuori Posto",
                "desc": "Percepisci immediatamente cosa non torna in un luogo o situazione.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Sai cosa stona e perché — il dettaglio e il suo significato",
                "failure": "Senti che qualcosa è sbagliato, ma non riesci a definirlo con precisione",
            },
            {
                "name": "Pugno di Ferro",
                "desc": "La tua autorità diventa schiacciante. Il PNG sente il peso della tua posizione.",
                "attr": "Presenza", "cost": "2 Ki",
                "success": "Il PNG cede su un punto importante — confessa, collabora, si ritira",
                "failure": "Il PNG è scosso, intimidito, ma non cede completamente",
            },
        ],
        "background": "Sei un veterano rispettato. Hai risolto casi importanti, hai costruito una reputazione solida. Ma la tua vita privata è un disastro.\n\nSei mesi fa tua moglie Yuko ti ha lasciato. \"Non ci sei mai,\" ha detto. Aveva ragione. Ora vivi in un piccolo appartamento a Fushimi, con le scatole ancora da disfare.\n\nTuo figlio Takeshi (8 anni) vive con la madre, ma hai l'affidamento condiviso. Il martedì e il giovedì Takeshi è con te. O dovrebbe esserlo.\n\nIl problema: il lavoro viene sempre prima. Le promesse non mantenute. Le cene saltate.",
        "problem_title": "La Famiglia",
        "problem_text": "Tua sorella Yamamoto Noriko è stata il tuo supporto durante il divorzio. Ha coperto le tue assenze, ha badato a Takeshi quando tu non potevi. Ma la pazienza di Noriko ha un limite.\n\nL'ultima volta che le hai chiesto aiuto, te lo ha detto chiaramente: \"Non sono la tua baby-sitter di riserva. Devi scegliere cosa conta davvero.\"\n\nSai che Noriko è stanca. Sai che il rapporto con Takeshi è fragile. E sai che ogni richiesta d'aiuto pesa sulla relazione con tua sorella.",
        "png_name": "YAMAMOTO Noriko",
        "png_role": "Sorella minore",
        "png_age": "38",
        "png_job": "Insegnante di ikebana",
        "png_wants": "Che tu ammetta di avere un problema. Non singoli episodi — il pattern.",
        "png_behavior": "Non urla. È peggio. Parla con calma, con quella delusione quieta che ferisce più della rabbia.",
        "enja_name": "TANAKA Shuichi",
        "enja_role": "Giornalista al Kyoto Shimbun",
        "enja_desc": "Vecchio compagno di università. Accesso agli archivi del giornale, voci di corridoio, contatti nella stampa.",
        "enja_contact": "Cellulare personale. Risponde sempre, anche di notte.",
        "enja_cost": "Ogni tanto una soffiata, niente di compromettente.",
        "tatemae": "Sei il capo. Calmo, competente, rispettato. Non mostri mai debolezza davanti alla squadra.",
        "honne": "Sei esausto. Ti senti in colpa per Takeshi. A volte ti chiedi se ne vale la pena.",
        "phrase": "\"Abbiamo un lavoro da fare. Concentriamoci su quello.\"",
        "pressure": "Ti chiudi. Parli meno. Diventi più brusco. Non deleghi — fai tutto tu perché \"è più veloce.\"",
        "weakness": "Quando qualcuno mette in dubbio le sue priorità tra lavoro e famiglia. Soprattutto se ha ragione.",
        "squad": [
            ("HONDA Ryota", "Brillante, ma gioca col fuoco. Un giorno si brucerà e non ci sarà nessuno a salvarlo."),
            ("NAKAMURA Shota", "Il migliore di noi negli interrogatori. Mi fido di lui più di chiunque altro."),
            ("SATO Yuki", "Mi ricorda me a quell'età. Spero che non faccia i miei stessi errori."),
            ("FUJITA Emi", "Competente, riservata, non chiede mai nulla. Questo mi preoccupa."),
        ],
        "traits": {
            "vizio": "Caffè nero, almeno sei tazze al giorno. Non tocca l'alcol da quando Yuko se n'è andata.",
            "tic": "Si massaggia il ponte del naso quando è stanco o frustrato.",
            "oggetto": "Un orologio Seiko automatico, regalo di suo padre per la promozione a ispettore.",
            "gusto": "Udon caldo da un chiosco vicino alla centrale. Sempre lo stesso, sempre in piedi.",
            "extra": ("Rituale", "Ogni mattina, prima di entrare in centrale, caffè Boss in lattina al distributore. Sempre lo stesso, sempre in piedi."),
        },
    },
    # ── PG_02 — HONDA Ryota ──
    {
        "filename": "Scheda_PG_02_Honda_Ryota.docx",
        "photo": "PG_02_Honda_Ryota_foto.png",
        "name_jp": "本田 涼太",
        "name": "HONDA Ryota",
        "age": "35",
        "role": "Analista Scene del Crimine",
        "rank": "Sergente (Junsa-bucho)",
        "years": "10",
        "attrs": [
            ("Distacco", "6", "Freddo sulla scena"),
            ("Pazienza", "4", "Impulsivo, cerchi la scorciatoia"),
            ("Silenzio", "4", "Normale"),
            ("Lucidità", "7", "Brillante — noti ogni dettaglio"),
            ("Ascolto", "4", "Normale"),
            ("Presenza", "6", "Sicuro di te, a volte troppo"),
        ],
        "gou": [
            {
                "name": "Occhio della Gru",
                "desc": "Vedi ciò che altri non vedono — il dettaglio che cambia tutto.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Noti il dettaglio nascosto e ne capisci immediatamente il significato",
                "failure": "Vedi qualcosa di anomalo, ma non riesci a capire cosa significa",
            },
            {
                "name": "Teatro delle Ombre",
                "desc": "Guardando una scena, visualizzi come si sono svolti gli eventi.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Vedi la sequenza completa degli eventi",
                "failure": "Vedi frammenti disconnessi, immagini sfocate",
            },
            {
                "name": "Cuore di Ghiaccio",
                "desc": "Ti distacchi completamente dalla situazione. Niente ti tocca.",
                "attr": "Distacco", "cost": "2 Ki",
                "success": "Per il resto della scena, agisci normalmente senza bisogno di tiri emotivi",
                "failure": "Resti lucido ma la tensione si accumula — il prossimo tiro emotivo ha -2",
            },
        ],
        "background": "Sei brillante. Lo sanno tutti, tu per primo. Vedi cose che altri non vedono. Ricostruisci scene del crimine come se avessi assistito in persona.\n\nIl problema è che la tua mente non si ferma mai. Cerchi stimoli. Adrenalina. Sfide.\n\nHai iniziato con le scommesse sportive tre anni fa. \"Solo per divertimento.\" Poi le cifre sono cresciute. Ora scommetti su tutto: baseball, sumo, corse di cavalli. Vinci spesso — sei intelligente. Ma quando perdi, perdi grosso.\n\nE ultimamente hai perso più di quanto hai vinto.",
        "problem_title": "Il Gioco d'Azzardo",
        "problem_text": "Devi soldi a Murakami, un allibratore indipendente. Non è yakuza — è quasi rispettabile. Ma i debiti sono debiti.\n\nMurakami ha un uomo che si occupa delle riscossioni: Goto Masaru. Goto non minaccia mai. È educato, ragionevole, professionale. Il che lo rende più inquietante.\n\nSai che hai debiti di gioco con persone che non dimenticano. Sai che Goto potrebbe presentarsi in qualsiasi momento. E sai che se qualcuno in centrale scopre che scommetti illegalmente, la tua carriera è finita.",
        "png_name": "GOTO Masaru",
        "png_role": "Riscossore per Murakami",
        "png_age": "45",
        "png_job": "\"Consulente finanziario\" (riscossore)",
        "png_wants": "I soldi, preferibilmente. Ma è pragmatico — potrebbe accettare alternative.",
        "png_behavior": "Parla di \"soluzioni\", \"accordi\", \"comprensione reciproca\". Sorride sempre. Le minacce sono solo sottintese.",
        "enja_name": "ODA Takumi",
        "enja_role": "Informatore",
        "enja_desc": "Ex testimone di un vecchio caso, ora piccolo ricettatore. Conosce il sottobosco di Kyoto, sente voci, sa chi vende cosa a chi.",
        "enja_contact": "Bazzica al Bar Shinjuku a Gion, tutte le sere dopo le 21.",
        "enja_cost": "Che tu continui a chiudere un occhio sui suoi affari minori.",
        "tatemae": "Sei il tecnico sicuro di sé. Fai battute, smorzi la tensione. Sembri sempre sotto controllo.",
        "honne": "L'adrenalina è una droga. Quando non c'è, ti senti vuoto. Le scommesse riempiono quel vuoto. E adesso il vuoto sta per inghiottirti.",
        "phrase": "\"Aspetta. Guarda qui. Vedi questo? Nessuno l'aveva notato.\"",
        "pressure": "Diventi più arrogante. Prendi rischi. Scommetti — letteralmente e metaforicamente. A volte funziona. A volte no.",
        "weakness": "Quando qualcuno lo chiama \"fortunato\" invece di riconoscere il suo talento. O quando i debiti vengono anche solo sfiorati.",
        "squad": [
            ("YAMAMOTO Kenji", "Il capo. Duro, giusto, ma ultimamente lo vedo stanco. Non glielo dirò mai."),
            ("NAKAMURA Shota", "Troppo lento per i miei gusti, ma quando trova qualcosa è sempre la cosa giusta."),
            ("SATO Yuki", "Il ragazzino del laboratorio. Bravo, ma deve smettere di chiedere il permesso per tutto."),
            ("FUJITA Emi", "Mi legge come un libro aperto. Questo mi innervosisce."),
        ],
        "traits": {
            "vizio": "Mild Seven. Fuma solo sulla scena del crimine, dice che lo aiuta a pensare.",
            "tic": "Fa roteare una moneta da 500 yen tra le dita quando ragiona.",
            "oggetto": "Una lente d'ingrandimento pieghevole nel taschino — vecchia, graffiata, non la cambierebbe mai.",
            "gusto": "Yakitori e birra Asahi al bancone. Mai al tavolo, sempre al bancone.",
            "extra": ("Superstizione", "Non inizia mai un sopralluogo dal lato destro della scena. Dice che porta sfortuna. Non sa spiegare perché."),
        },
    },
    # ── PG_03 — NAKAMURA Shota ──
    {
        "filename": "Scheda_PG_03_Nakamura_Shota.docx",
        "photo": "PG_03_Nakamura_Shota_foto.png",
        "name_jp": "中村 翔太",
        "name": "NAKAMURA Shota",
        "age": "38",
        "role": "Specialista Interrogatori",
        "rank": "Ispettore (Keibu-ho)",
        "years": "14",
        "attrs": [
            ("Distacco", "4", "Normale"),
            ("Pazienza", "7", "Punto di forza — puoi aspettare ore"),
            ("Silenzio", "5", "Sai quando tacere"),
            ("Lucidità", "4", "Normale"),
            ("Ascolto", "7", "Eccellente — senti le bugie"),
            ("Presenza", "4", "Non intimidisci, convinci"),
        ],
        "gou": [
            {
                "name": "Ombra della Verità",
                "desc": "Senti quando qualcuno mente. Non sai come, ma lo senti.",
                "attr": "Ascolto", "cost": "2 Ki",
                "success": "Sai che mente e su cosa sta mentendo",
                "failure": "Sai che mente, ma non riesci a capire su quale parte",
            },
            {
                "name": "Porta Socchiusa",
                "desc": "Le persone si aprono con te. Dicono più di quanto vorrebbero.",
                "attr": "Ascolto", "cost": "2 Ki",
                "success": "Il PNG rivela qualcosa che non voleva assolutamente dire",
                "failure": "Il PNG lascia trapelare qualcosa, ma si ferma prima di dire troppo",
            },
            {
                "name": "L'Ora Giusta",
                "desc": "Sai quando è il momento perfetto per parlare, agire, colpire.",
                "attr": "Pazienza", "cost": "2 Ki",
                "success": "La prossima azione è un successo automatico",
                "failure": "Hai +2 al prossimo tiro (senti che il momento è vicino, ma non perfetto)",
            },
        ],
        "background": "Sei quello che fa parlare le persone. Non con la forza — con la pazienza. Ti siedi, aspetti, ascolti. Prima o poi tutti parlano.\n\nSei il figlio maggiore di una famiglia modesta di Osaka. Hai lavorato duro per arrivare dove sei. Tua madre ne è fiera. Tuo padre è morto quando avevi 15 anni.\n\nTuo fratello minore Kazuo (34 anni) è l'opposto. Non ha mai tenuto un lavoro più di sei mesi. Vive di espedienti, piccoli imbrogli, soldi prestati e mai restituiti. Ogni tanto sparisce per mesi. Poi torna, con una nuova idea, un nuovo progetto, una nuova richiesta.",
        "problem_title": "Il Fratello",
        "problem_text": "Kazuo è un peso che ti porti dietro da sempre. Gli hai già prestato soldi — tre volte. Ogni volta \"l'ultima\". Ogni volta una bugia.\n\nTua moglie non sa di questi prestiti. Se lo scoprisse, ci sarebbero problemi anche a casa.\n\nMa Kazuo è sangue del tuo sangue. E quando ha bisogno, sa sempre dove trovare suo fratello.",
        "png_name": "NAKAMURA Kazuo",
        "png_role": "Fratello minore",
        "png_age": "34",
        "png_job": "Nessuno fisso, \"imprenditore\"",
        "png_wants": "Dipende dalla situazione. Ma di solito sono soldi, e di solito sono urgenti.",
        "png_behavior": "Prima la simpatia. Poi il senso di colpa. Poi la minaccia velata. Conosce i tuoi punti deboli.",
        "enja_name": "NAKAMURA Hideki",
        "enja_role": "Avvocato (Cugino)",
        "enja_desc": "Cugino di primo grado, avvocato penalista a Kyoto. Consulenza legale, accesso a fascicoli pubblici, contatti in tribunale.",
        "enja_contact": "Studio legale Nakamura & Associati, Via Kawaramachi. O al cellulare.",
        "enja_cost": "Niente di specifico — famiglia è famiglia.",
        "tatemae": "Sei quello calmo. Non ti agiti mai. Parli poco, ascolti molto. Quando parli, le persone tendono ad ascoltare.",
        "honne": "Sei stanco di essere quello responsabile. Di dover tenere insieme tutto. A volte vorresti mandare tutto al diavolo. Ma non lo farai mai.",
        "phrase": "\"Prendiamoci un momento. Raccontami tutto dall'inizio.\"",
        "pressure": "Diventi ancora più silenzioso. Ti ritiri in te stesso. Smetti di chiedere e inizi a osservare.",
        "weakness": "Quando qualcuno usa i legami di sangue come arma. Tocca troppo vicino a casa.",
        "squad": [
            ("YAMAMOTO Kenji", "Un capo che si porta il lavoro a casa e la casa al lavoro. Lo capisco fin troppo bene."),
            ("HONDA Ryota", "Geniale e imprudente. Lo tengo d'occhio, anche se lui non se ne accorge."),
            ("SATO Yuki", "Ha talento vero. Deve solo imparare che non tutti i problemi si risolvono col microscopio."),
            ("FUJITA Emi", "L'unica che capisce il silenzio come lo capisco io. Ci rispettiamo senza bisogno di parole."),
        ],
        "traits": {
            "vizio": "Tè verde. Ne beve litri durante gli interrogatori. Offre sempre una tazza all'interrogato.",
            "tic": "Annuisce lentamente anche quando non è d'accordo. Le persone lo trovano rassicurante.",
            "oggetto": "Un taccuino Moleskine nero, pieno di appunti in calligrafia minuscola. Non usa registratori.",
            "gusto": "Ramen miso con extra chashu. Conosce tutti i posti migliori di Kyoto a memoria.",
            "extra": ("Segreto", "Colleziona manga shoujo. Li tiene nascosti in un cassetto della scrivania. Se qualcuno li vedesse, morirebbe di vergogna."),
        },
    },
    # ── PG_04 — SATO Yuki ──
    {
        "filename": "Scheda_PG_04_Sato_Yuki.docx",
        "photo": "PG_04_Sato_Yuki_foto.png",
        "name_jp": "佐藤 勇気",
        "name": "SATO Yuki",
        "age": "27",
        "role": "Tecnico Analisi Tracce",
        "rank": "Agente Scelto (Junsa-cho)",
        "years": "3",
        "attrs": [
            ("Distacco", "4", "Ti fai coinvolgere troppo"),
            ("Pazienza", "6", "Metodico, preciso"),
            ("Silenzio", "4", "Tendi a parlare troppo quando sei nervoso"),
            ("Lucidità", "7", "Eccellente — mente analitica formidabile"),
            ("Ascolto", "5", "Attento ai dettagli"),
            ("Presenza", "5", "Entusiasta, energia contagiosa"),
        ],
        "gou": [
            {
                "name": "Palazzo della Memoria",
                "desc": "Puoi richiamare con precisione fotografica qualcosa che hai visto o sentito.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Ricordi il dettaglio e anche elementi periferici che non avevi notato consciamente",
                "failure": "Ricordi il dettaglio principale, ma sfocato o incompleto",
            },
            {
                "name": "Occhio della Gru",
                "desc": "Vedi ciò che altri non vedono — il dettaglio che cambia tutto.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Noti il dettaglio nascosto e ne capisci il significato",
                "failure": "Vedi qualcosa di anomalo, ma non riesci a interpretarlo",
            },
            {
                "name": "Pietra Fuori Posto",
                "desc": "Percepisci immediatamente cosa non torna — tracce, residui, anomalie.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Identifichi l'anomalia e la sua probabile origine",
                "failure": "Senti che qualcosa non quadra, ma non riesci a isolare cosa",
            },
        ],
        "background": "Sei il più giovane della squadra. Brillante, laureato con lode in chimica forense, pieno di entusiasmo. Sei quello che lavora fino a tardi, che controlla tre volte, che non si arrende finché non trova la risposta.\n\nIl problema è tua madre.\n\nSato Michiko non ha mai accettato che tu facessi il poliziotto. \"Uno spreco,\" dice. \"Con la tua laurea potevi lavorare in un'azienda vera. Guadagnare bene. Trovare una brava ragazza.\"\n\nTuo zio Tanaka Jiro ha un'azienda di import-export. Niente di entusiasmante, ma paga bene. Tua madre non smette di ripetere che \"lo zio ti aspetta\".",
        "problem_title": "La Madre",
        "problem_text": "Tua madre chiama. Spesso. A volte in centrale. \"È urgente,\" dice alla segretaria. Non è mai urgente.\n\nVuole sapere se hai ripensato al lavoro dello zio. Vuole sapere quando ti sposi. Vuole sapere perché \"sprechi la tua vita\" a fare il poliziotto.\n\nNon urla mai. Non minaccia. Usa il senso di colpa come un'arma di precisione.\n\nSai che non approva le tue scelte. Sai che ogni telefonata è una battaglia. E sai che non puoi riattaccare in faccia a tua madre — non in Giappone.",
        "png_name": "SATO Michiko",
        "png_role": "Madre",
        "png_age": "58",
        "png_job": "Casalinga, vedova da 10 anni",
        "png_wants": "Che tu \"sistemi la tua vita\". Matrimonio, lavoro stabile, nipotini.",
        "png_behavior": "Mai aggressiva direttamente. \"Io voglio solo il tuo bene.\" \"Tuo padre sarebbe così deluso.\" Il senso di colpa è la sua arma principale.",
        "enja_name": "KATO Hiroshi",
        "enja_role": "Tecnico Informatico",
        "enja_desc": "Amico d'infanzia, ora lavora in un'azienda tech. Recupero dati, analisi file, hacking leggero.",
        "enja_contact": "Telefono di casa. Risponde anche a tarda sera.",
        "enja_cost": "Storie dal lavoro, cene, compagnia.",
        "tatemae": "Sei il giovane promettente. Entusiasta, competente, sempre pronto ad aiutare.",
        "honne": "Sei esausto dal peso delle aspettative. Di tua madre, del lavoro, di te stesso. A volte ti chiedi se non avresti fatto meglio ad ascoltarla.",
        "phrase": "\"Datemi ancora un'ora, ho quasi finito l'analisi.\"",
        "pressure": "Parli troppo. Riempi i silenzi. Ti giustifichi anche quando non devi. Cerchi approvazione.",
        "weakness": "Quando qualcuno lo tratta da ragazzino o mette in dubbio la sua esperienza. Soprattutto se è un superiore.",
        "squad": [
            ("YAMAMOTO Kenji", "L'ispettore. Lo ammiro, ma a volte ho paura di deluderlo."),
            ("HONDA Ryota", "Il più figo della squadra. Vorrei avere la sua sicurezza. O almeno fingerla così bene."),
            ("NAKAMURA Shota", "Calmo come un lago. Quando parla con me, mi sento meno ansioso."),
            ("FUJITA Emi", "A volte mi guarda come se sapesse esattamente cosa sto pensando. Probabilmente lo sa."),
        ],
        "traits": {
            "vizio": "Gomme da masticare alla menta. Ne consuma un pacchetto al giorno in laboratorio.",
            "tic": "Si sistema gli occhiali spingendoli sul naso con l'indice, anche quando non scivolano.",
            "oggetto": "Un portachiavi a forma di struttura molecolare, regalo di laurea dei compagni.",
            "gusto": "Curry rice della mensa della centrale. Lo mangia quasi ogni giorno, senza vergogna.",
            "extra": ("Abitudine", "Conta i passi quando è nervoso. Non se ne accorge, ma i colleghi sì."),
        },
    },
    # ── PG_05 — FUJITA Emi ──
    {
        "filename": "Scheda_PG_05_Fujita_Emi.docx",
        "photo": "PG_05_Fujita_Emi_foto.png",
        "name_jp": "藤田 恵美",
        "name": "FUJITA Emi",
        "age": "36",
        "role": "Profiler / Psicologa Investigativa",
        "rank": "Ispettore (Keibu-ho)",
        "years": "8",
        "attrs": [
            ("Distacco", "4", "Normale"),
            ("Pazienza", "4", "Normale"),
            ("Silenzio", "7", "Osservi tutto senza reagire"),
            ("Lucidità", "6", "Solida capacità d'analisi"),
            ("Ascolto", "6", "Formazione psicologica — leggi le persone"),
            ("Presenza", "4", "Preferisci restare in ombra"),
        ],
        "gou": [
            {
                "name": "Specchio dell'Anima",
                "desc": "Vedi oltre la maschera. Senti le emozioni vere di una persona.",
                "attr": "Ascolto", "cost": "2 Ki",
                "success": "Senti l'emozione dominante e capisci perché la prova",
                "failure": "Senti l'emozione, ma non la sua origine",
            },
            {
                "name": "Tocco del Medico",
                "desc": "Leggi il corpo e la mente. Vedi i segni che altri ignorano.",
                "attr": "Lucidità", "cost": "2 Ki",
                "success": "Dettagli precisi — tipo di trauma, natura della paura, segni di stress cronico",
                "failure": "Capisci che qualcosa non va, ma non riesci a definirlo con precisione",
            },
            {
                "name": "Sussurro della Sera",
                "desc": "La tua voce calma. Le persone si tranquillizzano in tua presenza.",
                "attr": "Silenzio", "cost": "2 Ki",
                "success": "La persona si calma completamente e inizia a fidarsi di te",
                "failure": "La persona si calma, ma resta guardinga e sospettosa",
            },
        ],
        "background": "Sei arrivata in polizia dopo una laurea in psicologia e anni di pratica clinica. Capisci le persone. Vedi i pattern. Sai perché la gente fa quello che fa.\n\nSei anche una donna single in un ambiente dominato dagli uomini. Hai dovuto lavorare il doppio per guadagnarti il rispetto. Ora ce l'hai — ma è stato un prezzo.\n\nCinque anni fa hai avuto bisogno di soldi. Tuo padre era malato, le cure costavano. Hai chiesto un prestito a Iwamoto, un conoscente di famiglia che \"aiutava le persone\". Non era yakuza — solo un uomo con soldi e pochi scrupoli.\n\nTuo padre è morto comunque. Il debito è quasi saldato. Ma Iwamoto non sembra volerlo chiudere.",
        "problem_title": "Il Creditore",
        "problem_text": "Iwamoto Koji ti ha prestato soldi cinque anni fa. Hai restituito quasi tutto. Ma Iwamoto continua a presentarsi. \"Non c'è fretta.\" \"Siamo amici.\" \"Quando sarà il momento, ne parleremo.\"\n\nNon chiede mai nulla di specifico. Questo è il punto. L'ambiguità è l'arma. Vuole tenerti in debito — non per i soldi, ma per il potere di poter chiedere qualcosa un giorno.\n\nSai che Iwamoto non vuole chiudere il debito. Sai che ogni incontro è un promemoria: \"Ti tengo d'occhio\". E sai che se qualcuno in centrale lo vede parlare con te, ci saranno domande.",
        "png_name": "IWAMOTO Koji",
        "png_role": "Creditore",
        "png_age": "55",
        "png_job": "\"Uomo d'affari\" — prestiti privati, investimenti grigi",
        "png_wants": "Controllo. Il potere di poter chiedere un favore quando vuole.",
        "png_behavior": "Gentilissimo. Premuroso. Mai una parola su cosa potrebbe chiedere. L'ambiguità è l'arma.",
        "enja_name": "MORITA Akiko",
        "enja_role": "Medico Legale",
        "enja_desc": "Ex collega della facoltà di psicologia, ora medico legale. Autopsie, consulenze mediche, accesso ai dati sanitari.",
        "enja_contact": "Ospedale Universitario, reparto Medicina Legale. O al cellulare personale.",
        "enja_cost": "Caffè e chiacchierate, ogni tanto una consulenza psicologica informale.",
        "tatemae": "Sei professionale, competente, riservata. Non parli della tua vita privata. Non cerchi amicizie.",
        "honne": "Sei stufa di dover dimostrare il doppio. Sei stufa di Iwamoto. Sei stufa di essere sempre all'erta. Ma non lo mostrerai mai.",
        "phrase": "\"Non guardare cosa ha fatto. Guarda perché l'ha fatto.\"",
        "pressure": "Ti chiudi ancora di più. Diventi fredda, quasi clinica. Analizzi tutto, anche le tue emozioni — come se fossero di qualcun altro.",
        "weakness": "Quando un uomo la sminuisce o le dice che \"esagera\". Anni di lotta per il rispetto condensati in un istante.",
        "squad": [
            ("YAMAMOTO Kenji", "Un buon capo, ma si sta consumando. Non è il mio ruolo dirglielo, ma qualcuno dovrebbe."),
            ("HONDA Ryota", "Nasconde qualcosa. Non so cosa, ma il linguaggio del corpo non mente."),
            ("NAKAMURA Shota", "L'unico che non mi ha mai chiesto perché non sono sposata. Per questo lo rispetto."),
            ("SATO Yuki", "Giovane, entusiasta, fragile. Mi ricorda una versione di me che non esiste più."),
        ],
        "traits": {
            "vizio": "Whisky Nikka. Un dito, da sola, la sera dopo i casi pesanti. Non in compagnia.",
            "tic": "Inclina leggermente la testa a sinistra quando ascolta qualcuno mentire.",
            "oggetto": "Una penna stilografica Pilot nera. Scrive i profili sempre a mano, mai al computer.",
            "gusto": "Wagashi e tè matcha. Ha un debole per la pasticceria tradizionale vicino a Kiyomizu.",
            "extra": ("Rifugio", "Il tempio di Nanzen-ji, la sera. Si siede nel giardino di pietra e non pensa a nulla. L'unico posto dove si concede di non analizzare."),
        },
    },
]


def create_sheet(pg):
    doc = Document()

    # Page setup A4 with tight margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.3)
    section.bottom_margin = Cm(0.3)
    section.left_margin = Cm(0.3)
    section.right_margin = Cm(0.3)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ════════════════════════════════════════════
    # PAGE 1
    # ════════════════════════════════════════════

    # ── Top decorative line ──
    add_para(doc, "限界  ·  GENKAI v1.2  ·  限界", size=7, color=GOLD,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # ── Header: name block + photo (table 1x2) ──
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_tbl)

    # Left cell: name/kanji — vertically centered
    left_cell = header_tbl.rows[0].cells[0]
    left_cell.width = Cm(14)
    tc_left = left_cell._tc
    tcPr_left = tc_left.get_or_add_tcPr()
    tcPr_left.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))

    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, pg["name_jp"], size=28, color=DARK, font="Calibri")

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, pg["name"], bold=True, size=14, color=DARK)

    # Right cell: photo — tight fit, no padding
    right_cell = header_tbl.rows[0].cells[1]
    right_cell.width = Cm(3.2)
    tc_right = right_cell._tc
    tcPr_right = tc_right.get_or_add_tcPr()
    # Zero cell margins so photo fills the border exactly
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="0" w:type="dxa"/>'
        f'  <w:left w:w="0" w:type="dxa"/>'
        f'  <w:bottom w:w="0" w:type="dxa"/>'
        f'  <w:right w:w="0" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr_right.append(tcMar)
    tcPr_right.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
    # Thin gold border
    set_cell_borders(right_cell,
                     top=(4, GOLD_HEX), bottom=(4, GOLD_HEX),
                     left=(4, GOLD_HEX), right=(4, GOLD_HEX))

    # Insert photo or placeholder
    p_photo = right_cell.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    photo_path = OUT_DIR / pg.get("photo", "")
    if photo_path.exists():
        run = p_photo.add_run()
        run.add_picture(str(photo_path), width=Cm(3.2))
    else:
        r_photo = p_photo.add_run("FOTO")
        r_photo.font.size = Pt(12)
        r_photo.font.color.rgb = RGBColor(200, 200, 200)
        r_photo.font.name = "Calibri"

    add_separator(doc)

    # ── Info line ──
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(4)
    items = [
        ("Età ", pg["age"]),
        ("Ruolo ", pg["role"]),
        ("Grado ", pg["rank"]),
        ("Servizio ", f"{pg['years']} anni"),
    ]
    for i, (label, value) in enumerate(items):
        if i > 0:
            add_run(p3, "  ·  ", size=8, color=GOLD)
        add_run(p3, label, bold=True, size=8, color=MUTED)
        add_run(p3, value, size=8, color=DARK)

    add_thin_separator(doc)

    # ── ATTRIBUTI ──
    add_section_header(doc, "Attributi", "能力")

    tbl = doc.add_table(rows=7, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Elegant borders: gold top/bottom, no vertical, subtle warm horizontal
    tbl_el = tbl._tbl
    tblPr_t = tbl_el.tblPr if tbl_el.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr_t.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{GOLD_HEX}"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="D0C8B0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))
    tblPr_t.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="30" w:type="dxa"/>'
        f'  <w:bottom w:w="30" w:type="dxa"/>'
        f'  <w:left w:w="80" w:type="dxa"/>'
        f'  <w:right w:w="80" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))

    # Header row
    headers = [
        ("Attributo", WD_ALIGN_PARAGRAPH.LEFT),
        ("Base", WD_ALIGN_PARAGRAPH.CENTER),
        ("Finale", WD_ALIGN_PARAGRAPH.CENTER),
        ("Note", WD_ALIGN_PARAGRAPH.LEFT),
        ("Variazioni", WD_ALIGN_PARAGRAPH.CENTER),
    ]
    for j, (header, align) in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_shading(cell, DARK_HEX)
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    # Gold accent under header row
    for j in range(5):
        set_cell_borders(tbl.rows[0].cells[j], bottom=(6, GOLD_HEX))

    for i, (attr_name, attr_val, attr_note) in enumerate(pg["attrs"]):
        row = tbl.rows[i + 1]
        is_strong = int(attr_val) >= 6

        # Name
        cell0 = row.cells[0]
        p = cell0.paragraphs[0]
        r = p.add_run(attr_name)
        r.bold = is_strong
        r.font.size = Pt(9)
        r.font.name = "Calibri"
        r.font.color.rgb = DARK

        # Base value
        cell1 = row.cells[1]
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(attr_val)
        r.bold = is_strong
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        r.font.color.rgb = RED if is_strong else DARK

        # Finale (empty — player fills in)
        cell2 = row.cells[2]
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Note
        cell3 = row.cells[3]
        p = cell3.paragraphs[0]
        r = p.add_run(attr_note)
        r.italic = True
        r.font.size = Pt(8)
        r.font.name = "Calibri"
        r.font.color.rgb = MUTED

        # Variazioni (empty — player tracks changes)
        cell4 = row.cells[4]
        p = cell4.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_HEX)

    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(1.2)
        row.cells[2].width = Cm(1.2)
        row.cells[3].width = Cm(8.5)
        row.cells[4].width = Cm(3.6)

    # Points note — centered, subtle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "5 punti", bold=True, size=8, color=MUTED)
    add_run(p, " da distribuire  ·  max +2 per attributo  ·  max 8  ·  Scrivi il totale in ", size=8, color=MUTED)
    add_run(p, "Finale", bold=True, size=8, color=MUTED)

    add_thin_separator(doc)

    # ── KI ──
    add_section_header(doc, "Ki — Tenuta", "気")

    # Ki — elegant layout
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "Attributo Finale più basso + 1d6", italic=True, size=8, color=MUTED)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(0.3)
    add_run(p2, "Ki Max:  ________", bold=True, size=11, color=DARK)
    add_run(p2, "          ", size=11, color=DARK)
    add_run(p2, "≤ 3 Genkai", size=8, color=MUTED, italic=True)
    add_run(p2, "  ·  ", size=8, color=GOLD)
    add_run(p2, "= 1 Critico", size=8, color=MUTED, italic=True)
    add_run(p2, "  ·  ", size=8, color=GOLD)
    add_run(p2, "< 1 Fuori gioco", size=8, color=MUTED, italic=True)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(4)
    p3.paragraph_format.left_indent = Cm(0.3)
    add_run(p3, "Ki attuale:  ", bold=True, size=11, color=DARK)
    add_run(p3, "_______________________________________________________________", size=11, color=RGBColor(200, 200, 200))

    add_thin_separator(doc)

    # ── GOU ──
    add_section_header(doc, "Gou — Il Debito  (scegli uno)", "業")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "Il Gou funziona SEMPRE. Successo = dettagli precisi. Fallimento = dettagli vaghi.",
            italic=True, size=8, color=MUTED)

    for gou in pg["gou"]:
        add_gou_box(doc, gou["name"], gou["desc"], gou["attr"], gou["cost"],
                     gou["success"], gou["failure"])

    # Bottom marker
    add_separator(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Gou scelto:  ____________________________", size=9, color=MUTED)

    # ── Satori reminder ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "悟り  ", size=10, color=GOLD)
    add_run(p, "SATORI — Successo automatico, 1 volta per sessione:  ", size=8, color=MUTED)
    add_run(p, "□ Usato", size=9, color=DARK)

    # ════════════════════════════════════════════
    # PAGE 2
    # ════════════════════════════════════════════

    doc.add_page_break()

    # Top line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_run(p, pg["name_jp"], size=14, color=GOLD)
    add_run(p, "  ·  ", size=10, color=MUTED)
    add_run(p, pg["name"], bold=True, size=11, color=DARK)

    add_separator(doc)

    # ── CHI SEI ──
    add_section_header(doc, "Chi Sei", "人")

    for para_text in pg["background"].split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.3)
        add_run(p, para_text, size=9, color=DARK)

    add_thin_separator(doc)

    # ── IL TUO PROBLEMA ──
    add_section_header(doc, f"Il Tuo Problema — {pg['problem_title']}", "影")

    for para_text in pg["problem_text"].split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.3)
        add_run(p, para_text, size=9, color=DARK)

    add_thin_separator(doc)

    # ── PNG DEL PROBLEMA ──
    add_section_header(doc, pg["png_name"], "縁")

    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(tbl2)

    png_data = [
        ("Età", pg["png_age"]),
        ("Occupazione", pg["png_job"]),
        ("Relazione", pg["png_role"]),
        ("Cosa vuole", pg["png_wants"]),
    ]
    for i, (label, value) in enumerate(png_data):
        row = tbl2.rows[i]
        c0 = row.cells[0]
        p = c0.paragraphs[0]
        add_run(p, label, bold=True, size=8, color=MUTED)
        c0.width = Cm(2.5)

        c1 = row.cells[1]
        p = c1.paragraphs[0]
        add_run(p, value, size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, pg["png_behavior"], italic=True, size=9, color=DARK)

    add_thin_separator(doc)

    # ── CONOSCENZA ──
    add_section_header(doc, f"Conoscenza — {pg['enja_name']}", "縁者")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, pg["enja_role"], italic=True, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, pg["enja_desc"], size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Contatto: ", bold=True, size=8, color=MUTED)
    add_run(p, pg["enja_contact"], size=8, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Costo: ", bold=True, size=8, color=MUTED)
    add_run(p, pg["enja_cost"], size=8, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "1 volta per sessione senza conseguenze. La seconda volta → chiede qualcosa in cambio.",
            italic=True, size=7, color=MUTED)

    add_thin_separator(doc)

    # ── COME TI COMPORTI ──
    add_section_header(doc, "Come Ti Comporti", "面")

    # Tatemae
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "建前  In pubblico:  ", bold=True, size=9, color=GOLD)
    add_run(p, pg["tatemae"], size=9, color=DARK)

    # Honne
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "本音  In privato:  ", bold=True, size=9, color=GOLD)
    add_run(p, pg["honne"], size=9, color=DARK)

    # Phrase
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Frase tipica:  ", bold=True, size=9, color=MUTED)
    add_run(p, pg["phrase"], italic=True, size=9, color=DARK)

    # Pressure
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Sotto pressione:  ", bold=True, size=9, color=RED)
    add_run(p, pg["pressure"], size=9, color=DARK)

    # Weakness
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Debolezza:  ", bold=True, size=9, color=RED)
    add_run(p, pg["weakness"], size=9, color=DARK)

    # ── TRATTI PERSONALI ──
    add_thin_separator(doc)
    add_section_header(doc, "Tratti Personali", "癖")

    trait_labels = [
        ("Vizio", "vizio"),
        ("Tic", "tic"),
        ("Oggetto personale", "oggetto"),
        ("Gusto", "gusto"),
    ]
    for label, key in trait_labels:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{label}:  ", bold=True, size=9, color=MUTED)
        add_run(p, pg["traits"][key], size=9, color=DARK)

    # 5° tratto unico per PG
    extra_label, extra_text = pg["traits"]["extra"]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{extra_label}:  ", bold=True, size=9, color=MUTED)
    add_run(p, extra_text, size=9, color=DARK)

    # ── RAPPORTI NELLA SQUADRA ──
    add_thin_separator(doc)
    add_section_header(doc, "Rapporti nella Squadra", "絆")

    for name, quote in pg["squad"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{name}:  ", bold=True, size=9, color=MUTED)
        add_run(p, f"\u201c{quote}\u201d", italic=True, size=9, color=DARK)

    # ── Footer ──
    add_separator(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "GENKAI 限界 v1.2", size=7, color=MUTED)

    # Save
    out_path = OUT_DIR / pg["filename"]
    doc.save(str(out_path))
    print(f"OK — {out_path.name}")


def main():
    for pg in PG_DATA:
        create_sheet(pg)

if __name__ == "__main__":
    main()
