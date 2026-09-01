# -*- coding: utf-8 -*-
# Dossier del GM per «L'era glaciale del lavoro»: un Word solo con
# soluzione · cronistoria · schede PNG · luoghi · En · scheda distretto.
import io, sys, os, sqlite3, json, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.join(r'C:' + chr(92) + 'Public', '_Clienti', 'Maruga', 'Giochi', 'Vampiri', 'Vault', 'Vampiri', 'Investigare', 'Avventura Tanto Rumore'))
from contesto_tavolo import RUOLI, LOCATION, RUOLI_FILE, file_scheda  # fonte unica del contesto da tavolo
BASE = r'C:\Public\_Clienti\Maruga\Giochi\Vampiri\Vault\Vampiri\Investigare'
WIZ  = os.path.join(BASE, 'Wizard', 'codice', 'GenkaiWizard')
ID   = 'C3C15FF7-AFCE-4299-A49C-53B367CD29EA'
OUT  = os.path.join(BASE, 'Avventura Tanto Rumore', 'DOSSIER_GM.docx')

# ── dati ──
c = sqlite3.connect('file:' + os.path.join(WIZ, 'app.db').replace('\\','/') + '?mode=ro', uri=True)
titolo_caso, stato = c.execute("select Titolo, StatoJson from Progetti where Id=?", (ID,)).fetchone()
S = json.loads(stato)
CAST   = {p['id']: p for p in S['cast']}
LUOG   = {l['id']: l for l in S['luoghi']}
GRUP   = {g['id']: g for g in S['gruppi']}
SCHEDE = {s['personaId']: s for s in S['passo8']['schede']}
COLP   = set(S['passo5'].get('colpevoliIds') or [])
VITT   = S['passo2'].get('personaId')

def nome(i):
    if i in CAST: return (CAST[i]['cognome'] + ' ' + CAST[i]['nome']).strip()
    if i in GRUP: return GRUP[i]['nome']
    return ''
def luogo(i):
    l = LUOG.get(i)
    return (l.get('nome') or l.get('via') or '') if l else ''

_bib = {}
def lib(n):
    if n not in _bib:
        p = os.path.join(WIZ, 'Dati', 'biblioteche', n + '.json')
        _bib[n] = json.load(open(p, encoding='utf-8')) if os.path.exists(p) else {}
    return _bib[n]

def voce_lib(file, chiave, id_):
    """la voce di biblioteca con quell'id, o con quel nome (i campi liberi tengono il nome)"""
    if not id_: return None
    for x in (lib(file).get(chiave) or []):
        if x.get('id') == id_ or x.get('nome') == id_: return x
    return None
def nomeCat(file, chiave, id_):
    v = voce_lib(file, chiave, id_)
    return v['nome'] if v else (id_ or '')
def nome_omicidio(tid, sid):
    for t in (lib('tipologie_omicidio').get('tipologie') or []):
        if t.get('id') == tid:
            for s in (t.get('sottotipi') or []):
                if s.get('id') == sid: return f"{t['nome']} — {s['nome']}"
            return t['nome']
    return tid or ''
def nome_movente(mid):
    for m in (lib('moventi').get('moventi') or []):
        if m.get('id') == mid: return m['nome']
        for s in (m.get('sottocasi') or []):
            if s.get('id') == mid: return f"{m['nome']} — {s['nome']}"
    return mid or ''

# ── documento ──
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(1.6)
sec.left_margin = sec.right_margin = Cm(1.8)

ORO    = RGBColor(0xB4, 0x96, 0x50)
NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
GRIGIO = RGBColor(0x70, 0x6A, 0x5E)
ROSSO  = RGBColor(0x8B, 0x1A, 0x1A)

st = doc.styles['Normal']
st.font.name = 'Garamond'; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
st.paragraph_format.space_after = Pt(3)
st.paragraph_format.line_spacing = 1.06

def P(testo='', size=10.5, bold=False, italic=False, colore=None, dopo=3, prima=0, align=None, indent=0):
    p = doc.add_paragraph(); r = p.add_run(testo)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if colore is not None: r.font.color.rgb = colore
    p.paragraph_format.space_after = Pt(dopo); p.paragraph_format.space_before = Pt(prima)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if align is not None: p.alignment = align
    return p

def etichetta(p, lab, val, size=10):
    r = p.add_run(lab); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = NAVY
    r2 = p.add_run(str(val)); r2.font.size = Pt(size)
    return p

def linea(dopo=6, prima=6):
    p = doc.add_paragraph(); r = p.add_run('━' * 62)
    r.font.size = Pt(5); r.font.color.rgb = ORO
    p.paragraph_format.space_after = Pt(dopo); p.paragraph_format.space_before = Pt(prima)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def bordo_sotto(p, colore='B49650', sz=6):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    for k, v in (('val','single'), ('sz',str(sz)), ('color',colore), ('space','2')): bt.set(qn('w:'+k), v)
    b.append(bt); pPr.append(b)

def titolo_sez(testo, kanji='', prima=14):
    p = doc.add_paragraph()
    r = p.add_run(testo.upper()); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    if kanji:
        t = p.add_run('   ' + kanji); t.font.size = Pt(12); t.font.color.rgb = ORO
    p.paragraph_format.space_before = Pt(prima); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    bordo_sotto(p, sz=10)
    return p

def sottotit(testo, prima=6):
    return P(testo, size=9.5, bold=True, colore=ORO, prima=prima, dopo=2)

def paragrafi(testo, size=10.5, italic=False, indent=0.4):
    for blocco in [b for b in re.split(r'\n+', (testo or '').strip()) if b.strip()]:
        P(blocco.strip(), size=size, italic=italic, dopo=3, indent=indent)

def ripeti_intestazione(riga):
    """la prima riga della tabella si ristampa in cima a ogni pagina"""
    trPr = riga._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)

def fissa_larghezze(tab, larghezze):
    """Word rispetta le larghezze solo con l'autofit spento"""
    tab.autofit = False
    tblPr = tab._tbl.tblPr
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    for riga in tab.rows:
        for i, cel in enumerate(riga.cells):
            if i < len(larghezze): cel.width = Cm(larghezze[i])

def tabellina(righe, larghezze=(3.6, 13.4)):
    righe = [(l, v) for l, v in righe if str(v).strip()]
    if not righe: return None
    t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for lab, val in righe:
        c1, c2 = t.add_row().cells
        c1.width, c2.width = Cm(larghezze[0]), Cm(larghezze[1])
        r = c1.paragraphs[0].add_run(lab); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GRIGIO
        r2 = c2.paragraphs[0].add_run(str(val)); r2.font.size = Pt(10)
        for cc in (c1, c2): cc.paragraphs[0].paragraph_format.space_after = Pt(1)
    fissa_larghezze(t, larghezze)
    return t

# ═══════════ COPERTINA ═══════════
# titolo «Bakuon» col kanji, sottotitolo «Tanto Rumore per Nulla» (decisione utente 2026-08-16)
TITOLO_BREVE = titolo_caso.split('—')[0].strip() or 'Bakuon'
SOTTOTITOLO  = titolo_caso.split('—')[1].strip() if '—' in titolo_caso else ''
P('GENKAI ver: 1.3', size=9, colore=ORO, align=WD_ALIGN_PARAGRAPH.CENTER, dopo=2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(1)
r = p.add_run(TITOLO_BREVE.upper()); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
r2 = p.add_run('  爆音'); r2.font.size = Pt(22); r2.font.color.rgb = ORO
if SOTTOTITOLO:
    P(SOTTOTITOLO, size=13, italic=True, colore=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, dopo=3)
setup = S.get('setup', {})
P(f"Kyoto, {nomeCat('luoghi','quartieri', setup.get('quartiere'))} — maggio 1997",
  size=11, italic=True, colore=GRIGIO, align=WD_ALIGN_PARAGRAPH.CENTER, dopo=4)
linea(dopo=8, prima=6)
P('DOSSIER DEL GAME MASTER', size=11, bold=True, colore=ROSSO, align=WD_ALIGN_PARAGRAPH.CENTER, dopo=2)
P('Contiene la soluzione del caso. Non va mostrato ai giocatori.', size=9.5, italic=True,
  colore=GRIGIO, align=WD_ALIGN_PARAGRAPH.CENTER, dopo=8)

# ── il briefing del capo — per aprire (testo fisso, deciso dall'utente 2026-09-01) ──
titolo_sez('Il briefing del capo — per aprire', '訓示', prima=8)
P("Domenica 25 maggio 1997, mattina. Il capo ispettore alla squadra — da leggere ai giocatori così com'è:",
  size=9.5, italic=True, colore=GRIGIO, dopo=4)
P("«Sabato sera, 21:15, snack bar SnakUp, Shimogyō. Shimada Yuta, ventun anni, studente di medicina "
  "che faceva le serate lì per pagarsi gli studi: un colpo alla nuca, da dietro, oggetto contundente. "
  "Morto prima che arrivasse l'ambulanza. Quattro, forse cinque giovani: entrati, trenta secondi di lite, "
  "il colpo, e via in moto. Nessun nome, nessuna faccia — l'unica cosa che tutti ricordano è il rumore: "
  "quel rombo di scarichi tagliati che ha svegliato mezza via.", size=11, italic=True, dopo=3, indent=0.4)
P("Il titolare e la cameriera erano dentro. La madre e la sorella aspettano da stanotte che qualcuno "
  "gli spieghi perché. Il caso è vostro: portatemi chi ha dato quel colpo — prima che i giornali "
  "ci mangino vivi.»", size=11, italic=True, dopo=3, indent=0.4)

# ── il caso in breve ──
titolo_sez('Il caso in breve', '事件', prima=8)
q1, q2, q3, q4, q5, q6 = (S.get(k) or {} for k in ('passo1','passo2','passo3','passo4','passo5','passo6'))
P(q1.get('rigaUnica',''), size=12.5, bold=True, dopo=5)
uno_colp = list(COLP)[0] if COLP else None
tabellina([
    ('Come è morto', nome_omicidio(q1.get('tipologiaId'), q1.get('sottotipoId'))),
    ('La vittima', ' · '.join(x for x in [nome(VITT), f"{q2.get('eta','')} anni",
        nomeCat('professioni','professioni', q2.get('professioneId')), q2.get('postoNelMondo','')] if x)),
    ('Il colpevole', ' · '.join(x for x in [', '.join(nome(i) for i in COLP),
        (f"{CAST[uno_colp].get('eta','')} anni" if uno_colp else ''),
        (CAST[uno_colp].get('professione') or '') if uno_colp else ''] if x)),
    ('Il movente', nome_movente(q4.get('moventeId'))),
    ('Cos\u2019è successo', q4.get('descrizione','')),
    ('Chi è l\u2019assassino', q5.get('competenze','')),
    ('Lo sbaglio che lo tradisce', q5.get('erroreCoerente','')),
    ('Il legame con la vittima', q6.get('connessioneVittima','')),
    ('Cosa fa dopo il fatto', q6.get('dopoIlFatto','')),
])
probl = [p for p in (q3.get('problemi') or []) if (p.get('testo') or '').strip()]
if probl:
    sottotit('I nodi nella vita della vittima', prima=8)
    for p in probl:
        pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.4); pp.paragraph_format.space_after = Pt(2)
        r = pp.add_run(p['testo'].strip()); r.font.size = Pt(10)
        if p.get('potenzialeFalsaPista'):
            r2 = pp.add_run('   falsa pista'); r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = ROSSO
cand = [x for x in (q4.get('candidati') or []) if (x.get('perche') or '').strip() and x.get('personaId') not in COLP]
if cand:
    sottotit('Chi altro aveva la stessa ragione', prima=8)
    for x in cand:
        pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.4); pp.paragraph_format.space_after = Pt(2)
        etichetta(pp, nome(x['personaId']) + ': ', x['perche'].strip())

# ═══════════ CRONISTORIA ═══════════
doc.add_page_break()
titolo_sez('Cronistoria — la verità in ordine', '時系列', prima=0)
P('Quello che è successo davvero, che i giocatori lo scoprano o no. Le tracce sono ciò che resta e si può trovare.',
  size=9.5, italic=True, colore=GRIGIO, dopo=6)

eventi = sorted(S['passo7']['eventi'], key=lambda e: (e.get('quando') or ''))
for fase, lab in (('prima','Prima del fatto'), ('fatto','Il fatto'), ('dopo','Dopo il fatto')):
    ev = [e for e in eventi if e.get('fase') == fase]
    if not ev: continue
    p = P(lab, size=10.5, bold=True, colore=(ROSSO if fase == 'fatto' else ORO), prima=10, dopo=4)
    bordo_sotto(p)
    for e in ev:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
        r = p.add_run(e.get('quando','—')); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
        lg = luogo(e.get('luogoId'))
        if lg:
            r2 = p.add_run('   ' + lg); r2.font.size = Pt(9.5); r2.italic = True; r2.font.color.rgb = GRIGIO
        P(e.get('testo','').strip(), size=10.5, dopo=2, indent=0.5)
        chi = [nome(i) for i in (e.get('personeIds') or []) if nome(i)]
        if chi:
            pc = doc.add_paragraph(); pc.paragraph_format.left_indent = Cm(0.5); pc.paragraph_format.space_after = Pt(2)
            etichetta(pc, 'Presenti: ', ', '.join(chi), size=9.5)
        for tr in (e.get('generaTraccia') or []):
            if not (tr or '').strip(): continue
            pt = doc.add_paragraph(); pt.paragraph_format.left_indent = Cm(0.5); pt.paragraph_format.space_after = Pt(2)
            r = pt.add_run('Traccia   '); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = ROSSO
            r2 = pt.add_run(tr.strip()); r2.font.size = Pt(9.5); r2.italic = True
        doc.paragraphs[-1].paragraph_format.space_after = Pt(7)

# ═══════════ LE PERSONE ═══════════
doc.add_page_break()
titolo_sez('Le persone del caso', '人物', prima=0)
P('Chi sono, cosa sanno, cosa nascondono. Tutto quello che segue è verità del GM: quello che i giocatori '
  'sentiranno è solo la deposizione.', size=9.5, italic=True, colore=GRIGIO, dopo=4)

CERCHI = ('famiglia','lavoro','amici','altri','intersezione')
def cerchia_di(pid):
    for base, chi in (('passo3','della vittima'), ('passo6', "dell'assassino")):
        for cc in CERCHI:
            for r in (S.get(base, {}).get(cc) or []):
                if r.get('personaId') == pid and (r.get('relazione') or '').strip():
                    return f"{r['relazione']} — nel mondo {chi}"
    return ''
def ruolo_label(p):
    if p['id'] in RUOLI: return RUOLI[p['id']]
    if p['id'] == VITT: return 'la vittima'
    if p['id'] in COLP: return "l'assassino"
    return {'testimone':'testimone','falsaPista':'falsa pista','complice':'complice'}.get(p.get('ruoloNelCaso',''), '')
def gruppi_di(pid):
    return [g['nome'] for g in S['gruppi'] if pid in (g.get('membriIds') or [])]
def luoghi_suoi(pid):
    """i luoghi che il caso lega a questa persona (casa, lavoro, covo…)"""
    return [l.get('nome') or l.get('via') or '' for l in S['luoghi'] if l.get('personaId') == pid]

ordinati = sorted(CAST.values(), key=lambda p: (p['id'] != VITT, p['id'] not in COLP,
                                                p['id'] not in SCHEDE, p['cognome'] + p['nome']))
ETICH_VOCE = dict(intercalare='Intercalare', marcatore='Marcatore', appellativo='Come ti chiama',
                  abitudine='Abitudine', saluto='Saluto', rifiuto='Se rifiuta', minaccia='Se minaccia',
                  vizio='Vizio', tic='Tic', oggetto='Oggetto')
ETICH_CONT = dict(telefono='Telefono', cellulare='Cellulare', email='Email', altro='Altro', dove='Dove trovarlo')
senza_scheda = []

primo = True
for pers in ordinati:
    pid = pers['id']; sch = SCHEDE.get(pid, {})
    if not primo: linea(prima=10, dopo=5)
    primo = False
    # intestazione
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
    r = p.add_run(nome(pid)); r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = NAVY
    if pers.get('kanji'):
        rk = p.add_run('  ' + pers['kanji']); rk.font.size = Pt(11.5); rk.font.color.rgb = GRIGIO
    rl = ruolo_label(pers)
    if rl:
        rr = p.add_run('   ' + rl); rr.font.size = Pt(10); rr.italic = True
        rr.font.color.rgb = ROSSO if (pid in COLP or pid == VITT) else GRIGIO
    sotto = ' · '.join(x for x in [
        f"{pers['eta']} anni" if pers.get('eta') else '',
        {'m':'uomo','f':'donna'}.get(pers.get('genere'), ''),
        (nomeCat('professioni','professioni', q2.get('professioneId')) if pid == VITT else (pers.get('professione') or '')),
        cerchia_di(pid), ', '.join(gruppi_di(pid))] if x)
    if sotto: P(sotto, size=9.5, colore=GRIGIO, dopo=3)
    if (pers.get('note') or '').strip():
        P(re.sub(r'^\[ex Dati base\]\s*', '', pers['note'].strip()), size=10, italic=True, dopo=3, indent=0.4)

    testi = {k: (sch.get(k) or '').strip() for k in
             ('descrizioneFisica','cosaSa','cosaNonSa','cosaHaFatto','comportamento','deposizione')}
    SOSTANZA = ('cosaSa', 'cosaNonSa', 'cosaHaFatto', 'comportamento', 'deposizione')

    if pid == VITT:
        # la vittima non depone: quello che si sa di lei sta nei passi 2 e 3
        if testi['descrizioneFisica']:
            sottotit('Che aspetto aveva'); paragrafi(testi['descrizioneFisica'])
        if (q2.get('postoNelMondo') or '').strip():
            sottotit('Chi era'); paragrafi(q2['postoNelMondo'])
        vicini = []
        for cc in CERCHI:
            for r in (q3.get(cc) or []):
                if nome(r.get('personaId','')) and (r.get('relazione') or '').strip():
                    vicini.append(f"{nome(r['personaId'])} — {r['relazione']}")
        if vicini:
            sottotit('Chi aveva intorno')
            for v in vicini: P('· ' + v, size=10, dopo=1, indent=0.5)
        suoi = luoghi_suoi(pid)
        if suoi: P('Luoghi suoi: ' + ', '.join(suoi), size=9.5, colore=GRIGIO, prima=4, dopo=2, indent=0.4)
        continue

    if testi['descrizioneFisica']:
        sottotit('Che aspetto ha'); paragrafi(testi['descrizioneFisica'])
    if not any(testi[k] for k in SOSTANZA):
        # c'è il ritratto e l'aspetto, ma non cosa sa né come si comporta
        senza_scheda.append(nome(pid))
        suoi = luoghi_suoi(pid)
        if suoi: P('Luoghi suoi: ' + ', '.join(suoi), size=9.5, colore=GRIGIO, prima=4, dopo=2, indent=0.4)
        P('Scheda non compilata: manca cosa sa, come si comporta e la deposizione. Al tavolo si improvvisa.',
          size=9.5, italic=True, colore=ROSSO, dopo=2, indent=0.4)
        continue
    for chiave, lab in (('cosaSa','Cosa sa'), ('cosaNonSa','Cosa NON sa'),
                        ('cosaHaFatto','Cosa ha fatto davvero'),
                        ('comportamento','Come si comporta, cosa nasconde')):
        if testi[chiave]:
            sottotit(lab); paragrafi(testi[chiave])
    trg = [t for t in (sch.get('trigger') or []) if (t.get('se') or '').strip()]
    if trg:
        sottotit('Cosa lo fa cedere')
        for t in trg:
            P(t['se'].strip() + ((' → ' + t['allora'].strip()) if (t.get('allora') or '').strip() else ''),
              size=10, dopo=2, indent=0.4)
    voce   = {k: v for k, v in (sch.get('voce') or {}).items() if (v or '').strip()}
    tratti = {k: v for k, v in (sch.get('tratti') or {}).items() if (v or '').strip()}
    if voce or tratti:
        sottotit('Come parla, com\u2019è fatto')
        tabellina([(ETICH_VOCE.get(k, k), v) for k, v in list(voce.items()) + list(tratti.items())], (3.4, 13.6))
    cont = {k: v for k, v in (sch.get('contatti') or {}).items()
            if isinstance(v, str) and v.strip() and k != 'residenzaLuogoId'}
    ab   = luogo((sch.get('contatti') or {}).get('residenzaLuogoId'))
    suoi = [x for x in luoghi_suoi(pid) if x != ab]
    if cont or ab or suoi:
        sottotit('Dove si trova')
        tabellina(([('Abita', ab)] if ab else []) +
                  [(ETICH_CONT.get(k, k), v) for k, v in cont.items()] +
                  ([('Luoghi suoi', ', '.join(suoi))] if suoi else []), (3.4, 13.6))
    # ── tutto ciò che riguarda questa persona, nella sua scheda (richiesta utente 2026-09-01) ──
    evs = [e for e in S['passo7']['eventi'] if pid in (e.get('personeIds') or [])]
    if evs:
        sottotit('Nella storia (verità del GM)')
        for e in sorted(evs, key=lambda x: x.get('quando','')):
            dove = luogo(e.get('luogoId'))
            P('· ' + (e.get('quando','—') or '—') + (' · ' + dove if dove else '') + ' — ' + (e.get('testo','') or '').strip(),
              size=9.5, dopo=1, indent=0.5)
    _gr = [g['id'] for g in S['gruppi'] if pid in (g.get('membriIds') or [])]
    _fonti = [(tr, f) for tr in S['passo9']['tracce'] for f in (tr.get('fonti') or [])
              if f.get('attoreId') == pid or f.get('attoreId') in _gr]
    if _fonti:
        sottotit('Cosa può dare ai PG')
        _COME = {'interrogatorio': 'interrogandolo', 'richiestaEnte': 'con richiesta all' + chr(39) + 'ente', 'nulla': 'senza condizioni'}
        for tr, f in _fonti:
            _t = '· «' + tr['nome'] + '» (' + _COME.get(f.get('richiede',''), f.get('richiede','')) + ')'
            if (f.get('versione') or '').strip(): _t += ': ' + f['versione'].strip()
            if f.get('handout'): _t += ' — diventa handout: «' + (f.get('handoutTitolo') or tr['nome']) + '»'
            P(_t, size=9.5, dopo=1, indent=0.5)
    _rel = [r2 for r2 in S.get('relazioni', []) if pid in (r2.get('aId'), r2.get('bId'))]
    if _rel:
        sottotit('I suoi legami (En)')
        for r2 in _rel:
            _altro = r2['bId'] if r2['aId'] == pid else r2['aId']
            _da = r2['enAB'] if r2['aId'] == pid else r2['enBA']
            _ri = r2['enBA'] if r2['aId'] == pid else r2['enAB']
            _tipo = (r2.get('tipo') or '').strip()
            P('· con ' + nome(_altro) + (' (' + _tipo + ')' if _tipo else '') + ': En suo verso l' + chr(39) + 'altro ' + str(_da) + ' · in risposta ' + str(_ri),
              size=9.5, dopo=1, indent=0.5)
    stt = {k: v for k, v in (sch.get('stats') or {}).items() if k != 'esempio' and v not in (None, '')}
    if stt:
        pp = doc.add_paragraph(); pp.paragraph_format.space_before = Pt(5); pp.paragraph_format.space_after = Pt(2)
        etichetta(pp, 'Statistiche   ', ' · '.join(f'{k} {v}' for k, v in stt.items()), size=9.5)
        es = (sch.get('stats') or {}).get('esempio')
        if es:
            r = pp.add_run(f'   (profilo: {es})'); r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRIGIO
    if testi['deposizione']:
        sottotit('La sua deposizione')
        paragrafi(testi['deposizione'], size=10, italic=True)

# ═══════════ LUOGHI ═══════════
doc.add_page_break()
titolo_sez('I luoghi', '場所', prima=0)
for l in S['luoghi']:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
    r = p.add_run((l.get('icona','') + ' ' if l.get('icona') else '') + (l.get('nome') or '(senza nome)'))
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    if l.get('segretoPG'):
        rs = p.add_run('   i giocatori non sanno che esiste')
        rs.font.size = Pt(9); rs.italic = True; rs.font.color.rgb = ROSSO
    meta = ' · '.join(x for x in [nomeCat('luoghi','quartieri', l.get('quartiere')),
                                  nomeCat('luoghi','tipologie', l.get('tipologiaId'))] if x)
    if meta: P(meta, size=9.5, colore=GRIGIO, dopo=2)
    if (l.get('via') or '').strip(): P(l['via'].strip(), size=10, dopo=2, indent=0.4)
    chi = nome(l.get('personaId',''))
    if chi:
        pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.4); pp.paragraph_format.space_after = Pt(2)
        etichetta(pp, 'Legato a: ', chi, size=10)
    ctx = LOCATION.get((l.get('nome') or '').strip())
    if ctx:
        sottotit('Quando i PG arrivano', prima=3)
        P(ctx['arrivo'], size=10, italic=True, dopo=2, indent=0.4)
        sottotit('Quando entrano', prima=3)
        P(ctx['entrata'], size=10, italic=True, dopo=2, indent=0.4)
    tip = voce_lib('luoghi', 'tipologie', l.get('tipologiaId'))
    if tip:
        for campo, lab in (('cosaSiVede','Che posto è'), ('tracceTipiche','Cosa ci si trova, di solito')):
            v = tip.get(campo)
            if not v: continue
            v = ', '.join(v) if isinstance(v, list) else v
            pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.4); pp.paragraph_format.space_after = Pt(2)
            r = pp.add_run(lab + ': '); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GRIGIO
            r2 = pp.add_run(v); r2.font.size = Pt(9.5); r2.italic = True
    doc.paragraphs[-1].paragraph_format.space_after = Pt(9)

# ═══════════ GRUPPI E EN ═══════════
titolo_sez('Gruppi e legami', '縁', prima=12)
if S.get('gruppi'):
    sottotit('I gruppi', prima=2)
    for g in S['gruppi']:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(g.get('nome','')); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = NAVY
        meta = ' · '.join(x for x in [g.get('tipo',''), g.get('zona','')] if x)
        if meta:
            r2 = p.add_run('   ' + meta); r2.font.size = Pt(9.5); r2.italic = True; r2.font.color.rgb = GRIGIO
        if g.get('segretoPG'):
            r3 = p.add_run('   i giocatori non sanno che esiste')
            r3.font.size = Pt(9); r3.italic = True; r3.font.color.rgb = ROSSO
        if (g.get('descrizione') or '').strip(): P(g['descrizione'].strip(), size=10, dopo=2, indent=0.4)
        membri = [nome(m) for m in (g.get('membriIds') or []) if nome(m)]
        if membri:
            pm = doc.add_paragraph(); pm.paragraph_format.left_indent = Cm(0.4); pm.paragraph_format.space_after = Pt(3)
            etichetta(pm, 'Chi ne fa parte: ', ', '.join(membri), size=9.5)
        ente = SCHEDE.get(g['id'], {}).get('statsEnte')
        if ente:
            pe = doc.add_paragraph(); pe.paragraph_format.left_indent = Cm(0.4); pe.paragraph_format.space_after = Pt(6)
            etichetta(pe, 'Attributi dell\u2019ente: ',
                      ' · '.join(f'{k} {v}' for k, v in ente.items() if v not in (None, '')), size=9.5)

sottotit('L\u2019En — chi prova cosa verso chi', prima=10)
P('En positivo si sottrae dal dado, negativo si aggiunge. A ±4 e ±5 il legame è profondo: non lo smuovono i tiri, '
  'solo i fatti grossi.', size=9, italic=True, colore=GRIGIO, dopo=4)
t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.LEFT
COLW = (5.0, 5.0, 4.8, 1.8)
for i, h in enumerate(('Chi', 'Verso', 'Che legame', 'En')):
    cel = t.rows[0].cells[i]; cel.width = Cm(COLW[i])
    r = cel.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRIGIO
    cel.paragraphs[0].paragraph_format.space_after = Pt(1)
    bordo_sotto(cel.paragraphs[0], sz=6)
for rel in S.get('relazioni', []):
    a, b = nome(rel.get('aId','')), nome(rel.get('bId',''))
    if not a or not b: continue
    for (x, y, en) in ((a, b, rel.get('enAB')), (b, a, rel.get('enBA'))):
        if str(en).strip() in ('', 'None'): continue
        cs = t.add_row().cells
        for i, (txt, grasse) in enumerate(((x, True), (y, False), (rel.get('tipo','') or '—', False), (str(en), True))):
            cs[i].width = Cm(COLW[i])
            r = cs[i].paragraphs[0].add_run(txt); r.font.size = Pt(9.5); r.bold = grasse
            if i == 3:
                try: r.font.color.rgb = ROSSO if int(en) < 0 else NAVY
                except ValueError: pass
            cs[i].paragraphs[0].paragraph_format.space_after = Pt(1)
ripeti_intestazione(t.rows[0]); fissa_larghezze(t, COLW)

# ═══════════ INFORMAZIONI ═══════════
tracce = [t for t in (S.get('passo9', {}).get('tracce') or []) if (t.get('nome') or '').strip()]
if tracce:
    doc.add_page_break()
    titolo_sez('Le informazioni — chi sa cosa', '情報', prima=0)
    P('Quello che si può scoprire, e da chi. Gli indizi si danno sempre: il dado decide come ne esci, non se trovi.',
      size=9.5, italic=True, colore=GRIGIO, dopo=6)
    RICH = {'nulla':'basta chiedere', 'interrogatorio':'interrogatorio', 'richiestaEnte':'richiesta a un ente',
            'sopralluogo':'sopralluogo', 'perquisizione':'perquisizione', 'pedinamento':'pedinamento'}
    for tr in tracce:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
        r = p.add_run(tr['nome']); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
        if (tr.get('classificazione') or '').strip():
            r2 = p.add_run('   ' + tr['classificazione']); r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = ORO
        if (tr.get('testo') or '').strip(): paragrafi(tr['testo'])
        for f in (tr.get('fonti') or []):
            chi = nome(f.get('attoreId','')) or f.get('canale','')
            if not chi: continue
            pf = doc.add_paragraph(); pf.paragraph_format.left_indent = Cm(0.4); pf.paragraph_format.space_after = Pt(2)
            r = pf.add_run('da ' + chi); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
            ric = f.get('richiede','')
            if ric and ric != 'nulla':
                r2 = pf.add_run('  (' + RICH.get(ric, ric) + ')'); r2.font.size = Pt(9); r2.font.color.rgb = GRIGIO
            if f.get('handout'):
                r3 = pf.add_run('  handout'); r3.font.size = Pt(9); r3.italic = True; r3.font.color.rgb = ORO
            for campo in ('versione', 'handoutTitolo'):
                if (f.get(campo) or '').strip():
                    P(f[campo].strip(), size=9.5, italic=True, dopo=2, indent=0.8)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(9)

# ═══════════ CALENDARIO ═══════════
giorni = [g for g in (S.get('passo11', {}).get('giorni') or []) if (g.get('evento') or '').strip()]
if giorni:
    titolo_sez('Il calendario vivo', '暦', prima=12)
    P('Cosa succede comunque, che i giocatori guardino o no.', size=9.5, italic=True, colore=GRIGIO, dopo=4)
    for g in giorni:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
        r = p.add_run(f"Giorno {g.get('giorno','')}" + (f" · {g['momento']}" if (g.get('momento') or '').strip() else ''))
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
        P(g['evento'].strip(), size=10, dopo=3, indent=0.5)
        if (g.get('condizione') or '').strip():
            pc = doc.add_paragraph(); pc.paragraph_format.left_indent = Cm(0.5); pc.paragraph_format.space_after = Pt(6)
            etichetta(pc, 'Solo se: ', g['condizione'].strip(), size=9.5)

# ═══════════ SCHEDA DISTRETTO ═══════════
doc.add_page_break()
titolo_sez('La scheda del distretto', '所轄', prima=0)
P('Quando i PG delegano un\u2019operazione, il GM tira 2d6 sull\u2019attributo dell\u2019ente. '
  'La Corruzione non si tira mai davanti ai giocatori.', size=9.5, italic=True, colore=GRIGIO, dopo=4)

sd = open(os.path.join(BASE, 'Materiale', 'Scheda_Distretto.md'), encoding='utf-8').read()

def sezioni_md(testo):
    """le sezioni di secondo livello: [(titolo, corpo), …] — il corpo tiene i suoi ###"""
    tagli = [(m.start(), m.group(1).strip()) for m in re.finditer(r'^## (.+)$', testo, re.M)]
    fuori = [(s, t, (tagli[i+1][0] if i+1 < len(tagli) else len(testo))) for i, (s, t) in enumerate(tagli)]
    return [(t, testo[s:fine].split('\n', 1)[1].strip()) for s, t, fine in fuori]

# la testata del file (tipologia ente, sezione, inquadramento): sta prima del primo ##
testata = sd.split('\n## ', 1)[0].split('\n', 1)[1].strip()

def run_md(p, testo, size=10):
    """scrive testo markdown inline (**grassetto**, *corsivo*) dentro un paragrafo"""
    for pezzo in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', testo):
        if not pezzo: continue
        grassetto = pezzo.startswith('**') and pezzo.endswith('**') and len(pezzo) > 4
        corsivo = not grassetto and pezzo.startswith('*') and pezzo.endswith('*') and len(pezzo) > 2
        r = p.add_run(pezzo.strip('*') if (grassetto or corsivo) else pezzo)
        r.font.size = Pt(size); r.bold = grassetto; r.italic = corsivo

TESTO_UTILE = 17.4                                   # cm fra i margini

def larghezze_per(righe):
    """colonne proporzionali al testo più lungo che ciascuna contiene"""
    n = len(righe[0])
    pesi = [max(6, max(len(r[i]) for r in righe if i < len(r))) for i in range(n)]
    scala = TESTO_UTILE / sum(pesi)
    largh = [max(2.0, p * scala) for p in pesi]
    return [x * TESTO_UTILE / sum(largh) for x in largh]             # rinormalizza dopo i minimi

def scrivi_tabella(righe):
    largh = larghezze_per(righe)
    tab = doc.add_table(rows=0, cols=len(righe[0])); tab.alignment = WD_TABLE_ALIGNMENT.LEFT
    for n, celle in enumerate(righe):
        cs = tab.add_row().cells
        for i, cc in enumerate(celle):
            if i >= len(cs): break
            cs[i].width = Cm(largh[i])
            cs[i].paragraphs[0].paragraph_format.space_after = Pt(1)
            run_md(cs[i].paragraphs[0], cc, 9.5)
            if n == 0:
                for run in cs[i].paragraphs[0].runs:
                    run.bold = True; run.font.color.rgb = GRIGIO
                bordo_sotto(cs[i].paragraphs[0], sz=4)
    ripeti_intestazione(tab.rows[0]); fissa_larghezze(tab, largh)

def rendi_md(testo, size=10):
    buffer = []                                                      # righe della tabella in corso
    def chiudi():
        if buffer: scrivi_tabella(buffer[:]); buffer.clear()
    for riga in testo.split('\n'):
        r = re.sub(r'`([^`]*)`', r'\1', riga.strip())                # via i backtick
        if r.startswith('|'):
            celle = [x.strip() for x in r.strip('|').split('|')]
            if all(set(x) <= set('-: ') for x in celle): continue     # riga separatrice
            if buffer and len(celle) != len(buffer[0]): chiudi()
            buffer.append(celle); continue
        chiudi()
        if 'Scheda:' in r and '.md' in r: continue                    # puntatori al vault: non servono in stampa
        if not r or r.startswith('---'): continue
        if r.startswith('### '):
            P(r[4:].replace('*', '').strip(), size=10.5, bold=True, colore=ORO, prima=9, dopo=2); continue
        if r.startswith('#'):
            P(r.lstrip('#').replace('*', '').strip(), size=10.5, bold=True, colore=ORO, prima=8, dopo=2); continue
        if r.startswith('>'):                                         # nota di servizio
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.space_after = Pt(2)
            run_md(p, r.lstrip('> ').strip(), 9.5)
            for run in p.runs: run.italic = True; run.font.color.rgb = GRIGIO
            continue
        if r.startswith(('- ', '* ')):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.space_after = Pt(1)
            run_md(p, '· ' + r[2:], 9.5); continue
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        run_md(p, r, size)
    chiudi()

rendi_md(testata, size=10)

# titoli più brevi dove il file è verboso; la squadra dei PG si salta (schede già stampate a parte)
RINOMINA = {'Attributi del Distretto': 'Gli attributi',
            'Note descrittive degli attributi': 'Cosa vogliono dire',
            'Personale fisso del Distretto (PNG notevoli)': 'Chi ci lavora',
            'Azioni delegabili — riferimento generico': 'Cosa si può delegare',
            'Difficoltà — modificatori al tiro': 'Le difficoltà',
            'Tabella risultati (riepilogo dal Manuale Situazioni)': 'I risultati',
            'Esempi di tiri': 'Come viene, in pratica',
            'Procura della Repubblica — riferimento operativo': 'La Procura',
            'Quadro legale di riferimento': 'Cosa dice la legge, nel 1997',
            'Note di gioco al tavolo': 'Come si usa al tavolo'}
SALTA = ('La Squadra Investigativa Centrale (i 5 PG)',)

viste = []
for tit, corpo in sezioni_md(sd):
    if tit in SALTA or not corpo.strip(): continue
    viste.append(tit)
    p = P(RINOMINA.get(tit, tit), size=11.5, bold=True, colore=NAVY, prima=13, dopo=3); bordo_sotto(p)
    rendi_md(corpo)

# ═══════════ PIÈ DI PAGINA ═══════════
pf = sec.footer.paragraphs[0]
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pf.add_run(TITOLO_BREVE + '  —  dossier del GM,  pag. ')
r.font.size = Pt(8); r.font.color.rgb = GRIGIO; r.italic = True
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
rr = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
for tag, val in (('sz', '16'), ('color', '706A5E')):
    e = OxmlElement('w:' + tag); e.set(qn('w:val'), val); rpr.append(e)
rr.append(rpr); fld.append(rr); pf._p.append(fld)

doc.save(OUT)
print('creato:', OUT, '·', os.path.getsize(OUT)//1024, 'KB')
print('sezioni distretto riportate:', len(viste), '→', ' | '.join(viste))
if senza_scheda: print('!! senza scheda nel wizard:', ', '.join(senza_scheda))


# ═══ DOSSIER_GM.md per GMDASHBOARD: stesso contenuto del DOCX, rigenerato insieme ═══
OUT_MD = OUT[:-5] + '.md'
from docx import Document as _Doc
from docx.table import Table as _Tab
from docx.text.paragraph import Paragraph as _Par
_d = _Doc(OUT)
_md = []
def _cella(c): return ' '.join(c.text.split()).replace('|', chr(92)+'|')
def _riga_par(par):
    txt = par.text.strip()
    if not txt: return None
    if set(txt) <= set('━'): return '---'
    sizes = [r.font.size.pt for r in par.runs if r.font.size is not None]
    mx = max(sizes) if sizes else 10.5
    if mx >= 20: return '# ' + txt
    if any(r.bold and r.font.size is not None and r.font.size.pt >= 13 for r in par.runs): return '## ' + txt
    con_testo = [r for r in par.runs if r.text.strip()]
    if con_testo and all(r.bold for r in con_testo) and mx <= 10.5 and len(txt) < 70: return '### ' + txt
    pezzi = []
    for r in par.runs:
        t = r.text
        if not t: continue
        if (r.bold or r.italic) and t.strip():
            testa = t[:len(t)-len(t.lstrip())]
            coda = t[len(t.rstrip()):]
            mark = '**' if r.bold else '*'
            t = testa + mark + t.strip() + mark + coda
        pezzi.append(t)
    return ''.join(pezzi) if pezzi else txt
_skip_persone = False
for _el in _d.iter_inner_content():
    if isinstance(_el, _Par):
        _txt0 = _el.text.strip()
        if _txt0.startswith('LE PERSONE DEL CASO'):
            _md.append('## LE PERSONE DEL CASO   人物'); _md.append('')
            _md.append('*Le schede complete — una per file, con dentro tutto: aspetto, cosa sa, eventi, En, deposizione — sono in `PNG/`. Qui solo l' + chr(39) + 'indice.*'); _md.append('')
            for _p in ordinati:
                _md.append('- [[PNG/' + file_scheda(nome(_p['id']), _p['id']) + '|' + nome(_p['id']) + ']] — ' + (RUOLI.get(_p['id']) or ruolo_label(_p) or ''))
            _md.append('')
            _skip_persone = True
            continue
        if _skip_persone and _txt0.startswith('I LUOGHI'):
            _skip_persone = False
        if _skip_persone: continue
        _r = _riga_par(_el)
        if _r is not None: _md.append(_r); _md.append('')
    elif isinstance(_el, _Tab):
        if _skip_persone: continue
        _rt = [[_cella(c) for c in row.cells] for row in _el.rows]
        _rt = [r for r in _rt if any(x.strip() for x in r)]
        if not _rt: continue
        _nc = max(len(r) for r in _rt)
        if _nc == 2:
            for _lab, _val in _rt:
                _md.append(('- **' + _lab + '** — ' + _val) if _lab.strip() else ('- ' + _val))
        else:
            _md.append('| ' + ' | '.join(_rt[0]) + ' |')
            _md.append('|' + ' --- |' * _nc)
            for _r in _rt[1:]:
                _md.append('| ' + ' | '.join((_r + ['']*_nc)[:_nc]) + ' |')
        _md.append('')
open(OUT_MD, 'w', encoding='utf-8').write(chr(10).join(_md).strip() + chr(10))
print('creato:', OUT_MD, '·', os.path.getsize(OUT_MD)//1024, 'KB')