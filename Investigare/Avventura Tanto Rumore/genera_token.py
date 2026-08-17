# -*- coding: utf-8 -*-
# Token da tavolo per «L'era glaciale del lavoro»: un cartoncino per persona,
# col nome com'è scritto nel dossier e sotto chi è. 12 su una pagina A4, da ritagliare.
import io, sys, os, sqlite3, json, re, urllib.parse
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'C:\Public\_Clienti\Maruga\Giochi\Vampiri\Vault\Vampiri\Investigare'
WIZ  = os.path.join(BASE, 'Wizard', 'codice', 'GenkaiWizard')
ID   = 'C3C15FF7-AFCE-4299-A49C-53B367CD29EA'
OUT  = os.path.join(BASE, 'Avventura Tanto Rumore', 'Token', 'TOKEN_PERSONE.docx')
TMP  = os.path.join(__import__('tempfile').gettempdir(), 'genkai_token_foto')
os.makedirs(TMP, exist_ok=True)

c = sqlite3.connect('file:' + os.path.join(WIZ, 'app.db').replace('\\','/') + '?mode=ro', uri=True)
titolo_caso, stato = c.execute("select Titolo, StatoJson from Progetti where Id=?", (ID,)).fetchone()
S = json.loads(stato)
CAST   = {p['id']: p for p in S['cast']}
SCHEDE = {s['personaId']: s for s in S['passo8']['schede']}
COLP   = set(S['passo5'].get('colpevoliIds') or [])
VITT   = S['passo2'].get('personaId')
nome = lambda i: (CAST[i]['cognome'] + ' ' + CAST[i]['nome']).strip() if i in CAST else ''

# ── chi è ciascuno: la relazione scritta nel wizard ──
RUOLI = {}
for base in ('passo3', 'passo6'):
    for cerchio in ('famiglia', 'lavoro', 'amici', 'altri', 'intersezione'):
        for r in (S.get(base, {}).get(cerchio) or []):
            pid, rel = r.get('personaId'), (r.get('relazione') or '').strip()
            if pid and rel and pid not in RUOLI: RUOLI[pid] = rel[0].upper() + rel[1:]
RUOLI[VITT] = 'Vittima'
# il token sta sul tavolo davanti ai giocatori: al colpevole va l'identità pubblica,
# non la soluzione del caso (la nota del cast dice «Capo della banda»)
for pid in COLP:
    nota = re.sub(r'^\[ex Dati base\]\s*', '', (CAST[pid].get('note') or '')).strip()
    RUOLI[pid] = nota or CAST[pid].get('professione') or 'Sospettato'

def foto_di(pid):
    """il ritratto del wizard, ritagliato quadrato sul volto e alleggerito"""
    f = (SCHEDE.get(pid, {}) or {}).get('foto')
    if not f: return None
    rel = urllib.parse.unquote(f.split('?')[0]).lstrip('/')
    src = os.path.join(WIZ, 'wwwroot', *rel.split('/'))
    if not os.path.exists(src): return None
    im = Image.open(src).convert('RGB')
    w, h = im.size
    lato = min(w, h)
    im = im.crop(((w - lato)//2, 0, (w - lato)//2 + lato, lato))     # dall'alto: il volto
    im = im.resize((360, 360), Image.LANCZOS)
    dest = os.path.join(TMP, pid + '.jpg')
    im.save(dest, quality=88)
    return dest

# ordine: vittima, colpevole, poi il mondo della vittima, poi quello dell'assassino
def peso(p):
    pid = p['id']
    if pid == VITT: return (0, '')
    if pid in COLP: return (1, '')
    cerchio = p.get('cerchio', '')
    return (2 if cerchio.startswith('vittima') else 3, p['cognome'] + p['nome'])
persone = sorted(CAST.values(), key=peso)

# ── documento ──
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(1.0)
sec.left_margin = sec.right_margin = Cm(1.0)

ORO    = RGBColor(0xB4, 0x96, 0x50)
NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
GRIGIO = RGBColor(0x70, 0x6A, 0x5E)

st = doc.styles['Normal']
st.font.name = 'Garamond'; st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.0

COLONNE = 4
LARG, ALT = Cm(3.8), Cm(4.3)
FOTO_CM = Cm(2.5)

def bordi_taglio(cel, colore='C8C2B4'):
    """filo sottile su tutti i lati: è la linea su cui tagliare"""
    tcPr = cel._tc.get_or_add_tcPr()
    bd = OxmlElement('w:tcBorders')
    for lato in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + lato)
        for k, v in (('val','single'), ('sz','4'), ('color',colore), ('space','0')): e.set(qn('w:'+k), v)
        bd.append(e)
    tcPr.append(bd)

def margini_cella(cel, cm=0.08):
    tcPr = cel._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for lato in ('top', 'start', 'bottom', 'end'):
        e = OxmlElement('w:' + lato)
        e.set(qn('w:w'), str(int(cm * 567))); e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tcPr.append(mar)

def altezza_riga(riga, cm):
    trPr = riga._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), str(int(cm * 567))); h.set(qn('w:hRule'), 'exact')
    trPr.append(h)

def riga_p(cel, testo, size, bold=False, italic=False, colore=None, prima=0, dopo=0, eastasia=False):
    p = cel.add_paragraph() if cel.paragraphs[0].runs or cel.paragraphs[0].text else cel.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(prima); p.paragraph_format.space_after = Pt(dopo)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(testo); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if colore is not None: r.font.color.rgb = colore
    if eastasia:
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'Yu Mincho')
    return p

def cartoncino(cel, pers):
    pid = pers['id']
    bordi_taglio(cel); margini_cella(cel)
    cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cel.width = LARG
    f = foto_di(pid)
    if f:
        p = cel.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.add_run().add_picture(f, width=FOTO_CM, height=FOTO_CM)
    else:
        # senza ritratto: il cognome in kanji, grande, al posto della foto
        kanji = (pers.get('kanji') or '').split(' ')[0]
        p = riga_p(cel, kanji or '—', 30, colore=RGBColor(0xD8, 0xD2, 0xC4), prima=20, dopo=16, eastasia=True)
    riga_p(cel, nome(pid), 9.5, bold=True, colore=NAVY, prima=0, dopo=0)
    if pers.get('kanji'):
        riga_p(cel, pers['kanji'], 7, colore=GRIGIO, dopo=0, eastasia=True)
    ruolo = RUOLI.get(pid, '')
    corpo = 7 if len(ruolo) <= 16 else (6.3 if len(ruolo) <= 24 else 5.6)
    riga_p(cel, ruolo, corpo, italic=True, colore=ORO, dopo=0)

tab = doc.add_table(rows=0, cols=COLONNE)
tab.alignment = WD_TABLE_ALIGNMENT.CENTER
tab.autofit = False
lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tab._tbl.tblPr.append(lay)

for n in range(0, len(persone), COLONNE):
    riga = tab.add_row(); altezza_riga(riga, 4.3)
    for i in range(COLONNE):
        cel = riga.cells[i]; cel.width = LARG
        if n + i < len(persone): cartoncino(cel, persone[n + i])
        else: bordi_taglio(cel, 'EFEBE2')

# Word pretende un paragrafo dopo una tabella finale: lo do io, invisibile,
# se no ne aggiunge uno suo e la pagina diventa due
coda = doc.add_paragraph()
coda.paragraph_format.space_before = Pt(0); coda.paragraph_format.space_after = Pt(0)
coda.paragraph_format.line_spacing = Pt(1)
coda.add_run().font.size = Pt(1)

doc.save(OUT)
senza_foto = [nome(p['id']) for p in persone if not (SCHEDE.get(p['id'], {}) or {}).get('foto')]
print('creato:', OUT, '·', os.path.getsize(OUT)//1024, 'KB')
print('token:', len(persone), '· righe:', len(tab.rows))
print('senza ritratto (cognome in kanji):', ', '.join(senza_foto) if senza_foto else 'nessuno')
for p in persone: print(f"   {nome(p['id']):24s} → {RUOLI.get(p['id'],'')}")
