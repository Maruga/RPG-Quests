# -*- coding: utf-8 -*-
# Foglio da tavolo: lo scontro con pistola, manganello e pugni + coperture.
# Tutto ripreso da GENKAI_Combattimento.md v2.1 — nessuna regola aggiunta.
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'C:\Public\_Clienti\Maruga\Giochi\Vampiri\Vault\Vampiri\Investigare'
OUT  = os.path.join(BASE, 'Avventura Tanto Rumore', 'SCONTRO_FOGLIO_TAVOLO.docx')

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(1.0)
sec.left_margin = sec.right_margin = Cm(1.3)

ORO    = RGBColor(0xB4, 0x96, 0x50)
NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
GRIGIO = RGBColor(0x70, 0x6A, 0x5E)
ROSSO  = RGBColor(0x8B, 0x1A, 0x1A)
UTILE  = 18.4

st = doc.styles['Normal']
st.font.name = 'Garamond'; st.font.size = Pt(9.2)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
st.paragraph_format.space_after = Pt(2)
st.paragraph_format.line_spacing = 1.0

def bordo_sotto(p, colore='B49650', sz=6):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    for k, v in (('val','single'), ('sz',str(sz)), ('color',colore), ('space','2')): bt.set(qn('w:'+k), v)
    b.append(bt); pPr.append(b)

def run_md(p, testo, size=9.2, colore=None):
    """**grassetto** e *corsivo* dentro il paragrafo"""
    import re
    for pezzo in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', testo):
        if not pezzo: continue
        g = pezzo.startswith('**') and pezzo.endswith('**')
        c = not g and pezzo.startswith('*') and pezzo.endswith('*')
        r = p.add_run(pezzo.strip('*') if (g or c) else pezzo)
        r.font.size = Pt(size); r.bold = g; r.italic = c
        if colore is not None: r.font.color.rgb = colore

def P(testo='', size=9.2, dopo=2, prima=0, indent=0, colore=None, align=None):
    p = doc.add_paragraph()
    run_md(p, testo, size, colore)
    p.paragraph_format.space_after = Pt(dopo); p.paragraph_format.space_before = Pt(prima)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if align is not None: p.alignment = align
    return p

def titolo(testo, kanji='', prima=6):
    p = doc.add_paragraph()
    r = p.add_run(testo.upper()); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = NAVY
    if kanji:
        k = p.add_run('   ' + kanji); k.font.size = Pt(10.5); k.font.color.rgb = ORO
    p.paragraph_format.space_before = Pt(prima); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    bordo_sotto(p, sz=8)
    return p

def punto(testo, size=9.2, dopo=1):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.space_after = Pt(dopo)
    r = p.add_run('· '); r.font.size = Pt(size); r.font.color.rgb = ORO; r.bold = True
    run_md(p, testo, size)
    return p

def tabella(righe, largh, size=9.2):
    t = doc.add_table(rows=0, cols=len(largh)); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); t._tbl.tblPr.append(lay)
    for n, celle in enumerate(righe):
        cs = t.add_row().cells
        for i, testo in enumerate(celle):
            cs[i].width = Cm(largh[i])
            par = cs[i].paragraphs[0]; par.paragraph_format.space_after = Pt(1)
            run_md(par, testo, size)
            if n == 0:
                for r in par.runs: r.bold = True; r.font.color.rgb = GRIGIO
                bordo_sotto(par, sz=4)
    for riga in t.rows:
        for i, cel in enumerate(riga.cells): cel.width = Cm(largh[i])
    return t

def riquadro(testo, size=9.5):
    """la formula, incorniciata"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr')
    for lato in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + lato)
        for k, v in (('val','single'), ('sz','6'), ('color','B49650'), ('space','4')): e.set(qn('w:'+k), v)
        b.append(e)
    pPr.append(b)
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'FBF8F1'); pPr.append(sh)
    run_md(p, testo, size)
    return p

# ═══════════ testata ═══════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
r = p.add_run('GENKAI 限界 — LO SCONTRO'); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
r2 = p.add_run('   衝突'); r2.font.size = Pt(13); r2.font.color.rgb = ORO
P('Pistola · manganello · pugni · coperture — foglio da tavolo, dal manuale *Shōtotsu* v2.1',
  size=9, colore=GRIGIO, align=WD_ALIGN_PARAGRAPH.CENTER, dopo=4)

# ═══════════ lo scambio ═══════════
titolo('Lo scambio', prima=4)
P('Ogni scambio **tutti tirano 2d6**, e quel tiro fa due cose insieme:', dopo=2)
punto('**Iniziativa** = dado **più basso** + velocità dell\u2019arma → agisce chi ha il totale **più basso**')
punto('**Somma dei due dadi** = il tuo tiro: **attacco** se agisci, **difesa** se sei bersaglio')
P('Riesci se la somma è **≤ attributo**. Lo **scarto** è attributo − somma. '
  '**Chi subisce danno perde la propria azione** in quello scambio. '
  'Poi si ritira: l\u2019iniziativa si ridecide ogni volta.', prima=2)
P('Chi agisce può **fare altro** invece di attaccare: ripararsi, spostarsi, ricaricare. '
  'Parità di iniziativa: spareggia il secondo dado puro, poi le azioni sono simultanee.', size=9)

# ═══════════ armi ═══════════
titolo('Le tre armi')
tabella([
    ['Arma', 'Attributo', 'Velocità', 'Danno'],
    ['**Pugni** (Lotta — tutti ce l\u2019hanno a 1 dall\u2019accademia)', 'Presenza', '1 in mano · 1 sempre', '**1**'],
    ['**Manganello** (keibō 警棒)', 'Silenzio', '2 da sfoderare · 1 in mano', '**2**'],
    ['**Pistola / revolver**', 'Lucidità', '3 dalla fondina · 2 in mano · 4 ricarica · 0 puntata', '**3**'],
], [8.4, 2.6, 5.4, 1.6])
P('La velocità **si somma al dado basso**: pistola nella fondina = +3 (estrai e spari nello stesso scambio), '
  'in mano = **+2**. Il manganello sfoderato (1) batte la pistola in fondina (3) — e anche quella in mano (2): '
  'la lama è più rapida del grilletto non ancora puntato. La pistola torna davanti quando è **puntata** '
  '(*Sotto Tiro*: velocità 0, iniziativa = dado basso puro).', prima=3)
punto('**Pistola, fino a 3 colpi** in uno scambio, con **un tiro solo di 2d6** per tutta la sequenza. '
      '**Il tempo si somma**: 1º colpo a iniziativa +2, 2º a +4, 3º a +6. **La fretta si paga**: '
      '+2 alla somma se dichiari due colpi, +3 se tre — ma il malus **cala di 1 a ogni colpo dopo il primo** '
      '(il primo traccia la traiettoria, se il bersaglio è ancora lì). Se incassi danno, la sequenza si ferma lì.')
punto('**Sparare a contatto** (gli sei addosso): **+1/+2** alla somma — è caotico. A distanza lontana, idem.')
punto('Ricaricare è come sfoderare: «ricarico e sparo» usa la velocità di ricarica (pistola: dado basso **+4**).')

# ═══════════ danno ═══════════
titolo('Il danno')
riquadro('danno  =  scarto attaccante  +  danno arma  −  scarto difensore  −  Assorbe          (minimo 0)', 10.5)
P('Si perde in **Ki**. Chi attacca tira sull\u2019**attributo dell\u2019arma**; chi è bersaglio difende così:', dopo=2)
tabella([
    ['Ti trovi', 'Difendi su', 'Perché'],
    ['Sotto il fuoco di un\u2019arma', '**Lucidità**', 'ripararti senza farti prendere dal panico'],
    ['Un pugno, un coltello', '**Pazienza**', 'aspettare il momento giusto per deviare'],
    ['Il caos, un\u2019esplosione', '**Lucidità o Distacco**', 'buttarti via, o la freddezza di chi si defila'],
], [5.0, 4.2, 8.8])
P('Se la difesa **riesce**, il suo scarto assorbe. Se **fallisce**, assorbe zero — ma l\u2019Assorbe fisso vale lo stesso.', prima=2)

# ═══════════ coperture ═══════════
titolo('Copertura e giubbotto — l\u2019Assorbe')
P('L\u2019Assorbe è **fisso e automatico**: non si tira, vale sempre — anche quando la difesa fallisce — '
  'e vale **contro tutti** gli attacchi. (Se sono in tre contro uno: la difesa attiva copre **un solo** attacco a scelta, '
  'l\u2019Assorbe li copre tutti.)', dopo=3)
tabella([
    ['Protezione', 'Tempo per mettersi / indossare', 'Assorbe'],
    ['**Dietro un muro**', '1 — reazione istintiva, se è a portata', '**5**'],
    ['**Giubbotto antiproiettile**', '4 — un\u2019azione composta: si indossa **prima**', '**3**'],
    ['**Tavolo di legno ribaltato**', '2 — un movimento', '**1**'],
], [5.6, 10.4, 2.0])
P('*Altre coperture: il GM sceglie da 1 a 5 sulla scala tavolo → muro — portiera d’auto 2, pilastro di cemento 4.*', size=9, prima=1)
punto('**Il tempo conta**: quel numero è una velocità, come le armi. La copertura protegge **dal tuo momento '
      'd\u2019iniziativa in poi** — se l\u2019altro spara prima, ti trova ancora scoperto, e se ti fa danno perdi l\u2019azione: '
      'dietro quel muro non ci sei arrivato.')
punto('**Azione o difesa**: raggiungere una copertura è un\u2019**azione** (guadagni l\u2019Assorbe). Buttarsi *mentre ti sparano* '
      'è la normale **difesa** (la somma che assorbe). Una dichiarazione vale una delle due — il GM chiede quale.')
punto('Buio fitto o fumo: **+2/+3** alla somma di chi ti spara.')

# ═══════════ da non dimenticare ═══════════
titolo('Le cinque cose che si dimenticano')
for n, t in enumerate((
    '**Genkai sospeso** durante lo scontro: l\u2019adrenalina copre tutto. Si valuta **a fine scontro** — '
    'chi resta a **Ki ≤ 3** quando cala il silenzio, crolla lì.',
    '**Il danno non ha pavimento**: a **Ki 0 o sotto un PG è morto**. Il fermarsi a 1 vale per costi e pressioni, '
    'non per le ferite. I PNG invece hanno una **Riserva** (comparsa 3 · duro 6 · professionista 9): a 0 sono fuori '
    'combattimento, e come ne escono lo decidi tu.',
    '**Niente Nami né Kiwami, e il soroban non si muove.** Qui il critico è l\u20191+1: **+1d6** di danni se attacchi, '
    'di assorbimento se difendi. Il **6+6** vale solo quando **agisci** (mai in difesa): tira 1d6 sulla tabella degli '
    'imprevisti — arma inceppata, presa scivolata, arma a terra.',
    '**Stringere i denti**: se il danno netto è **1 o 2** puoi non perdere l\u2019azione pagando **1 punto dell\u2019attributo '
    'che stai usando** (non si può se è già a 4; il punto la notte si ripara). Da **3 in su** ti ferma e basta.',
    '**Dopo**: oltre al Genkai, il GM può chiedere **un** tiro di pressione (Distacco o Silenzio) a chi ha appena visto '
    'qualcosa di grosso. E un poliziotto che estrae lo scrive nel rapporto; che spara apre un fascicolo.',
), 1):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'{n}.  '); r.bold = True; r.font.size = Pt(9.2); r.font.color.rgb = ORO
    run_md(p, t, 9.2)

# ═══════════ esempio ═══════════
titolo('Uno scambio, per intero')
P('*PG: Lucidità 7, pistola in mano. Il tizio: Lucidità 6, pistola in mano anche lui, dietro un tavolo ribaltato (Assorbe 1).*', size=9, dopo=1)
P('Il PG tira **2 e 5** → iniziativa 2+2 = **4**; il tizio tira 4 e 3 → 3+2 = **5**: agisce il PG. '
  'Attacco: somma **7 ≤ Lucidità 7** → colpito, **scarto 0**, +3 di pistola = **3 in arrivo**. '
  'Il tizio difende (sotto il fuoco → Lucidità): somma **7 > 6**, fallisce e assorbe zero; il tavolo però toglie 1 → '
  '**prende 2 Ki**, e avendo subito danno **perde la sua azione**. '
  'Fosse stato dietro un **muro** (Assorbe 5), non gli sarebbe arrivato niente.', size=9)

doc.save(OUT)
print('creato:', OUT, '·', os.path.getsize(OUT)//1024, 'KB')
