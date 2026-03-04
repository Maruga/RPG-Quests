"""
Combina tutti gli handout HTML in un unico file Word con page break tra ogni handout.
Estrae il contenuto testuale mantenendo tabelle, elenchi e formattazione base.
"""

import os
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

HANDOUT_DIR = Path(__file__).parent
OUTPUT = HANDOUT_DIR / "Handout_Ultima_Cena_di_Tanaka.docx"

FILES = [
    "01_Planimetria_Villa.html",
    "02_Rapporto_Preliminare.html",
    "03_Lista_Presenti.html",
    "04_Ristrutturazione_Societaria.html",
    "05_Agenda_Tanaka.html",
    "06_Referto_Scientifica.html",
    "07_Menu_Kaiseki.html",
    "08_Biglietto_Nishida.html",
]

# ── Helpers ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_paragraph(doc, text, bold=False, italic=False, size=11,
                         color=None, align=None, space_after=6, space_before=0,
                         font_name="Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_header_block(doc, org_text, title_text, subtitle_text=""):
    """Add a dark header block similar to the HTML doc-header."""
    # Organization line
    add_styled_paragraph(doc, org_text, bold=False, size=8,
                         color=(192, 160, 96),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    # Title
    add_styled_paragraph(doc, title_text, bold=True, size=14,
                         color=(26, 26, 46),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if subtitle_text:
        add_styled_paragraph(doc, subtitle_text, italic=True, size=8,
                             color=(128, 128, 128),
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    # Separator line
    add_styled_paragraph(doc, "━" * 70, size=8, color=(192, 160, 96),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

def add_section_title(doc, text):
    add_styled_paragraph(doc, text.upper(), bold=True, size=11,
                         color=(26, 26, 46), space_before=12, space_after=6)
    # thin line
    add_styled_paragraph(doc, "─" * 60, size=6, color=(26, 26, 46),
                         space_after=6, space_before=0)

def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.left_indent = Cm(indent)
    return p

def html_table_to_word(doc, html_table, header_bg="1A1A2E"):
    """Convert an HTML <table> to a Word table."""
    rows_data = []
    for tr in html_table.find_all('tr'):
        cells = tr.find_all(['th', 'td'])
        row = []
        is_header = any(c.name == 'th' for c in cells)
        for c in cells:
            colspan = int(c.get('colspan', 1))
            text = c.get_text(separator='\n', strip=True)
            row.append((text, colspan, c.name == 'th', 'group-header' in (c.parent.get('class') or []),
                         'victim-row' in (c.parent.get('class') or []),
                         'highlight-row' in (c.parent.get('class') or []),
                         'result-positive' in (c.get('class') or []),
                         'result-negative' in (c.get('class') or [])))
        if row:
            rows_data.append(row)

    if not rows_data:
        return

    max_cols = max(sum(c[1] for c in row) for row in rows_data)
    tbl = doc.add_table(rows=0, cols=max_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    for row_data in rows_data:
        row = tbl.add_row()
        col_idx = 0
        for text, colspan, is_th, is_group, is_victim, is_highlight, is_pos, is_neg in row_data:
            if col_idx >= max_cols:
                break
            cell = row.cells[col_idx]
            # Merge if colspan > 1
            if colspan > 1 and col_idx + colspan - 1 < max_cols:
                cell.merge(row.cells[col_idx + colspan - 1])
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"

            if is_th:
                run.bold = True
                run.font.color.rgb = RGBColor(232, 224, 208)
                set_cell_bg(cell, header_bg)
            elif is_group:
                run.bold = True
                run.font.size = Pt(8)
                set_cell_bg(cell, "E8E4DA")
            elif is_victim:
                set_cell_bg(cell, "FFF0F0")
            elif is_highlight:
                run.bold = True
                set_cell_bg(cell, "FFFFF0")

            if is_pos:
                run.bold = True
                run.font.color.rgb = RGBColor(139, 0, 0)
            elif is_neg:
                run.font.color.rgb = RGBColor(46, 122, 46)

            col_idx += colspan

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_page_break(doc):
    doc.add_page_break()


# ── Handout processors ──────────────────────────────────

def process_01_planimetria(doc, soup):
    add_header_block(doc,
        "京都府警察本部 — POLIZIA PREFETTURALE DI KYOTO",
        "PLANIMETRIA — VILLA TANAKA",
        "Via Higashiyama-ku, Kyoto | Rilevamento del 14 novembre 1997 | Allegato A")

    add_section_title(doc, "Piano Terra — 一階")

    # Represent the floor plan as a structured text layout
    rooms_pt = [
        ("1. INGRESSO / GENKAN (玄関)", "Guardaroba ospiti — scarpe e giacche\n[ Giacca di Tanaka con EpiPen ]", False),
        ("2. SALA PRINCIPALE / OHIROMA (大広間)", "Tavolo da cena — zona salotto — impianto stereo\n[ Incenso con wintergreen ] — [ Bottiglia Juyondai ]", False),
        ("5. STUDIO / SHOSAI (書斎)", "Porta chiusa\n[ Documenti ristrutturazione ] — [ Agenda — n. Nishida ] — [ Cartellina Ogawa ]", False),
        ("", "— Corridoio di servizio — la musica dalla sala copre i rumori —", True),
        ("4. CUCINA / DAIDOKORO (台所)", "Cuoco Endo + 2 aiuti\n[ Menù kaiseki — pulito ]", False),
        ("3. ⚠ BAGNO OSPITI / OTEARAI (お手洗い) — SCENA DEL CRIMINE ⚠", "[ Sapone doppio strato ] — [ Graffi porta — chiavistello ]", True),
        ("10. GIARDINO / NIWA (庭)", "Accesso dalla sala — Yuki e Nakamura visti qui", False),
    ]

    tbl = doc.add_table(rows=len(rooms_pt), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, detail, highlight) in enumerate(rooms_pt):
        row = tbl.rows[i]
        c0 = row.cells[0]
        c1 = row.cells[1]
        r0 = c0.paragraphs[0].add_run(name)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.name = "Calibri"
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9)
        r1.font.name = "Calibri"
        if highlight:
            r0.font.color.rgb = RGBColor(139, 0, 0)
            set_cell_bg(c0, "FFF0F0")
            set_cell_bg(c1, "FFF0F0")

    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()

    add_styled_paragraph(doc,
        "Il bagno ospiti si trova in fondo al corridoio — lontano dal guardaroba e dalla sala principale",
        italic=True, size=9, color=(139, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)

    add_section_title(doc, "Primo Piano — 二階")

    rooms_pp = [
        ("7. CAMERA PADRONALE (寝室)", "[ Lettere Nakamura — comodino Yuki ]\n[ Farmaci antistaminici Tanaka ]"),
        ("8. BAGNO PADRONALE (浴室)", "[ Sapone con logo Tanaka Cosmetics ]\nConfronto con bagno ospiti"),
        ("9. CAMERA OSPITI (客室)", "Niente di rilevante"),
    ]

    tbl2 = doc.add_table(rows=len(rooms_pp), cols=2)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, detail) in enumerate(rooms_pp):
        row = tbl2.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(name)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.name = "Calibri"
        r1 = row.cells[1].paragraphs[0].add_run(detail)
        r1.font.size = Pt(9)
        r1.font.name = "Calibri"
    for row in tbl2.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()

    add_section_title(doc, "Legenda")
    legend = [
        ("🔴 Rosso", "Scena del crimine (accesso ristretto)"),
        ("🟡 Giallo", "Stanza con indizi rilevanti"),
        ("🟢 Verde", "Esterno"),
        ("⬜ Grigio", "Area di passaggio"),
    ]
    for sym, desc in legend:
        add_styled_paragraph(doc, f"{sym}  {desc}", size=9, space_after=2)

    add_styled_paragraph(doc,
        "Allegato A — Rilievo planimetrico Villa Tanaka — Fascicolo 97-KPD-1114 — Riservato",
        size=8, color=(160, 160, 160), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)


def process_generic(doc, soup, header_bg="1A1A2E"):
    """Generic processor: extract header, sections, tables, lists from HTML."""
    # Header
    header = soup.find(class_='doc-header')
    if header:
        org = header.find(class_='org')
        title = header.find(class_='title')
        subtitle = header.find(class_='subtitle')
        add_header_block(doc,
            org.get_text(strip=True) if org else "",
            title.get_text(strip=True) if title else "",
            subtitle.get_text(strip=True) if subtitle else "")
    else:
        # Check for other header styles (menu, agenda, company)
        doc_header = soup.find(class_='menu-header') or soup.find(class_='agenda-header') or soup.find(class_='doc-header')
        if doc_header:
            texts = [t.get_text(strip=True) for t in doc_header.children
                     if hasattr(t, 'get_text') and t.get_text(strip=True)]
            for i, t in enumerate(texts):
                add_styled_paragraph(doc, t, bold=(i == 1 or i == 0),
                                     size=14 if i <= 1 else 9,
                                     align=WD_ALIGN_PARAGRAPH.CENTER,
                                     space_after=4)
            add_styled_paragraph(doc, "━" * 70, size=8, color=(192, 160, 96),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Classificazione
    classif = soup.find(class_='classificazione')
    if classif:
        add_styled_paragraph(doc, classif.get_text(strip=True), bold=True,
                             size=9, color=(139, 0, 0),
                             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)

    # Confidential
    conf = soup.find(class_='confidential')
    if conf:
        add_styled_paragraph(doc, conf.get_text(strip=True), italic=True,
                             size=8, color=(160, 160, 160),
                             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)

    # Metadata grid
    meta = soup.find(class_='metadata')
    if meta:
        divs = meta.find_all('div', recursive=False)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i in range(0, len(divs), 2):
            row = tbl.add_row()
            for j in range(2):
                if i + j < len(divs):
                    text = divs[i+j].get_text(separator=': ', strip=True)
                    cell = row.cells[j]
                    run = cell.paragraphs[0].add_run(text)
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Process main content - sections, tables, special boxes
    document_div = soup.find(class_='document') or soup.find(class_='menu-inner') or soup.find('body')

    for element in document_div.children:
        if isinstance(element, NavigableString):
            continue
        if not hasattr(element, 'get'):
            continue

        classes = element.get('class', [])

        # Skip already processed
        if any(c in classes for c in ['doc-header', 'metadata', 'classificazione',
                                       'confidential', 'timbro', 'footer', 'hanko',
                                       'signature-section', 'signature-area',
                                       'menu-header', 'agenda-header']):
            continue

        # Section
        if 'section' in classes:
            sec_title = element.find(class_='section-title')
            if sec_title:
                add_section_title(doc, sec_title.get_text(strip=True))

            for child in element.children:
                if isinstance(child, NavigableString):
                    continue
                if not hasattr(child, 'get'):
                    continue
                child_classes = child.get('class', [])

                if 'section-title' in child_classes:
                    continue  # already done
                elif 'subsection-title' in child_classes:
                    add_styled_paragraph(doc, child.get_text(strip=True),
                                         bold=True, size=10, space_before=6)
                elif child.name == 'p':
                    add_styled_paragraph(doc, child.get_text(strip=True), size=10)
                elif child.name == 'ul':
                    for li in child.find_all('li'):
                        add_bullet(doc, li.get_text(strip=True))
                elif child.name == 'table':
                    html_table_to_word(doc, child, header_bg)
                elif 'critical-box' in child_classes or 'highlight-box' in child_classes or 'note-box' in child_classes:
                    text = child.get_text(separator=' ', strip=True)
                    is_critical = 'critical-box' in child_classes
                    p = add_styled_paragraph(doc, text, bold=is_critical, size=10,
                                             color=(139, 0, 0) if is_critical else (100, 80, 40),
                                             space_before=6)
                    # Left border effect via indentation
                    p.paragraph_format.left_indent = Cm(0.5)
                elif 'diagram' in child_classes:
                    # Soap diagram
                    add_styled_paragraph(doc, "[ Diagramma sezione trasversale del sapone ]",
                                         bold=True, size=10,
                                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
                    add_styled_paragraph(doc, "STRATO ESTERNO: Olio di camelia — neutro — innocuo — parzialmente consumato",
                                         size=9, color=(46, 122, 46),
                                         align=WD_ALIGN_PARAGRAPH.CENTER)
                    add_styled_paragraph(doc, "- - - - - - - - - - - -", size=8,
                                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
                    add_styled_paragraph(doc, "STRATO INTERNO: Metil-salicilato (wintergreen) — concentrazione elevata",
                                         size=9, color=(139, 0, 0),
                                         align=WD_ALIGN_PARAGRAPH.CENTER)

        # Decision box
        elif 'decision-box' in classes:
            text = element.get_text(separator=' ', strip=True)
            p = add_styled_paragraph(doc, text, bold=True, size=10, space_before=8)
            p.paragraph_format.left_indent = Cm(0.5)

        # Conclusion box
        elif 'conclusion-box' in classes:
            text = element.get_text(separator=' ', strip=True)
            p = add_styled_paragraph(doc, text, bold=True, size=10,
                                     color=(26, 46, 26), space_before=8)
            p.paragraph_format.left_indent = Cm(0.5)

        # Section title (standalone)
        elif 'section-title' in classes:
            add_section_title(doc, element.get_text(strip=True))

        # Orgchart
        elif 'orgchart' in classes:
            add_styled_paragraph(doc, "[ Organigramma ]", bold=True, size=10,
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
            for box in element.find_all(class_='orgchart-box'):
                text = box.get_text(separator=' — ', strip=True)
                is_pres = 'president' in (box.get('class') or [])
                is_prom = 'promoted' in (box.get('class') or [])
                color = (26, 26, 46) if is_pres else ((192, 160, 96) if is_prom else (100, 100, 100))
                add_styled_paragraph(doc, text, bold=is_pres or is_prom, size=10,
                                     color=color, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Standalone table
        elif element.name == 'table':
            html_table_to_word(doc, element, header_bg)

        # Standalone paragraph
        elif element.name == 'p':
            add_styled_paragraph(doc, element.get_text(strip=True), size=10)

        # Standalone list
        elif element.name == 'ul':
            for li in element.find_all('li'):
                add_bullet(doc, li.get_text(strip=True))

        # Notes area
        elif 'notes-area' in classes:
            title_el = element.find(class_='notes-title')
            if title_el:
                add_section_title(doc, title_el.get_text(strip=True))
            for _ in range(5):
                add_styled_paragraph(doc, "_" * 80, size=8, color=(200, 200, 200),
                                     space_after=8)

    # Footer
    footer = soup.find(class_='footer')
    if footer:
        add_styled_paragraph(doc, footer.get_text(strip=True), size=8,
                             color=(160, 160, 160),
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16)


def process_05_agenda(doc, soup):
    """Special processor for the agenda handout."""
    # Header
    header = soup.find(class_='agenda-header')
    if header:
        month = header.find(class_='agenda-month')
        year = header.find(class_='agenda-year')
        add_styled_paragraph(doc, month.get_text(strip=True) if month else "",
                             size=10, color=(90, 48, 16),
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_styled_paragraph(doc, year.get_text(strip=True) if year else "",
                             bold=True, size=20, color=(90, 48, 16),
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Weeks
    for week in soup.find_all(class_='week-section'):
        wh = week.find(class_='week-header')
        if wh:
            add_styled_paragraph(doc, wh.get_text(strip=True), bold=True, size=9,
                                 color=(139, 69, 19), space_before=8, space_after=4)

        for entry in week.find_all(class_='day-entry'):
            date_el = entry.find(class_='day-date')
            content_el = entry.find(class_='day-content')
            date_text = date_el.get_text(strip=True) if date_el else ""
            content_text = content_el.get_text(strip=True) if content_el else ""

            if not content_text:
                add_styled_paragraph(doc, f"{date_text:<12}", size=9,
                                     color=(90, 48, 16), space_after=1)
                continue

            is_red = content_el and 'handwritten-red' in (content_el.get('class') or [])
            color = (139, 0, 0) if is_red else (26, 26, 110)

            p = doc.add_paragraph()
            r_date = p.add_run(f"{date_text:<12}")
            r_date.bold = True
            r_date.font.size = Pt(9)
            r_date.font.color.rgb = RGBColor(90, 48, 16)
            r_date.font.name = "Calibri"
            r_content = p.add_run(content_text)
            r_content.font.size = Pt(10)
            r_content.font.color.rgb = RGBColor(*color)
            r_content.italic = True
            r_content.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)

        # Circled note
        circled = week.find(class_='circled-note')
        if circled:
            text = circled.get_text(separator=' — ', strip=True)
            p = add_styled_paragraph(doc, f"‼ {text}", bold=True, size=10,
                                     color=(139, 0, 0), space_before=4)
            p.paragraph_format.left_indent = Cm(1)

    # Business card
    card = soup.find(class_='business-card-attached')
    if card:
        add_styled_paragraph(doc, "━" * 40, size=8,
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
        add_styled_paragraph(doc, "[ Biglietto da visita incollato alla pagina ]",
                             italic=True, size=9, color=(128, 128, 128),
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        card_data = [
            ("Nome", card.find(class_='card-name').get_text(strip=True) if card.find(class_='card-name') else ""),
            ("Nome JP", card.find(class_='card-name-jp').get_text(strip=True) if card.find(class_='card-name-jp') else ""),
            ("Titolo", card.find(class_='card-title').get_text(strip=True) if card.find(class_='card-title') else ""),
            ("Telefono", card.find(class_='card-phone').get_text(strip=True) if card.find(class_='card-phone') else ""),
            ("Indirizzo", card.find(class_='card-address').get_text(strip=True) if card.find(class_='card-address') else ""),
        ]
        for i, (label, value) in enumerate(card_data):
            row = tbl.rows[i]
            r0 = row.cells[0].paragraphs[0].add_run(label)
            r0.bold = True
            r0.font.size = Pt(9)
            r0.font.name = "Calibri"
            r1 = row.cells[1].paragraphs[0].add_run(value)
            r1.font.size = Pt(9)
            r1.font.name = "Calibri"
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(8)


def process_07_menu(doc, soup):
    """Special processor for the kaiseki menu."""
    # Header
    menu_header = soup.find(class_='menu-header')
    if menu_header:
        date = menu_header.find(class_='menu-date')
        title_jp = menu_header.find(class_='menu-title-jp')
        title = menu_header.find(class_='menu-title')
        subtitle = menu_header.find(class_='menu-subtitle')

        if date:
            add_styled_paragraph(doc, date.get_text(strip=True), size=9,
                                 color=(160, 160, 160),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        if title_jp:
            add_styled_paragraph(doc, title_jp.get_text(strip=True), bold=True, size=22,
                                 color=(44, 44, 44),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        if title:
            add_styled_paragraph(doc, title.get_text(strip=True), size=11,
                                 color=(139, 115, 85),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        if subtitle:
            add_styled_paragraph(doc, subtitle.get_text(strip=True), italic=True,
                                 size=9, color=(170, 170, 170),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_styled_paragraph(doc, "━" * 30, size=8, color=(192, 160, 96),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Courses
    for course in soup.find_all(class_='course'):
        name_jp = course.find(class_='course-name-jp')
        name = course.find(class_='course-name')
        desc = course.find(class_='course-desc')

        if name_jp:
            add_styled_paragraph(doc, name_jp.get_text(strip=True), bold=True, size=12,
                                 color=(44, 44, 44),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
        if name:
            add_styled_paragraph(doc, name.get_text(strip=True), size=9,
                                 color=(192, 160, 96),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        if desc:
            text = desc.get_text(separator='\n', strip=True)
            add_styled_paragraph(doc, text, italic=True, size=10,
                                 color=(85, 85, 85),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_styled_paragraph(doc, "━" * 30, size=8, color=(192, 160, 96),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Sake section
    sake = soup.find(class_='sake-section')
    if sake:
        sake_jp = sake.find(class_='sake-title-jp')
        sake_title = sake.find(class_='sake-title')
        if sake_jp:
            add_styled_paragraph(doc, sake_jp.get_text(strip=True), bold=True, size=14,
                                 color=(44, 44, 44),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        if sake_title:
            add_styled_paragraph(doc, sake_title.get_text(strip=True), size=9,
                                 color=(192, 160, 96),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        for item in sake.find_all(class_='sake-item'):
            add_styled_paragraph(doc, item.get_text(strip=True), size=10,
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
            detail = item.find_next_sibling(class_='sake-detail')
            if detail:
                add_styled_paragraph(doc, detail.get_text(strip=True), italic=True,
                                     size=9, color=(136, 136, 136),
                                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Footer
    mf = soup.find(class_='menu-footer')
    if mf:
        for child in mf.children:
            if hasattr(child, 'get_text') and child.get_text(strip=True):
                text = child.get_text(strip=True)
                add_styled_paragraph(doc, text, size=8,
                                     color=(170, 170, 170),
                                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)


def process_08_biglietto(doc, soup):
    """Special processor for the business card handout."""
    add_styled_paragraph(doc, "✂ Ritagliare lungo i bordi ✂", size=8,
                         color=(170, 170, 170),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_section_title(doc, "Fronte")
    front = soup.find(class_='meishi-front')
    if front:
        name_jp = front.find(class_='name-jp')
        name_r = front.find(class_='name-romaji')
        title_jp = front.find(class_='title-jp')
        title_it = front.find(class_='title-it')
        license_el = front.find(class_='license')

        if name_jp:
            add_styled_paragraph(doc, name_jp.get_text(strip=True), bold=True, size=18,
                                 color=(26, 26, 46),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        if name_r:
            add_styled_paragraph(doc, name_r.get_text(strip=True), size=9,
                                 color=(136, 136, 136),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        if title_jp:
            add_styled_paragraph(doc, title_jp.get_text(strip=True), size=10,
                                 color=(68, 68, 68),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
        if title_it:
            add_styled_paragraph(doc, title_it.get_text(strip=True), size=9,
                                 color=(150, 150, 150),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        if license_el:
            add_styled_paragraph(doc, license_el.get_text(strip=True), size=7,
                                 color=(170, 170, 170),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_section_title(doc, "Retro — Contatti")
    back = soup.find(class_='meishi-back')
    if back:
        for cl in back.find_all(class_='contact-line'):
            label = cl.find(class_='contact-label')
            value = cl.find(class_='contact-value')
            text = f"{label.get_text(strip=True)}: {value.get_text(strip=True)}" if label and value else cl.get_text(strip=True)
            add_styled_paragraph(doc, text, size=10,
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

        addr = back.find(class_='address-section')
        if addr:
            add_styled_paragraph(doc, addr.get_text(separator='\n', strip=True),
                                 size=9, color=(100, 100, 100),
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)

        hw = back.find(class_='handwritten-note')
        if hw:
            add_styled_paragraph(doc, f"✐ {hw.get_text(strip=True)}", italic=True,
                                 size=10, color=(26, 26, 110),
                                 align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=8)

    add_styled_paragraph(doc, "✂ Ritagliare lungo i bordi ✂", size=8,
                         color=(170, 170, 170),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)


# ── Main ─────────────────────────────────────────────────

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set A4 margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "L'ULTIMA CENA DI TANAKA", bold=True, size=24,
                         color=(26, 26, 46),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "田中の最後の晩餐", size=16,
                         color=(100, 100, 100),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_styled_paragraph(doc, "━" * 40, size=10, color=(192, 160, 96),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_styled_paragraph(doc, "HANDOUT PER I GIOCATORI", bold=True, size=14,
                         color=(26, 26, 46),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_styled_paragraph(doc, "GENKAI 限界 v1.2", size=11,
                         color=(100, 100, 100),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Table of contents
    add_styled_paragraph(doc, "CONTENUTO", bold=True, size=12,
                         color=(26, 26, 46),
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    toc_items = [
        "01 — Planimetria Villa Tanaka",
        "02 — Rapporto Preliminare di Intervento",
        "03 — Elenco Persone Presenti",
        "04 — Ristrutturazione Societaria Tanaka Cosmetics",
        "05 — Agenda di Tanaka (Settembre 1997)",
        "06 — Referto Analisi Scientifica",
        "07 — Menù Kaiseki della Cena",
        "08 — Biglietto da Visita Nishida Goro",
    ]
    for item in toc_items:
        add_styled_paragraph(doc, item, size=11, space_after=4,
                             align=WD_ALIGN_PARAGRAPH.CENTER)

    # Process special handlers
    SPECIAL = {
        "01_Planimetria_Villa.html": process_01_planimetria,
        "05_Agenda_Tanaka.html": process_05_agenda,
        "07_Menu_Kaiseki.html": process_07_menu,
        "08_Biglietto_Nishida.html": process_08_biglietto,
    }

    for filename in FILES:
        filepath = HANDOUT_DIR / filename
        if not filepath.exists():
            print(f"WARN: {filename} not found, skipping")
            continue

        with open(filepath, 'r', encoding='utf-8') as f:
            html = f.read()

        soup = BeautifulSoup(html, 'html.parser')

        # Page break before each handout
        add_page_break(doc)

        if filename in SPECIAL:
            SPECIAL[filename](doc, soup)
        elif filename == "06_Referto_Scientifica.html":
            process_generic(doc, soup, header_bg="1A2E1A")
        else:
            process_generic(doc, soup)

    doc.save(str(OUTPUT))
    print(f"OK — Salvato: {OUTPUT}")


if __name__ == "__main__":
    main()
