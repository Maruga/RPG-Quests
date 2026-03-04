"""
Genera schede PNG ricorrenti in stile giapponese elegante per GENKAI 限界.
Pagina 1: Statistiche, Ki (fisso), Gou (fisso)
Pagina 2: Background, problema, PNG del problema, conoscenza, comportamento
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

OUT_DIR = Path(__file__).parent

# ── Colors ──
DARK = RGBColor(26, 26, 46)       # #1A1A2E
GOLD = RGBColor(180, 150, 80)     # #B49650
MUTED = RGBColor(120, 120, 120)   # #787878
RED = RGBColor(139, 0, 0)         # #8B0000
WHITE = RGBColor(255, 255, 255)
CREAM_HEX = "FAF8F0"
DARK_HEX = "1A1A2E"
GOLD_HEX = "B49650"
LIGHT_HEX = "F5F3EB"

# ── Helpers ──

def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            sz, color = val
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{sz}" '
                f'w:space="0" w:color="{color}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_borders(table, color="1A1A2E", size=4):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_run(paragraph, text, bold=False, italic=False, size=10, color=DARK, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run

def add_para(doc, text="", bold=False, italic=False, size=10, color=DARK,
             align=None, space_after=4, space_before=0, font="Calibri"):
    p = doc.add_paragraph()
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color, font=font)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_separator(doc, char="━", length=50, color=GOLD):
    add_para(doc, char * length, size=7, color=color,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, space_before=6)

def add_thin_separator(doc, color=GOLD):
    add_para(doc, "─" * 65, size=5, color=color,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=4)

def add_section_header(doc, text, kanji=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(2)
    if kanji:
        add_run(p, f"{kanji}  ", size=11, color=GOLD, font="Calibri")
    add_run(p, text.upper(), bold=True, size=10, color=DARK, font="Calibri")
    add_para(doc, "─" * 55, size=5, color=GOLD,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, space_before=0)


def add_gou_box(doc, name, kanji_desc, attribute, cost, success, failure):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(7)
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.3)
    add_run(p, "▸ ", size=9, color=GOLD)
    add_run(p, name, bold=True, size=10, color=DARK)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.left_indent = Cm(0.6)
    add_run(p2, kanji_desc, italic=True, size=9, color=MUTED)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(1)
    p3.paragraph_format.left_indent = Cm(0.6)
    add_run(p3, "Attributo: ", bold=True, size=8, color=MUTED)
    add_run(p3, attribute, size=8, color=DARK)
    add_run(p3, "  ·  ", size=8, color=GOLD)
    add_run(p3, "Costo: ", bold=True, size=8, color=MUTED)
    add_run(p3, cost, size=8, color=RED)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(1)
    p4.paragraph_format.left_indent = Cm(0.6)
    add_run(p4, "Successo: ", bold=True, size=8, color=MUTED)
    add_run(p4, success, size=8, color=DARK)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(4)
    p5.paragraph_format.left_indent = Cm(0.6)
    add_run(p5, "Fallimento: ", bold=True, size=8, color=MUTED)
    add_run(p5, failure, size=8, color=DARK)


# ── PNG Data ──

PNG_DATA = [
    # ── PNG_01 — TANIGUCHI Kenji ──
    {
        "filename": "Scheda_PNG_01_Taniguchi_Kenji.docx",
        "photo": "PNG_01_Taniguchi_Kenji_foto.png",
        "name_jp": "谷口 健二",
        "name": "TANIGUCHI Kenji",
        "age": "56",
        "role": "Commissario, Sezione Omicidi",
        "rank": "Keibu (警部) — Commissario",
        "years": "34",
        "ki": "9",
        "attrs": [
            ("Distacco", "7", "Ha visto tutto, niente lo sorprende più"),
            ("Pazienza", "6", "Impaziente per natura, temperato dall'esperienza"),
            ("Silenzio", "8", "Non mostra mai le sue carte — la voce non tradisce nulla"),
            ("Lucidità", "7", "Trent'anni di casi gli hanno affinato lo sguardo"),
            ("Ascolto", "5", "Più comandante che ascoltatore — delega l'ascolto ai suoi"),
            ("Presenza", "8", "Riempie la stanza senza alzare la voce"),
        ],
        "gou": {
            "name": "Voce del Tuono",
            "kanji": "雷の声",
            "desc": "La tua presenza diventa schiacciante. Il PNG sente il peso della tua autorità.",
            "attr": "Presenza (8)", "cost": "2 Ki",
            "success": "Il PNG è completamente sottomesso, cede o si ritira",
            "failure": "Il PNG è intimidito e scosso, ma non cede completamente",
        },
        "background": "Sei un veterano. Hai scalato i ranghi della polizia in trent'anni di servizio, da pattuglia a commissario. Hai risolto casi importanti, hai protetto la tua squadra, hai tenuto a bada media e politici.\n\nSei la roccia della Sezione Omicidi. I tuoi uomini ti rispettano. I tuoi superiori si fidano. Nessuno vede le crepe.\n\nMa le crepe ci sono. A casa, tua moglie Fumiko sta peggiorando. La demenza precoce sta avanzando. E tu fingi che vada tutto bene.",
        "problem_title": "La Moglie",
        "problem_text": "Tua moglie Taniguchi Fumiko ha iniziato a dimenticare le cose. Prima le chiavi, poi i nomi, poi gli appuntamenti. Tre mesi fa il dottore ha detto la parola che non volevi sentire: demenza precoce. Ha 54 anni.\n\nNon l'hai detto a nessuno. Al lavoro sei lo stesso di sempre. Ma la sera torni a casa e trovi i fornelli accesi, le pentole vuote, Fumiko che sorride come se niente fosse.\n\nIl problema: se al lavoro scoprono della malattia di Fumiko, qualcuno potrebbe usarla per metterti in pensione anticipata. E ogni ora che resti in ufficio, ti chiedi se Fumiko sta bene.",
        "png_name": "TANIGUCHI Fumiko",
        "png_role": "Moglie da 30 anni",
        "png_age": "54",
        "png_job": "Ex insegnante di calligrafia, ora a casa",
        "png_wants": "Che Kenji smetta di fingere che vada tutto bene. Non ha paura della malattia — ha paura di affrontarla da sola.",
        "png_behavior": "Nei momenti lucidi è la donna forte e ironica di sempre. Nei momenti di vuoto si confonde, ripete le domande, si arrabbia con se stessa. Non fa scene — si vergogna.",
        "enja_name": "MORIYAMA Genzo",
        "enja_role": "Procuratore Capo in pensione",
        "enja_desc": "Ex procuratore capo di Kyoto. Ha lavorato con Taniguchi per vent'anni. Conoscenze nel mondo legale e politico, sa come funziona il sistema.",
        "enja_contact": "Telefono di casa. Risponde sempre, anche se finge di essere seccato.",
        "enja_cost": "Nulla di materiale — si aspetta che Taniguchi faccia sempre la cosa giusta.",
        "tatemae": "Il commissario. Freddo, efficiente, inattaccabile. Le decisioni sono prese, i dubbi non esistono.",
        "honne": "Sei esausto. La notte dormi poco — ascolti Fumiko muoversi per casa. Ti senti in colpa per ogni ora passata in ufficio.",
        "phrase": "\"Abbiamo un lavoro da fare.\" (la dice anche a se stesso)",
        "pressure": "Ti chiudi ancora di più. Fumi di più. Le frasi diventano più corte. Il tamburellare accelera.",
        "weakness": "Qualsiasi riferimento alla malattia, alla vecchiaia, al perdere il controllo.",
        "district": [
            ("I PG", "Superiore diretto. Li rispetta se lavorano bene. Diretto, esigente, mai crudele."),
            ("Yamada Tetsuo", "Lo conosce come agente affidabile. \"Yamada è un bravo poliziotto.\""),
            ("Ito Daisuke", "Rispetto professionale reciproco. \"Se lo dice Ito, è così.\""),
        ],
        "traits": {
            "vizio": "Fuma — Seven Stars, sempre lo stesso pacchetto. Non ha mai provato a smettere.",
            "tic": "Tamburella le dita sulla scrivania quando pensa. Più tamburella, più è serio.",
            "oggetto": "Accendino Zippo in ottone, regalo di Fumiko. Inciso: \"K+F 1967\".",
            "gusto": "Tè verde bollente, sempre. Mai caffè. \"Il caffè è per chi ha fretta.\"",
            "extra": ("Rituale", "Ogni mattina arriva in ufficio mezz'ora prima di tutti. Legge i rapporti della notte in silenzio."),
        },
    },
    # ── PNG_02 — YAMADA Tetsuo ──
    {
        "filename": "Scheda_PNG_02_Yamada_Tetsuo.docx",
        "photo": "PNG_02_Yamada_Tetsuo_foto.png",
        "name_jp": "山田 哲夫",
        "name": "YAMADA Tetsuo",
        "age": "35",
        "role": "Accompagnatore e supporto operativo",
        "rank": "Junsa-buchō (巡査部長) — Sergente",
        "years": "10",
        "ki": "8",
        "attrs": [
            ("Distacco", "6", "Solido, ma non freddo — le cose lo toccano"),
            ("Pazienza", "8", "Aspetta senza agitarsi, mai. Il suo tratto migliore"),
            ("Silenzio", "7", "Calmo, voce ferma, non reagisce d'impulso"),
            ("Lucidità", "7", "Buon occhio da poliziotto di strada — nota i dettagli"),
            ("Ascolto", "5", "Ascolta per obbedire, non per leggere le persone"),
            ("Presenza", "6", "Rispettato ma non imponente"),
        ],
        "gou": {
            "name": "Palazzo della Memoria",
            "kanji": "記憶の宮殿",
            "desc": "Puoi richiamare con precisione fotografica qualcosa che hai visto o sentito.",
            "attr": "Lucidità (7)", "cost": "2 Ki",
            "success": "Ricordi il dettaglio e anche elementi periferici che non avevi notato consciamente",
            "failure": "Ricordi il dettaglio principale, ma sfocato o incompleto",
        },
        "background": "Sei poliziotto da dieci anni. Hai iniziato come agente di pattuglia e lo sei ancora — non per mancanza di capacità, ma perché ami la strada. O almeno, è quello che dici.\n\nConosci Kyoto come le tue tasche. Ogni quartiere, ogni scorciatoia, ogni izakaya dove si trovano informazioni. Quando la Sezione Omicidi apre un caso, il commissario Taniguchi ti assegna come supporto.\n\nSei quello che facilita tutto: spostamenti, contatti, archivi, territorio. Un bravo poliziotto. Ma non un detective.",
        "problem_title": "L'Esame",
        "problem_text": "Hai fallito l'esame per detective due volte. La prima a 28 anni, la seconda a 31. Non per stupidità — per ansia. Davanti al foglio d'esame ti si blocca tutto.\n\nLa tua fidanzata Kawano Miki ti vuole bene, ma è stanca di aspettare. Ha 32 anni, vuole sposarsi, vuole una famiglia. E un agente di pattuglia che non avanza di grado non è il futuro che aveva in mente.\n\nMiki vuole che riprovi l'esame. Tu hai paura di fallire una terza volta. E ogni volta che i PG ti trattano da \"semplice agente\", la ferita si riapre.",
        "png_name": "KAWANO Miki",
        "png_role": "Fidanzata da 4 anni, conviventi da 2",
        "png_age": "32",
        "png_job": "Impiegata alla Posta Centrale di Kyoto",
        "png_wants": "Che Tetsuo riprovi l'esame. Non per i soldi — per lui, per il futuro, per dimostrare che vale.",
        "png_behavior": "Non fa scenate. È pratica, diretta, un po' tagliente quando è stanca. Quando Tetsuo rientra tardi, sospira forte e non dice nulla — il che è peggio.",
        "enja_name": "KATO Jiro",
        "enja_role": "Proprietario di izakaya, ex informatore",
        "enja_desc": "Proprietario dell'izakaya \"Tsukimi\" a Gion. Ex informatore della polizia. Conosce voci dalla strada, movimenti nel quartiere.",
        "enja_contact": "Di persona, al locale. Non al telefono — \"le pareti hanno orecchie\".",
        "enja_cost": "Che la polizia non ficchi il naso nel suo locale.",
        "tatemae": "L'agente perfetto. Affidabile, discreto, sempre un passo indietro. Mai una lamentela.",
        "honne": "Ti chiedi se sarai un agente di pattuglia per tutta la vita. La notte ripassi mentalmente le domande dell'esame che non hai saputo rispondere.",
        "phrase": "\"Se posso permettermi...\" (prima di ogni opinione, come a chiedere il permesso di esistere)",
        "pressure": "Diventi più silenzioso. Annuisci di più. Scrivi di più sul taccuino. Stringi il taccuino più forte.",
        "weakness": "Qualsiasi commento sulla sua posizione, sul fatto che \"è solo un agente\", o paragoni con detective più giovani.",
        "district": [
            ("I PG", "Li supporta come superiori. Li rispetta. \"A disposizione, signore.\""),
            ("Commissario Taniguchi", "Lo rispetta profondamente. Taniguchi sa dell'esame e non ne parla."),
            ("Ito Daisuke", "Lo conosce dalla centrale. \"Ito-san è il migliore nel suo campo.\""),
        ],
        "traits": {
            "vizio": "Onigiri al tonno dal konbini. Ne mangia almeno due al giorno.",
            "tic": "Si tocca il taccuino nella tasca interna quando è nervoso.",
            "oggetto": "Il taccuino a spirale. Ne ha riempiti una ventina in dieci anni.",
            "gusto": "Ramen miso con extra chashu dalla bancarella di fronte alla centrale.",
            "extra": ("Rituale", "Ogni mattina controlla l'auto di servizio — olio, pneumatici, carburante."),
        },
    },
    # ── PNG_03 — ITO Daisuke ──
    {
        "filename": "Scheda_PNG_03_Ito_Daisuke.docx",
        "photo": "PNG_03_Ito_Daisuke_foto.png",
        "name_jp": "伊藤 大介",
        "name": "ITO Daisuke",
        "age": "52",
        "role": "Responsabile Scientifica (Kanshiki-ka)",
        "rank": "Kanshiki-kan Classe 1 (鑑識官一級)",
        "years": "15+",
        "ki": "8",
        "attrs": [
            ("Distacco", "8", "Imperturbabile — centinaia di scene del crimine"),
            ("Pazienza", "7", "Metodico, lento, preciso. La fretta è nemica della scienza"),
            ("Silenzio", "7", "Riservato per natura — parla solo quando è rilevante"),
            ("Lucidità", "8", "Il suo strumento principale — vede nei dettagli ciò che altri ignorano"),
            ("Ascolto", "4", "Non gli interessano le persone — gli interessano le prove"),
            ("Presenza", "5", "Non carismatico, ma nessuno mette in dubbio la sua competenza"),
        ],
        "gou": {
            "name": "Occhio della Gru",
            "kanji": "鶴の目",
            "desc": "Vedi ciò che altri non vedono — il dettaglio che cambia tutto.",
            "attr": "Lucidità (8)", "cost": "2 Ki",
            "success": "Noti il dettaglio nascosto e ne capisci immediatamente il significato",
            "failure": "Vedi qualcosa di anomalo, ma non riesci a capire cosa significa",
        },
        "background": "Sei nella Kanshiki-ka da oltre 15 anni. Prima eri chimico industriale, poi sei entrato nella polizia per curiosità e ci sei rimasto per passione.\n\nHai formato decine di tecnici, migliorato le procedure della prefettura, testimoniato in centinaia di processi. I tuoi rapporti non sono mai stati contestati con successo.\n\nLe prove non mentono, le persone sì. Per questo preferisci le prove.",
        "problem_title": "La Figlia",
        "problem_text": "Tua figlia Ito Haruka ha 24 anni e non ti parla da due. L'ultima volta ti ha detto: \"Non sei mai stato un padre. Sei stato un laboratorio con le gambe.\"\n\nAveva ragione. Quando era bambina, tu eri al lavoro. Tua moglie Sachiko ha tenuto insieme la famiglia finché ha potuto, poi ha smesso. Il divorzio è stato pulito, silenzioso.\n\nHaruka vive a Osaka. Ogni tanto Sachiko ti aggiorna — \"Sta bene, Daisuke\" — e tu annuisci e non chiedi altro. Ma la verità è che vorresti chiamarla. Non sai cosa dire.",
        "png_name": "ITO Sachiko",
        "png_role": "Ex moglie, divorzio 6 anni fa",
        "png_age": "50",
        "png_job": "Insegnante di scuola media a Kyoto",
        "png_wants": "Che Daisuke faccia qualcosa per Haruka. Ha tenuto insieme la famiglia per vent'anni. Ha smesso. Ma non ha smesso di sperare.",
        "png_behavior": "Diretta, pratica, senza rancore ma senza pazienza. Chiama solo quando c'è qualcosa di concreto su Haruka.",
        "enja_name": "Prof. HONDA Akira",
        "enja_role": "Chimico, Università di Kyoto",
        "enja_desc": "Professore di Chimica Analitica. Ex compagno di studi, consulente informale. Analisi avanzate, consulenze tossicologiche, attrezzature universitarie.",
        "enja_contact": "Telefono dell'ufficio universitario. Solo in orario di lavoro.",
        "enja_cost": "Pubblicare i risultati (anonimizzati) nei suoi paper accademici.",
        "tatemae": "Il professionista. Freddo, preciso, impersonale. I dati parlano, le opinioni no.",
        "honne": "Sei solo. L'appartamento è ordinato come un laboratorio. A volte prendi il telefono per chiamare Haruka. Non componi mai il numero.",
        "phrase": "\"I dati sono questi.\" (perché i dati non feriscono, non deludono, non se ne vanno)",
        "pressure": "Ti chiudi nel metodo. Controlli tre volte. Rallenti. Hai sbagliato una volta, vent'anni fa, e non te lo sei mai perdonato.",
        "weakness": "Casi che coinvolgono giovani donne dell'età di Haruka. Casi in cui la vittima è stata ignorata o abbandonata.",
        "district": [
            ("I PG", "Rapporto professionale. Li tratta come colleghi, non come superiori."),
            ("Commissario Taniguchi", "Rispetto reciproco. Taniguchi non lo disturba senza motivo."),
            ("Yamada Tetsuo", "Lo conosce dalla centrale. Lo saluta, niente di più."),
        ],
        "traits": {
            "vizio": "Tè verde fortissimo — quasi amaro. Lo prepara lui stesso col thermos.",
            "tic": "Si toglie e rimette gli occhiali quando pensa. Quando dice \"Interessante\", gli occhiali sono in mano.",
            "oggetto": "Lente d'ingrandimento tascabile Nikon, vecchia di trent'anni.",
            "gusto": "Bento preparato da solo — sempre uguale: riso, salmone grigliato, tsukemono.",
            "extra": ("Rituale", "Prima di lavorare su una scena, fa un giro completo della stanza senza toccare nulla. Solo guardando."),
        },
    },
]


def create_sheet(png):
    doc = Document()

    # Page setup A4 with tight margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.3)
    section.bottom_margin = Cm(0.3)
    section.left_margin = Cm(0.3)
    section.right_margin = Cm(0.3)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ════════════════════════════════════════════
    # PAGE 1
    # ════════════════════════════════════════════

    # ── Top decorative line ──
    p = add_para(doc, "限界  ·  GENKAI v1.2  ·  限界", size=7, color=GOLD,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    # PNG label
    add_para(doc, "▸ SCHEDA PNG RICORRENTE ◂", size=7, color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # ── Header: name block + photo ──
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_tbl)

    left_cell = header_tbl.rows[0].cells[0]
    left_cell.width = Cm(14)
    tc_left = left_cell._tc
    tcPr_left = tc_left.get_or_add_tcPr()
    tcPr_left.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))

    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, png["name_jp"], size=28, color=DARK, font="Calibri")

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, png["name"], bold=True, size=14, color=DARK)

    right_cell = header_tbl.rows[0].cells[1]
    right_cell.width = Cm(3.2)
    tc_right = right_cell._tc
    tcPr_right = tc_right.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="0" w:type="dxa"/>'
        f'  <w:left w:w="0" w:type="dxa"/>'
        f'  <w:bottom w:w="0" w:type="dxa"/>'
        f'  <w:right w:w="0" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr_right.append(tcMar)
    tcPr_right.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
    set_cell_borders(right_cell,
                     top=(4, GOLD_HEX), bottom=(4, GOLD_HEX),
                     left=(4, GOLD_HEX), right=(4, GOLD_HEX))

    p_photo = right_cell.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    photo_path = OUT_DIR / png.get("photo", "")
    if photo_path.exists():
        run = p_photo.add_run()
        run.add_picture(str(photo_path), width=Cm(3.2))
    else:
        r_photo = p_photo.add_run("FOTO")
        r_photo.font.size = Pt(12)
        r_photo.font.color.rgb = RGBColor(200, 200, 200)
        r_photo.font.name = "Calibri"

    add_separator(doc)

    # ── Info line ──
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(4)
    items = [
        ("Età ", png["age"]),
        ("Ruolo ", png["role"]),
        ("Grado ", png["rank"]),
        ("Servizio ", f"{png['years']} anni"),
    ]
    for i, (label, value) in enumerate(items):
        if i > 0:
            add_run(p3, "  ·  ", size=8, color=GOLD)
        add_run(p3, label, bold=True, size=8, color=MUTED)
        add_run(p3, value, size=8, color=DARK)

    add_thin_separator(doc)

    # ── ATTRIBUTI ──
    add_section_header(doc, "Attributi", "能力")

    # PNG: 4 columns (no "Finale" or "Variazioni" — all points distributed)
    tbl = doc.add_table(rows=7, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_el = tbl._tbl
    tblPr_t = tbl_el.tblPr if tbl_el.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr_t.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{GOLD_HEX}"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="D0C8B0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))
    tblPr_t.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="30" w:type="dxa"/>'
        f'  <w:bottom w:w="30" w:type="dxa"/>'
        f'  <w:left w:w="80" w:type="dxa"/>'
        f'  <w:right w:w="80" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))

    headers = [
        ("Attributo", WD_ALIGN_PARAGRAPH.LEFT),
        ("Valore", WD_ALIGN_PARAGRAPH.CENTER),
        ("Note", WD_ALIGN_PARAGRAPH.LEFT),
    ]
    for j, (header, align) in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_shading(cell, DARK_HEX)
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    for j in range(3):
        set_cell_borders(tbl.rows[0].cells[j], bottom=(6, GOLD_HEX))

    for i, (attr_name, attr_val, attr_note) in enumerate(png["attrs"]):
        row = tbl.rows[i + 1]
        is_strong = int(attr_val) >= 7

        cell0 = row.cells[0]
        p = cell0.paragraphs[0]
        r = p.add_run(attr_name)
        r.bold = is_strong
        r.font.size = Pt(9)
        r.font.name = "Calibri"
        r.font.color.rgb = DARK

        cell1 = row.cells[1]
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(attr_val)
        r.bold = is_strong
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        r.font.color.rgb = RED if is_strong else DARK

        cell2 = row.cells[2]
        p = cell2.paragraphs[0]
        r = p.add_run(attr_note)
        r.italic = True
        r.font.size = Pt(8)
        r.font.name = "Calibri"
        r.font.color.rgb = MUTED

        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_HEX)

    for row in tbl.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(13)

    # Distribution note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Distribuzione: ", bold=True, size=8, color=MUTED)
    add_run(p, "base 4 × 6 = 24  +  12 punti  +  3 bonus (d6)  =  39 totale", size=8, color=MUTED)

    add_thin_separator(doc)

    # ── KI (fixed) ──
    add_section_header(doc, "Ki — Tenuta", "気")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, f"Ki Max:  {png['ki']}", bold=True, size=11, color=DARK)
    add_run(p, "          ", size=11, color=DARK)
    add_run(p, "≤ 3 Genkai", size=8, color=MUTED, italic=True)
    add_run(p, "  ·  ", size=8, color=GOLD)
    add_run(p, "= 1 Critico", size=8, color=MUTED, italic=True)
    add_run(p, "  ·  ", size=8, color=GOLD)
    add_run(p, "< 1 Fuori gioco", size=8, color=MUTED, italic=True)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(4)
    p3.paragraph_format.left_indent = Cm(0.3)
    add_run(p3, "Ki attuale:  ", bold=True, size=11, color=DARK)
    add_run(p3, "_______________________________________________________________", size=11, color=RGBColor(200, 200, 200))

    add_thin_separator(doc)

    # ── GOU (fixed, single) ──
    add_section_header(doc, "Gou — Il Debito", "業")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "Il Gou funziona SEMPRE. Successo = dettagli precisi. Fallimento = dettagli vaghi.",
            italic=True, size=8, color=MUTED)

    gou = png["gou"]
    add_gou_box(doc, f"{gou['name']}  {gou['kanji']}", gou["desc"],
                gou["attr"], gou["cost"], gou["success"], gou["failure"])

    # ── Satori reminder ──
    add_separator(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "悟り  ", size=10, color=GOLD)
    add_run(p, "SATORI — Successo automatico, 1 volta per sessione:  ", size=8, color=MUTED)
    add_run(p, "□ Usato", size=9, color=DARK)

    # ════════════════════════════════════════════
    # PAGE 2
    # ════════════════════════════════════════════

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_run(p, png["name_jp"], size=14, color=GOLD)
    add_run(p, "  ·  ", size=10, color=MUTED)
    add_run(p, png["name"], bold=True, size=11, color=DARK)

    add_separator(doc)

    # ── CHI SEI ──
    add_section_header(doc, "Chi Sei", "人")

    for para_text in png["background"].split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.3)
        add_run(p, para_text, size=9, color=DARK)

    add_thin_separator(doc)

    # ── IL TUO PROBLEMA ──
    add_section_header(doc, f"Il Tuo Problema — {png['problem_title']}", "影")

    for para_text in png["problem_text"].split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.3)
        add_run(p, para_text, size=9, color=DARK)

    add_thin_separator(doc)

    # ── PNG DEL PROBLEMA ──
    add_section_header(doc, png["png_name"], "縁")

    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(tbl2)

    png_info = [
        ("Età", png["png_age"]),
        ("Occupazione", png["png_job"]),
        ("Relazione", png["png_role"]),
        ("Cosa vuole", png["png_wants"]),
    ]
    for i, (label, value) in enumerate(png_info):
        row = tbl2.rows[i]
        c0 = row.cells[0]
        p = c0.paragraphs[0]
        add_run(p, label, bold=True, size=8, color=MUTED)
        c0.width = Cm(2.5)

        c1 = row.cells[1]
        p = c1.paragraphs[0]
        add_run(p, value, size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, png["png_behavior"], italic=True, size=9, color=DARK)

    add_thin_separator(doc)

    # ── CONOSCENZA ──
    add_section_header(doc, f"Conoscenza — {png['enja_name']}", "縁者")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, png["enja_role"], italic=True, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, png["enja_desc"], size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Contatto: ", bold=True, size=8, color=MUTED)
    add_run(p, png["enja_contact"], size=8, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Costo: ", bold=True, size=8, color=MUTED)
    add_run(p, png["enja_cost"], size=8, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "1 volta per sessione senza conseguenze. La seconda volta → chiede qualcosa in cambio.",
            italic=True, size=7, color=MUTED)

    add_thin_separator(doc)

    # ── COME TI COMPORTI ──
    add_section_header(doc, "Come Ti Comporti", "面")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "建前  In pubblico:  ", bold=True, size=9, color=GOLD)
    add_run(p, png["tatemae"], size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "本音  In privato:  ", bold=True, size=9, color=GOLD)
    add_run(p, png["honne"], size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Frase tipica:  ", bold=True, size=9, color=MUTED)
    add_run(p, png["phrase"], italic=True, size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Sotto pressione:  ", bold=True, size=9, color=RED)
    add_run(p, png["pressure"], size=9, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Debolezza:  ", bold=True, size=9, color=RED)
    add_run(p, png["weakness"], size=9, color=DARK)

    # ── TRATTI PERSONALI ──
    add_thin_separator(doc)
    add_section_header(doc, "Tratti Personali", "癖")

    trait_labels = [
        ("Vizio", "vizio"),
        ("Tic", "tic"),
        ("Oggetto personale", "oggetto"),
        ("Gusto", "gusto"),
    ]
    for label, key in trait_labels:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{label}:  ", bold=True, size=9, color=MUTED)
        add_run(p, png["traits"][key], size=9, color=DARK)

    extra_label, extra_text = png["traits"]["extra"]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{extra_label}:  ", bold=True, size=9, color=MUTED)
    add_run(p, extra_text, size=9, color=DARK)

    # ── RAPPORTI NEL DISTRETTO ──
    add_thin_separator(doc)
    add_section_header(doc, "Rapporti nel Distretto", "絆")

    for name, quote in png["district"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{name}:  ", bold=True, size=9, color=MUTED)
        add_run(p, f"\u201c{quote}\u201d", italic=True, size=9, color=DARK)

    # ── Footer ──
    add_separator(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "GENKAI 限界 v1.2  ·  PNG RICORRENTE", size=7, color=MUTED)

    # Save
    out_path = OUT_DIR / png["filename"]
    doc.save(str(out_path))
    print(f"OK — {out_path.name}")


def main():
    for png in PNG_DATA:
        create_sheet(png)

if __name__ == "__main__":
    main()
