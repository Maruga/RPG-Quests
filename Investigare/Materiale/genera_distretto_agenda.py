# -*- coding: utf-8 -*-
"""
Genera Distretto_Agenda.docx — versione Word modificabile della scheda distretto.
Output: Distretto_Agenda.docx accanto a questo script.

Sezioni:
  1. Header burocratico
  2. Attributi del Distretto (tabella 6 colonne)
  3. Azioni delegabili (tabella riferimento)
  4. Personale fisso del distretto (4 PNG senza statistiche di gioco)
  5. Agenda — calendario operativo (vuoto, compilabile)
  6. Note libere
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# =====================================================
# DATI
# =====================================================

ATTRIBUTI = [
    ("Organico", "8", "Squadra di sorveglianza (+1)"),
    ("Efficienza", "7", "—"),
    ("Velocità", "6", "—"),
    ("Risorse", "8", "Archivio esteso (+1)"),
    ("Rete", "7", "—"),
    ("Corruzione", "5", "28% attivazione"),
]

AZIONI = [
    ("Pedinamento / appostamento / sorveglianza", "Organico (8)", "+1 sorv.", "Operazione continua, autorizzazione interna"),
    ("Inseguimento veicolare / supporto arresto", "Organico (8)", "—", "Operazione singola, urgenza"),
    ("Perquisizione di un luogo", "Organico (8)", "—", "Richiede mandato → Procura. 4–6h + autorizzazione"),
    ("Ricerca anagrafica / precedenti / patrimoniali", "Risorse (8)", "+1 arch.", "1–4 ore"),
    ("Acquisizione tabulati telefonici", "Risorse (8)", "+1 arch.", "24–48h + autorizzazione"),
    ("Acquisizione filmati TVCC / scontrini", "Risorse (8)", "+1 arch.", "12–24h"),
    ("Elaborazione digitale filmato/audio", "Efficienza (7)", "—", "24–48h, Sez. Analisi Video"),
    ("Analisi tossicologica / impronte / balistica", "Efficienza (7)", "—", "24–72h, laboratori esterni"),
    ("Analisi documentale / grafologia", "Efficienza (7)", "—", "48h"),
    ("Localizzazione di una persona", "Rete (7)", "—", "4–8h"),
    ("Raccolta informazioni su un soggetto", "Rete (7)", "—", "Variabile"),
    ("Verifica documento / identità", "Efficienza (7)", "—", "1–2h"),
]

PNG = [
    {
        "nome": "TANIGUCHI Osamu",
        "kanji": "谷口 治",
        "ruolo": "Commissario / Capo Sezione Omicidi",
        "grado": "Keibu (警部) — 56 anni — 18+ anni di servizio",
        "desc": "Superiore diretto dei PG. Pragmatico, esigente, protettivo. Riceve briefing e rapporti, autorizza risorse e mandati attraverso il PM, gestisce le pressioni mediatiche e politiche. Frase tipica: «Sedetevi. Raccontate.»",
        "contatto": "Ufficio Piano 2 — Cercapersone, risponde <10 min",
    },
    {
        "nome": "YAMADA Tetsuo",
        "kanji": "山田 哲夫",
        "ruolo": "Sergente — Accompagnatore PG",
        "grado": "Junsa-buchō (巡査部長) — 35 anni — 10+ anni servizio",
        "desc": "Agente affidabile assegnato 24/7 ai PG come supporto operativo. Guida, prende appunti, fa da riferimento territoriale. Non investiga ma è prezioso per la logistica e i sopralluoghi sul campo.",
        "contatto": "Sempre disponibile durante l'indagine — Cercapersone <20 min",
    },
    {
        "nome": "ITO Daisuke",
        "kanji": "伊藤 大介",
        "ruolo": "Responsabile Polizia Scientifica",
        "grado": "Kanshiki-kan Cl.1 (鑑識官一級) — 52 anni — 15+ anni",
        "desc": "Capo della squadra scientifica (Kanshiki-ka). Metodico, fattuale, mai speculativo. Riferisce dati, mai opinioni. Punto di contatto coi laboratori esterni per medicina legale, tossicologia, balistica.",
        "contatto": "Lab seminterrato — Su scena del crimine in 45 min",
    },
    {
        "nome": "WATANABE Hideo",
        "kanji": "渡辺 秀夫",
        "ruolo": "Procuratore — Pubblico Ministero",
        "grado": "Kenji (検事) — 51 anni — magistrato dal 1971",
        "desc": "Magistrato del PM di riferimento. Imparziale, scrupoloso, lento per principio. Firma o respinge mandati e autorizzazioni. Non si contatta direttamente — sempre tramite Taniguchi. Frase tipica: «Lo valuterò.»",
        "contatto": "Procura Tribunale Prefetturale, 4° piano — solo orari ufficio",
    },
]

# =====================================================
# UTILITIES
# =====================================================

def set_cell_shading(cell, color_hex):
    """Imposta sfondo cella (color_hex senza #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="555555", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)

def add_horizontal_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '333333')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Courier New'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_horizontal_rule(p)
    return p

# =====================================================
# DOCUMENTO
# =====================================================

def build_document(output_path):
    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Stile base Courier New
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(10)

    # ===== HEADER =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("捜査一課 — SEZIONE INVESTIGATIVA CENTRALE")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SCHEDA DISTRETTO + AGENDA OPERATIVA")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Polizia Prefetturale — quadro operativo e cronologia indagini")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    add_horizontal_rule(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("RISERVATO")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x8b, 0x00, 0x00)

    # ===== ATTRIBUTI =====
    add_section_title(doc, "Attributi del Distretto")
    tbl = doc.add_table(rows=3, cols=6)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col_idx, (label, value, spec) in enumerate(ATTRIBUTI):
        # riga 0: label
        c = tbl.cell(0, col_idx)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c, "EEEEEE")
        set_cell_borders(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # riga 1: value
        c = tbl.cell(1, col_idx)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(20)
        # riga 2: spec
        c = tbl.cell(2, col_idx)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(spec)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ===== AZIONI DELEGABILI =====
    add_section_title(doc, "Azioni delegabili — riferimento operativo")
    tbl = doc.add_table(rows=1 + len(AZIONI), cols=4)
    tbl.autofit = False
    widths = [Cm(7.5), Cm(3.0), Cm(2.2), Cm(5.0)]
    for col_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[col_idx].width = w
    # header
    headers = ["Azione", "Attributo", "Bonus spec.", "Note / tempi"]
    for col_idx, h in enumerate(headers):
        c = tbl.rows[0].cells[col_idx]
        set_cell_shading(c, "DDDDDD")
        set_cell_borders(c)
        p = c.paragraphs[0]
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(8)
        if col_idx in (1, 2):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # body
    for row_idx, (azione, attr, bonus, note) in enumerate(AZIONI, start=1):
        cells = [azione, attr, bonus, note]
        for col_idx, txt in enumerate(cells):
            c = tbl.rows[row_idx].cells[col_idx]
            set_cell_borders(c)
            if row_idx % 2 == 0:
                set_cell_shading(c, "F8F8F4")
            p = c.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(9)
            if col_idx in (1, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== PERSONALE FISSO =====
    add_section_title(doc, "Personale fisso del distretto")
    # tabella 2x2 di card
    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False
    for row in tbl.rows:
        row.cells[0].width = Cm(8.5)
        row.cells[1].width = Cm(8.5)
    for idx, png in enumerate(PNG):
        r_idx = idx // 2
        c_idx = idx % 2
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_borders(cell)
        set_cell_shading(cell, "F9F6EC")
        # Pulisci paragrafo iniziale
        cell.text = ""
        # Nome + kanji
        p = cell.paragraphs[0]
        r = p.add_run(png["nome"])
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run("  " + png["kanji"])
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # Ruolo
        p = cell.add_paragraph()
        r = p.add_run(png["ruolo"].upper())
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        r.bold = True
        # Grado
        p = cell.add_paragraph()
        r = p.add_run(png["grado"])
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # Descrizione
        p = cell.add_paragraph()
        r = p.add_run(png["desc"])
        r.font.size = Pt(9)
        # Separatore
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        add_horizontal_rule(p)
        # Contatto
        p = cell.add_paragraph()
        r = p.add_run(png["contatto"])
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ===== AGENDA — CALENDARIO =====
    add_section_title(doc, "Agenda — calendario operativo")
    p = doc.add_paragraph()
    r = p.add_run("Compilare giorno per giorno: fatti, interrogatori, decisioni, consegne, eventi.")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Mese / Anno modificabili
    p = doc.add_paragraph()
    r = p.add_run("Mese / Anno: ____________________________     Caso: ____________________________")
    r.font.size = Pt(10)
    r.bold = True

    # Tabella agenda 31 righe vuote (giorno + giorno settimana + descrizione)
    tbl = doc.add_table(rows=32, cols=3)
    tbl.autofit = False
    widths = [Cm(1.2), Cm(2.2), Cm(14.1)]
    for col_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[col_idx].width = w
    # header
    headers = ["Giorno", "Sett.", "Eventi / annotazioni"]
    for col_idx, h in enumerate(headers):
        c = tbl.rows[0].cells[col_idx]
        set_cell_shading(c, "DDDDDD")
        set_cell_borders(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(8)
    # 31 righe vuote, colonna 1 pre-popolata con il numero
    for d in range(1, 32):
        row = tbl.rows[d]
        # giorno num
        c = row.cells[0]
        set_cell_borders(c)
        set_cell_shading(c, "F0F0F0")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(d))
        r.bold = True
        r.font.size = Pt(11)
        # sett (vuoto)
        c = row.cells[1]
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # eventi (vuoto)
        c = row.cells[2]
        set_cell_borders(c)
        # forza altezza minima della riga
        tr_pr = row._tr.get_or_add_trPr()
        trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), '480')  # ~ 24pt
        trh.set(qn('w:hRule'), 'atLeast')
        tr_pr.append(trh)

    # ===== NOTE LIBERE =====
    add_section_title(doc, "Note libere")
    tbl = doc.add_table(rows=1, cols=1)
    c = tbl.rows[0].cells[0]
    set_cell_borders(c)
    set_cell_shading(c, "FAFAF4")
    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    trh = OxmlElement('w:trHeight')
    trh.set(qn('w:val'), '4000')  # area ampia
    trh.set(qn('w:hRule'), 'atLeast')
    tr_pr.append(trh)
    p = c.paragraphs[0]
    r = p.add_run("Spazio libero per teorie, ipotesi, collegamenti, schemi, link tra sospettati, idee in corso…")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    # ===== FOOTER =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    add_horizontal_rule(p)
    r = p.add_run("Sezione Investigativa Centrale — Scheda Distretto + Agenda Operativa — uso GM/PG")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    r.italic = True

    doc.save(output_path)
    print(f"OK -> {output_path}")


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Distretto_Agenda.docx")
    build_document(out)
