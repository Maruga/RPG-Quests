# -*- coding: utf-8 -*-
"""RICHIAMO — Il libro del GM: aggrega tutti i file canone in un unico .docx
con indice, per condurre l'avventura dal tablet.
Converte il markdown dei file (heading, tabelle, liste, citazioni, bold/italic)
e accoda KAGE_FAMIGLIE.docx. Esclude i prompt di generazione immagini."""

import re, os, csv, io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Public\_Clienti\Maruga\Giochi\Vampiri\Vault\Vampiri\Investigare\Avventura Richiamo"
OUT = os.path.join(BASE, "RICHIAMO_Libro_GM.docx")

INK = RGBColor(0x1E, 0x1E, 0x1E)
DARKRED = RGBColor(0x8B, 0x1A, 0x1A)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

def set_eastasia(style_or_run, name='Yu Mincho'):
    rpr = style_or_run.font.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), name)

# stile base
st = doc.styles['Normal']
st.font.name = 'Georgia'; st.font.size = Pt(10); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(3); st.paragraph_format.line_spacing = 1.06
set_eastasia(st)

# heading nativi ristilizzati (servono al TOC)
for name, size, color in [('Heading 1', 17, INK), ('Heading 2', 13.5, DARKRED), ('Heading 3', 11.5, INK)]:
    hs = doc.styles[name]
    hs.font.name = 'Georgia'; hs.font.size = Pt(size); hs.font.bold = True
    hs.font.color.rgb = color
    hs.paragraph_format.space_before = Pt(10 if name != 'Heading 1' else 4)
    hs.paragraph_format.space_after = Pt(3)
    set_eastasia(hs)

# footer con numero pagina
def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = '1'
    r.append(t); fld.append(r)
    p._p.append(fld)
    for run in p.runs: run.font.size = Pt(8)
add_page_number(doc.sections[0])

# ---------- inline markdown ----------
TOKEN = re.compile(r'(\*\*.+?\*\*|(?<!\*)\*[^*\n]+?\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\)|~~.+?~~)')
LINK = re.compile(r'\[([^\]]+)\]\([^)]+\)')

def render_inline(p, text, size=None, force_italic=False, color=None):
    parts = TOKEN.split(text)
    for part in parts:
        if not part: continue
        bold = italic = strike = code = False
        chunk = part
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            bold = True; chunk = part[2:-2]
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            italic = True; chunk = part[1:-1]
        elif part.startswith('`') and part.endswith('`'):
            code = True; chunk = part[1:-1]
        elif part.startswith('~~') and part.endswith('~~'):
            strike = True; chunk = part[2:-2]
        m = LINK.match(chunk)  # link dentro il chunk: tengo solo il testo
        chunk = LINK.sub(r'\1', chunk)
        # bold annidato dentro corsivo e viceversa: passata grezza sufficiente
        chunk = chunk.replace('**', '')
        r = p.add_run(chunk)
        r.bold = bold or None
        r.italic = (italic or force_italic) or None
        if strike: r.font.strike = True
        if code:
            r.font.name = 'Consolas'; r.font.size = Pt((size or 10) - 0.5)
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color

def quote_border(p):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), '999999')
    pbdr.append(left); ppr.append(pbdr)
    p.paragraph_format.left_indent = Cm(0.35)

# ---------- blocco markdown ----------
TABLE_ROW = re.compile(r'^\s*\|.*\|\s*$')
TABLE_SEP = re.compile(r'^\s*\|[\s:\-|]+\|\s*$')
BULLET = re.compile(r'^\s*[-·]\s+')
NUMBERED = re.compile(r'^\s*\d+\.\s+')

def emit_table(rows):
    parsed = []
    for raw in rows:
        if TABLE_SEP.match(raw): continue
        cells = [c.strip() for c in raw.strip().strip('|').split('|')]
        parsed.append(cells)
    if not parsed: return
    ncols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.autofit = True
    for i, cells in enumerate(parsed):
        for j in range(ncols):
            cell = tbl.cell(i, j)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            text = cells[j] if j < len(cells) else ''
            render_inline(cell.paragraphs[0], text, size=8.5)
            if i == 0:
                for r in cell.paragraphs[0].runs: r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def emit_markdown(lines, skip_h1=True):
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        if TABLE_ROW.match(line):
            rows = []
            while i < len(lines) and TABLE_ROW.match(lines[i].rstrip('\n')):
                rows.append(lines[i].rstrip('\n')); i += 1
            emit_table(rows)
            continue
        stripped = line.strip()
        if not stripped:
            i += 1; continue
        if stripped in ('---', '***', '___'):
            i += 1; continue
        if stripped.startswith('# '):
            if not skip_h1:
                doc.add_heading(re.sub(r'[*`]', '', stripped[2:]).strip(), level=2)
            i += 1; continue
        if stripped.startswith('## '):
            h = doc.add_heading('', level=2)
            render_inline(h, stripped[3:], size=13.5)
            for r in h.runs: r.font.color.rgb = DARKRED; r.bold = True
            i += 1; continue
        if stripped.startswith('### '):
            h = doc.add_heading('', level=3)
            render_inline(h, stripped[4:], size=11.5)
            for r in h.runs: r.bold = True
            i += 1; continue
        if stripped.startswith('> '):
            p = doc.add_paragraph(); quote_border(p)
            render_inline(p, stripped[2:], size=9.5)
            i += 1; continue
        if stripped == '>':
            i += 1; continue
        if BULLET.match(line):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            render_inline(p, BULLET.sub('', line, count=1), size=10)
            i += 1; continue
        if NUMBERED.match(line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            render_inline(p, NUMBERED.sub('', line, count=1), size=10)
            i += 1; continue
        p = doc.add_paragraph()
        render_inline(p, stripped)
        i += 1

def load_md(fname, cut_ranges=None):
    """cut_ranges: lista di (start_marker, end_marker|None): esclude da start
    (inclusa) fino a end (esclusa); end None = fino a fine file."""
    with io.open(os.path.join(BASE, fname), encoding='utf-8') as f:
        lines = f.readlines()
    if not cut_ranges: return lines
    out, skipping, end_marker = [], False, None
    for ln in lines:
        s = ln.strip()
        if skipping:
            if end_marker and s.startswith(end_marker):
                skipping = False; end_marker = None
                out.append(ln)
            continue
        started = False
        for start, end in cut_ranges:
            if s.startswith(start):
                skipping = True; end_marker = end; started = True; break
        if not started:
            out.append(ln)
    return out

def chapter(title, subtitle=None):
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for r in h.runs: r.font.color.rgb = INK
    if subtitle:
        p = doc.add_paragraph()
        render_inline(p, subtitle, size=9.5, force_italic=True, color=GREY)

# ============================ FRONTESPIZIO ============================
cover = os.path.join(BASE, 'Immagini', 'Copertina.png')
if os.path.exists(cover):
    doc.add_picture(cover, width=Cm(17.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RICHIAMO'); r.font.size = Pt(34); r.bold = True; r.font.name = 'Georgia'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Il libro del GM — Giappone, 11–14 novembre 1986'); r.font.size = Pt(13); r.italic = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("un'avventura fuori dal normale per GENKAI 限界"); r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SOLO PER IL GM — i giocatori non devono leggere questo documento'); r.font.size = Pt(10); r.bold = True; r.font.color.rgb = DARKRED

# ============================ INDICE ============================
doc.add_page_break()
doc.add_heading('Indice', level=1)
toc_static = [
    ('1.', 'Guida rapida — come si usa questo libro'),
    ('2.', 'La copertina e il gōgai (da leggere ai giocatori)'),
    ('3.', 'La sequenza di gioco — i quattro giorni'),
    ('4.', 'La linea maestra e le tre vie'),
    ('5.', 'La storia — la verità del GM'),
    ('6.', 'Il Nushi — la scheda del mostro'),
    ('7.', 'Gli indizi'),
    ('8.', 'Gli scontri'),
    ('9.', "L'arma segreta — la pazienza"),
    ('10.', 'Le basi militari'),
    ('11.', 'I bollettini ai giocatori'),
    ('12.', 'Le famiglie a casa'),
    ('13.', 'La coda — il secondo dormiente (da leggere alla fine)'),
    ('A.', 'Appendice — il nastro dei marcatori (dov\u2019è adesso?)'),
    ('B.', 'Appendice — materiali fuori da questo documento'),
]
for num, t in toc_static:
    p = doc.add_paragraph()
    r = p.add_run(num + '  '); r.bold = True
    p.add_run(t)
    p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('Indice con numeri di pagina (aprire una volta in Word su PC e aggiornare i campi):')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
# campo TOC
tp = doc.add_paragraph()
fld_b = OxmlElement('w:r'); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); fc.set(qn('w:dirty'), 'true'); fld_b.append(fc)
instr_r = OxmlElement('w:r'); instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' TOC \\o "1-2" \\h \\z \\u '; instr_r.append(instr)
sep_r = OxmlElement('w:r'); fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate'); sep_r.append(fs)
txt_r = OxmlElement('w:r'); tt = OxmlElement('w:t'); tt.text = '(campo indice — aggiornare in Word)'; txt_r.append(tt)
end_r = OxmlElement('w:r'); fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end'); end_r.append(fe)
for el in (fld_b, instr_r, sep_r, txt_r, end_r): tp._p.append(el)

# ============================ GUIDA RAPIDA ============================
chapter('1 · Guida rapida', 'come si usa questo libro al tavolo')
g = [
    "**I nomi**: il mostro = **Nushi** 主 «il Padrone» (*uminari* = la parola dei pescatori per il suo rombo) · l'organizzazione = **Ryūgū-kai** 竜宮会 «Società del Palazzo del Drago» — mai identificata.",
    "**Le date**: 1d = **martedì 11 novembre 1986** (emersione 06:05, Genkai distrutta 07:00) · 2d = mercoledì 12 (consegna dell'oro, 10:00) · 3d = giovedì 13 (**il giorno nero** di Wakasa) · 4d = venerdì 14 (Kashiwazaki-Kariwa, 14:03) · sabato 15, 17:25: la coda (cap. 13).",
    "**Il tempo si tiene così**: la bestia fa **12 km/h, sempre** — un marcatore ogni 2 ore (appendice A); bollettini ai giocatori alle 09:00 e alle 21:00 (cap. 11). Il GM traccia un numero solo: «i km del Nushi».",
    "**Le due vie di vittoria**: demolire TUTTI i camioncini-faro (torna all'acqua) · oppure il disturbatore al capannone di Maizuru (torna al mare dal punto più vicino). Ogni faro demolito prima dello sbarco = una centrale salvata.",
    "**Cosa non c'è qui dentro**: gli handout grafici (già stampati: i 9 bollettini, il programma cantieri, la mappa del capannone) e le immagini — elenco in appendice B. I prompt di generazione restano nei file .md.",
    "**Sistema**: GENKAI — Lo Scontro/Shōtotsu v2.1 + tiri standard (2d6 ≤ attributo) e pressioni che costano Ki. **I Gou si usano normalmente** (ogni pregen ne ha uno). Kage meccanico e soroban restano fuori; le telefonate del cap. 12 sono scene libere.",
    "**I PG**: le 5 schede esistenti — selezionati dal loro procuratore, una delle squadre migliori; in elicottero alla Sala 1 di Karatsu alle ~11:00 del giorno 1.",
]
for t in g:
    p = doc.add_paragraph(style='List Bullet'); render_inline(p, t)

# ============================ CAPITOLI DA MD ============================
chapter('2 · La copertina e il gōgai', 'da mostrare/leggere ai giocatori in apertura (pre-sessione via WhatsApp)')
emit_markdown(load_md('COPERTINA.md'))

chapter('3 · La sequenza di gioco', 'il riepilogo dei quattro giorni — la spina dorsale della conduzione')
emit_markdown(load_md('Avventura.md'))

chapter('4 · La linea maestra e le tre vie', 'il quadro giorno per giorno: nastro, fari, storia — e le leve dei PG')
schema = load_md('SCHEMA.md')
start = next(i for i, l in enumerate(schema) if l.strip().startswith('## 3. La linea maestra'))
end = next(i for i, l in enumerate(schema) if l.strip().startswith('## 5. Proposte APERTE'))
emit_markdown(schema[start:end])

chapter('5 · La storia — la verità del GM', 'Ryūgū-kai, il piano, il riscatto, lo strumento, la regola del buio, lo Stato')
emit_markdown(load_md('STORIA.md'))

chapter('6 · Il Nushi — la scheda del mostro', 'dimensioni, fisiologia, punti deboli, fasce, scia, percorso')
emit_markdown(load_md('SCHEDA_NUSHI.md', cut_ranges=[('## 11. Prompt per le immagini', None)]))

chapter('7 · Gli indizi', 'camioncini · oro · strumento · organizzazione')
emit_markdown(load_md('INDIZI.md', cut_ranges=[('### Prompt per le foto', '## Indizio 2')]))

chapter('8 · Gli scontri', 'le tre squadre su Shōtotsu v2.1 — camioncino, oro, capannone')
emit_markdown(load_md('SCONTRI.md', cut_ranges=[('### Prompt foto del capannone', '*Coda comune')]))

chapter("9 · L'arma segreta — la pazienza", 'MAI mostrare al tavolo: scala 0-12, lo scarico, l\u2019impulso, la trappola di Wakasa')
emit_markdown(load_md('arma-kaiju.md'))

chapter('10 · Le basi militari', 'che cosa sono le armi, che cosa danno le basi, la nave-esca')
emit_markdown(load_md('BASI_MILITARI.md'))

chapter('11 · I bollettini ai giocatori', 'il canone degli eventi (timeline NON contrastata) — con costanti e note GM')
emit_markdown(load_md('bollettini.md'))

# ============================ KAGE FAMIGLIE (append docx) ============================
chapter('12 · Le famiglie a casa', 'le scene-telefonata di Yamamoto, Nakamura e Sato — dal dossier KAGE_FAMIGLIE')
src = Document(os.path.join(BASE, 'KAGE_FAMIGLIE.docx'))
first = True
for sp in src.paragraphs:
    if first and sp.text.strip().startswith('RICHIAMO'):
        first = False; continue
    first = False
    if not sp.text.strip():
        continue
    style_name = sp.style.name if sp.style.name in ('List Bullet', 'List Number') else None
    np = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    np.paragraph_format.space_after = Pt(3)
    for r in sp.runs:
        nr = np.add_run(r.text)
        nr.bold = r.bold; nr.italic = r.italic
        if r.font.size: nr.font.size = r.font.size
        if r.font.color and r.font.color.rgb: nr.font.color.rgb = r.font.color.rgb

# ============================ DORMIENTE ============================
chapter('13 · La coda — il secondo dormiente', 'da leggere al tavolo DOPO il finale, qualunque sia')
emit_markdown(load_md('SECONDO_DORMIENTE.md'))

# ============================ APPENDICE A: NASTRO ============================
chapter('A · Il nastro dei marcatori', '«dov\u2019è adesso?» — un marcatore ogni 2 ore, 12 km/h esatti (km reali = colonna Km/10)')
rows = [['Quando', 'Km', 'Dove']]
with io.open(os.path.join(BASE, 'percorso-marcatori-2h.csv'), encoding='utf-8') as f:
    rdr = list(csv.reader(f))
for rec in rdr[1:]:
    if len(rec) < 6: continue
    name = rec[1].strip()
    desc = ', '.join(x.strip() for x in rec[4:-1])
    try: km = float(rec[-1]) / 10.0
    except ValueError: continue
    kms = ('%g' % km)
    rows.append([name, kms, desc])
tbl = doc.add_table(rows=len(rows), cols=3)
tbl.style = 'Table Grid'
for i, rec in enumerate(rows):
    for j, val in enumerate(rec):
        c = tbl.cell(i, j); c.paragraphs[0].paragraph_format.space_after = Pt(0)
        run = c.paragraphs[0].add_run(val); run.font.size = Pt(8.5)
        if i == 0: run.bold = True

# ============================ APPENDICE B ============================
chapter('B · Materiali fuori da questo documento', 'handout stampati, immagini, mappe')
b = [
    "**Handout stampati** (già pronti): i **9 bollettini A4** con timbri (BOLLETTINI_STAMPA.html) · il **programma cantieri del capo-settore** con Kashiwazaki cerchiata (HANDOUT_MAPPA_CAPOSETTORE.html — da consegnare piegato dentro una carta stradale vera) · **l'esame del camioncino catturato** (HANDOUT_CAMIONCINO.html: rapporto tecnico + lettera di bordo; data e protocolli da compilare a penna) · la **mappa tattica del capannone** (MAPPA_CAPANNONE_A3.html / cartina fotografica).",
    "**Il messaggio di riscatto**: il testo è nel cap. 3 (sequenza, Giorno 1 ore 12:00) — da leggere o ricopiare a mano.",
    "**Le immagini** (cartella Immagini/): Copertina · 2 foto aeree (briefing) · Mappa_01/02 della V di Genkai («quello che si sa») · le 7 foto dei camioncini (F1-F7: porto, costiera, ponticello, recinzione, casello, zoom-fondina, interno) · le 3+2 del capannone (aerea-cartella, dal mare, cancello, cartina dall'alto, interno) · il kaiju sulla costa · P2/P10/P13 quando generate.",
    "**La mappa d'insieme**: Google My Maps del percorso (link in SCHEMA.md) + punti-mappa-avventura.csv per l'import.",
    "**I dossier Kage completi della campagna**: pg/Kage/*.md (qui dentro c'è solo la versione-Richiamo, cap. 12).",
]
for t in b:
    p = doc.add_paragraph(style='List Bullet'); render_inline(p, t)

doc.save(OUT)
print('OK ->', OUT)
