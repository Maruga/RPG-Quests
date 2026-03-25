from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(0.8)
section.bottom_margin = Cm(0.8)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Arial Narrow'
style.font.size = Pt(9)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Font constants
F_TITLE = 'Impact'       # Titoli, callsign, headers
F_BODY = 'Arial Narrow'  # Corpo testo, descrizioni
F_MONO = 'Courier New'   # Stat block, tracker, dati numerici

BASE = os.path.abspath(os.path.dirname(__file__))
STAMP = os.path.normpath(os.path.join(BASE, '..', 'AI', 'CLASSIFIED.png'))
PATCH = os.path.normpath(os.path.join(BASE, '..', 'AI', 'patch-prometheus.png'))
PHOTO = os.path.normpath(os.path.join(BASE, '..', 'Immagini', 'Operatori', 'Mikhail_Ghost_Russia.jpg'))

def shd(cell, color):
    s = cell._element.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    cell._element.get_or_add_tcPr().append(s)

def noborder(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = borders.makeelement(qn(f'w:{b}'), {qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'})
        borders.append(el)
    tblPr.append(borders)

def hdr(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = F_TITLE
    if level == 1:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '8', qn('w:space'): '1', qn('w:color'): '1a1a1a'})
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

def txt(text, bold=False, italic=False, size=9, color=None, align=None, sa=2, font=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.name = font or F_BODY
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(sa)

def edge(name, mech):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{name}: ')
    run.bold = True
    run.font.name = F_BODY
    run.font.size = Pt(8.5)
    run = p.add_run(mech)
    run.font.name = F_BODY
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd(c, '2c2c2c')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = F_MONO; r.font.size = Pt(8)
            if ri % 2 == 1: shd(c, 'f0f0f0')
    return table

def redacted(text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    rPr = run._element.get_or_add_rPr()
    hl = rPr.makeelement(qn('w:highlight'), {qn('w:val'): 'black'})
    rPr.append(hl)

def blackbar(sa=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run('██████████████████████████████████████████████████████')
    run.font.name = 'Courier New'; run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# ==========================================
# PAGINA 1 — FRONTE
# ==========================================

# Flag bar — Russia: bianco / blu / rosso
flag = doc.add_table(rows=1, cols=3)
flag.alignment = WD_TABLE_ALIGNMENT.CENTER
for c in flag.rows[0].cells:
    c.height = Cm(0.3); c.text = ''
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
shd(flag.rows[0].cells[0], 'FFFFFF')
shd(flag.rows[0].cells[1], '0039A6')
shd(flag.rows[0].cells[2], 'D52B1E')
noborder(flag)

# Header: patch + title + stamp
h = doc.add_table(rows=1, cols=3)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
noborder(h)
h.rows[0].cells[0].width = Cm(3); h.rows[0].cells[0].text = ''
p = h.rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(PATCH, width=Cm(2.3))

h.rows[0].cells[1].width = Cm(11); h.rows[0].cells[1].text = ''
p = h.rows[0].cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CLASSIFIED\n'); r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x99,0,0)
r = p.add_run('OPERAZIONE PROMETHEUS — DOSSIER 01/05'); r.font.name = F_BODY; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

h.rows[0].cells[2].width = Cm(3.5); h.rows[0].cells[2].text = ''
p = h.rows[0].cells[2].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run().add_picture(STAMP, width=Cm(3))

# Callsign + photo
lay = doc.add_table(rows=1, cols=2)
lay.alignment = WD_TABLE_ALIGNMENT.LEFT
noborder(lay)

left = lay.rows[0].cells[0]; left.width = Cm(12); left.text = ''
p = left.paragraphs[0]; p.paragraph_format.space_before = Pt(4)
r = p.add_run('"GHOST"'); r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(32)
p = left.add_paragraph()
r = p.add_run('MIKHAIL VORONOV'); r.font.name = F_TITLE; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x44,0x44,0x44)
p = left.add_paragraph(); p.paragraph_format.space_before = Pt(4)
for lbl, val in [('NAZIONALITA\'','Russia'), ('ETA\'','52'), ('UNITA\'','Spetsnaz GRU'), ('RUOLO','Scout / Infiltratore'), ('RANGO','Eroico')]:
    r = p.add_run(f'{lbl}: '); r.bold = True; r.font.name = F_BODY; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    r = p.add_run(f'{val}\n'); r.font.name = F_BODY; r.font.size = Pt(8)

right = lay.rows[0].cells[1]; right.width = Cm(5); right.text = ''
p = right.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run().add_picture(PHOTO, width=Cm(4.5))

# Background
hdr('Background Operativo')
txt('30 anni Spetsnaz GRU. Afghanistan, Cecenia, Georgia, Siria. Ricognizione profonda e eliminazioni silenziose. Silenzioso, pragmatico. Sorride raramente e sempre nel momento sbagliato. Beve vodka. Incubi ricorrenti.', size=8.5, sa=2)

# Attributi
hdr('Attributi e Derivati')
txt('Agilita\' d10 | Forza d8 | Intelligenza d6 | Spirito d6 | Vigore d8      Passo 8 | Parata 6 | Robustezza 10 (4)', bold=True, size=9, sa=2, font=F_MONO)

hdr('Abilita\'', level=2)
txt('Atletica d8 | Combattere d8 | Comuni d6 | Furtivita\' d12 | Percezione d10 | Persuasione d4 | Sopravvivenza d8 | Sparare d10', size=8.5, sa=2, font=F_MONO)

# Vantaggi
hdr('Vantaggi')
edge('Allerta', '+2 a Percezione')
edge('Scatto', 'Passo +2 (gia\' incluso nel Passo)')
edge('Schivata', '-1 ai tiri a distanza contro di te')
edge('Mira Ferma', '+1 a Sparare se non ti muovi')
edge('Scattante', 'Ridraw carte Iniziativa di 5 o meno')
edge('Sangue Freddo', 'Peschi 2 carte Iniziativa, tieni la migliore')
edge('Grinta', 'Ignori 1 livello di penalita\' Ferita')
edge('Senza Pieta\'', 'Spendi un Benny per ritirare i dadi danno')

# Svantaggi
hdr('Svantaggi')
edge('Segreto (Maggiore)', 'Lavaggio del cervello, ricordi cancellati. Se scoprono cosa ti hanno fatto, perdi un Benny')
edge('Cattiva Abitudine (Minore)', 'Alcol — bevi vodka regolarmente')
edge('Peculiarita\' (Minore)', 'Sorride nel momento sbagliato')

# Armamento
hdr('Armamento')
tbl(['ARMA','GITTATA','DANNO','PA','CDT','COLPI','NOTE'],
    [['AS Val','15/30/60','2d8','2','3','20','Auto, Silenziata'],
     ['SR-1 Vektor','12/24/48','2d6+1','2','1','18','Semi-auto'],
     ['NR-43 coltello','Mischia','For+d4','—','—','—','—']])

txt('Armatura: Giubbotto tattico +4 torso | Elmetto +4 testa — Equip: Visore notturno | Kit ricognizione | Dosimetro | Mappa Zona annotata', size=8, sa=2)

# Tracker
hdr('Tracker')
tracker = doc.add_table(rows=5, cols=2)
tracker.alignment = WD_TABLE_ALIGNMENT.LEFT
noborder(tracker)
data = [
    ('FERITE',  '[ ] [ ] [ ]  -1 per Ferita a tutti i tiri'),
    ('FATICA',  '[ ] [ ] [ ]  -1 per livello. A 3 = Incapacitato'),
    ('BENNIES', '[ ] [ ] [ ]  Ritirare tiro | Assorbire Ferita | Ritirare danno'),
    ('AS VAL',  '20: [ ][ ][ ][ ] (casella=5)'),
    ('SR-1',    '18: [ ][ ][ ] (casella=6)'),
]
for i, (lbl, val) in enumerate(data):
    c0 = tracker.rows[i].cells[0]; c0.text = ''; c0.width = Cm(2)
    r = c0.paragraphs[0].add_run(lbl); r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    c1 = tracker.rows[i].cells[1]; c1.text = ''
    r = c1.paragraphs[0].add_run(val); r.font.name = F_MONO; r.font.size = Pt(7.5)

# ==========================================
# PAGINA 2 — RETRO (SEGRETO)
# ==========================================
doc.add_page_break()

# Flag bar
flag2 = doc.add_table(rows=1, cols=3)
flag2.alignment = WD_TABLE_ALIGNMENT.CENTER
for c in flag2.rows[0].cells:
    c.height = Cm(0.3); c.text = ''
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
shd(flag2.rows[0].cells[0], 'FFFFFF')
shd(flag2.rows[0].cells[1], '0039A6')
shd(flag2.rows[0].cells[2], 'D52B1E')
noborder(flag2)

# Stamp grande
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.space_before = Pt(4)
p.add_run().add_picture(STAMP, width=Cm(7))

txt('RISERVATO — SOLO PER TE — NON CONDIVIDERE', bold=True, size=14,
    color=RGBColor(0xBF, 0x0A, 0x30), align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

hdr('Il Tuo Segreto')
txt('', sa=6)
txt('Sei gia\' stato in questo bunker.', size=12, bold=True, sa=4)
txt('Era il 1986. Avevi 23 anni. Spetsnaz.', size=11, sa=2)
txt('Hai visto il laboratorio. Le creature. I tuoi compagni morire.', size=11, sa=4)

txt('Sei arrivato al Piano -2. Ferito gravemente, portato fuori in barella.', size=11, sa=4)
txt('Poi il lavaggio del cervello: farmaci, ipnosi, elettroshock.', size=11, sa=2)
txt('I ricordi sono stati cancellati.', size=11, bold=True, sa=6)
txt('Ma stanno tornando.', size=14, bold=True, sa=10)

hdr('Quello Che Non Sai')
txt('L\'italiano — Marco "Torcia" Ferrante — era all\'esterno nel 1986. Ti ha visto sulla barella. Ti ha lasciato passare, disobbedendo agli ordini.', size=10.5, sa=4)
txt('Tu non lo sai. Non lo ricordi. Ma qualcosa in lui ti sembra familiare.', size=10.5, italic=True, sa=8)

hdr('Durante La Sessione')
txt('Avrai flashback sempre piu\' intensi man mano che scendi nel bunker.', size=10.5, sa=4)
txt('Il GM ti passera\' fogliettini con immagini, suoni, frammenti di memoria.', size=10.5, sa=4)
txt('Corridoi che riconosci. Stanze gia\' viste. Nomi che affiorano.', size=10.5, sa=8)

txt('', sa=6)
txt('Hai incubi ricorrenti su qualcosa che non riesci a ricordare.', bold=True, size=12, sa=4)
txt('Corridoi. Luci rosse. Urla.', size=11, sa=4)
txt('La Zona ti e\' familiare. Non sai perche\'.', size=11, bold=True, sa=10)

hdr('Come Giocare')
txt('Reagisci ai flashback — confusione, disagio, paura. Non sai cosa significano.', size=10, sa=4)
txt('Se qualcosa ti sembra familiare, dillo ad alta voce: "Questo posto... lo conosco."', size=10, sa=4)
txt('Non rivelare agli altri il contenuto dei fogliettini del GM.', size=10, sa=10)

txt('CLASSIFIED // OPERAZIONE PROMETHEUS // NATO SECRET', bold=True, size=7,
    color=RGBColor(0x99,0,0), align=WD_ALIGN_PARAGRAPH.CENTER)

# SAVE
out = os.path.join(BASE, 'RU_Mikhail_Ghost_Scheda.docx')
doc.save(out)
print(f'Salvata: {out}')
