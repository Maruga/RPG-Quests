from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(0.8)
section.bottom_margin = Cm(0.8)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Arial Narrow'
style.font.size = Pt(9)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Font constants
F_TITLE = 'Impact'       # Titoli, callsign, headers
F_BODY = 'Arial Narrow'  # Corpo testo, descrizioni
F_MONO = 'Courier New'   # Stat block, tracker, dati numerici

BASE = os.path.abspath(os.path.dirname(__file__))
STAMP = os.path.normpath(os.path.join(BASE, '..', 'AI', 'CLASSIFIED.png'))
PATCH = os.path.normpath(os.path.join(BASE, '..', 'AI', 'patch-prometheus.png'))
PHOTO = os.path.normpath(os.path.join(BASE, '..', 'Immagini', 'Operatori', 'Marco_Torcia_Italia.jpg'))

def shd(cell, color):
    s = cell._element.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    cell._element.get_or_add_tcPr().append(s)

def noborder(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = borders.makeelement(qn(f'w:{b}'), {qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'})
        borders.append(el)
    tblPr.append(borders)

def hdr(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = F_TITLE
    if level == 1:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '8', qn('w:space'): '1', qn('w:color'): '1a1a1a'})
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

def txt(text, bold=False, italic=False, size=9, color=None, align=None, sa=2, font=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.name = font or F_BODY
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(sa)

def edge(name, mech):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{name}: ')
    run.bold = True
    run.font.name = F_BODY
    run.font.size = Pt(8.5)
    run = p.add_run(mech)
    run.font.name = F_BODY
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd(c, '2c2c2c')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = F_MONO; r.font.size = Pt(8)
            if ri % 2 == 1: shd(c, 'f0f0f0')
    return table

def redacted(text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    rPr = run._element.get_or_add_rPr()
    hl = rPr.makeelement(qn('w:highlight'), {qn('w:val'): 'black'})
    rPr.append(hl)

def blackbar(sa=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run('██████████████████████████████████████████████████████')
    run.font.name = 'Courier New'; run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# ==========================================
# PAGINA 1 — FRONTE
# ==========================================

# Flag bar — Italia: verde / bianco / rosso
flag = doc.add_table(rows=1, cols=3)
flag.alignment = WD_TABLE_ALIGNMENT.CENTER
for c in flag.rows[0].cells:
    c.height = Cm(0.3); c.text = ''
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
shd(flag.rows[0].cells[0], '009246')
shd(flag.rows[0].cells[1], 'FFFFFF')
shd(flag.rows[0].cells[2], 'CE2B37')
noborder(flag)

# Header: patch + title + stamp
h = doc.add_table(rows=1, cols=3)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
noborder(h)
h.rows[0].cells[0].width = Cm(3); h.rows[0].cells[0].text = ''
p = h.rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(PATCH, width=Cm(2.3))

h.rows[0].cells[1].width = Cm(11); h.rows[0].cells[1].text = ''
p = h.rows[0].cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CLASSIFIED\n'); r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x99,0,0)
r = p.add_run('OPERAZIONE PROMETHEUS — DOSSIER 02/05'); r.font.name = F_BODY; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

h.rows[0].cells[2].width = Cm(3.5); h.rows[0].cells[2].text = ''
p = h.rows[0].cells[2].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run().add_picture(STAMP, width=Cm(3))

# Callsign + photo
lay = doc.add_table(rows=1, cols=2)
lay.alignment = WD_TABLE_ALIGNMENT.LEFT
noborder(lay)

left = lay.rows[0].cells[0]; left.width = Cm(12); left.text = ''
p = left.paragraphs[0]; p.paragraph_format.space_before = Pt(4)
r = p.add_run('"TORCIA"'); r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(32)
p = left.add_paragraph()
r = p.add_run('MARCO FERRANTE'); r.font.name = F_TITLE; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x44,0x44,0x44)
p = left.add_paragraph(); p.paragraph_format.space_before = Pt(4)
for lbl, val in [('NAZIONALITA\'','Italia'), ('ETA\'','49'), ('UNITA\'','9° Rgt. d\'Assalto "Col Moschin"'), ('RUOLO','DMR / Lanciafiamme'), ('RANGO','Eroico')]:
    r = p.add_run(f'{lbl}: '); r.bold = True; r.font.name = F_BODY; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    r = p.add_run(f'{val}\n'); r.font.name = F_BODY; r.font.size = Pt(8)

right = lay.rows[0].cells[1]; right.width = Cm(5); right.text = ''
p = right.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run().add_picture(PHOTO, width=Cm(4.5))

# Background
hdr('Background Operativo')
txt('29 anni Col Moschin. Afghanistan, Iraq, Libano, Libia. Soprannome "Torcia" da un incidente in Somalia con un lanciafiamme. Ironico, rilassato in superficie, professionista letale sotto. Canticchia opera durante i combattimenti.', size=8.5, sa=2)

# Attributi
hdr('Attributi e Derivati')
txt('Agilita\' d8 | Forza d8 | Intelligenza d6 | Spirito d8 | Vigore d8      Passo 6 | Parata 5 | Robustezza 10 (4)', bold=True, size=9, sa=2, font=F_MONO)

hdr('Abilita\'', level=2)
txt('Atletica d8 | Combattere d6 | Comuni d6 | Furtivita\' d8 | Intimidire d6 | Percezione d8 | Persuasione d4 | Sopravvivenza d6 | Sparare d10', size=8.5, sa=2, font=F_MONO)

# Vantaggi
hdr('Vantaggi')
edge('Tiratore Scelto', '+1 danno a Gittata Corta, +2 con Incremento')
edge('Grinta Migliorata', 'Ignori 2 livelli di penalita\' Ferita')
edge('Sangue Freddo', 'Peschi 2 carte Iniziativa, tieni la migliore')
edge('Duro a Morire', 'Ignori la penalita\' Ferita ai tiri Vigore per non morire')
edge('Ferrea Volonta\'', '+2 a Spirito e Intelligenza per resistere a Test')
edge('Senza Pieta\'', 'Spendi un Benny per ritirare i dadi danno')
edge('Mira Ferma', '+1 a Sparare se non ti muovi')

# Svantaggi
hdr('Svantaggi')
edge('Segreto (Maggiore)', 'I fatti del 1986 — quello che hai fatto e quello che ti e\' successo. Se scoprono, perdi un Benny')
edge('Cattiva Abitudine (Minore)', 'Fuma costantemente')
edge('Peculiarita\' (Minore)', 'Canticchia opera in combattimento')

# Armamento
hdr('Armamento')
tbl(['ARMA','GITTATA','DANNO','PA','CDT','COLPI','NOTE'],
    [['HK417','30/60/120','2d8+1','2','1','20','Semi-auto'],
     ['LPO-50','Cono Medio','3d6','—','1','3','Lanciafiamme'],
     ['Beretta 92FS','12/24/48','2d6','1','1','15','Semi-auto']])

txt('Armatura: Giubbotto tattico +4 torso | Elmetto +4 testa — Equip: Ottica Leupold Mk 4 | Binocolo telemetro | Taniche extra combustibile x2 | Dosimetro | Sigarette', size=8, sa=2)

# Tracker
hdr('Tracker')
tracker = doc.add_table(rows=5, cols=2)
tracker.alignment = WD_TABLE_ALIGNMENT.LEFT
noborder(tracker)
data = [
    ('FERITE',  '[ ] [ ] [ ]  -1 per Ferita a tutti i tiri'),
    ('FATICA',  '[ ] [ ] [ ]  -1 per livello. A 3 = Incapacitato'),
    ('BENNIES', '[ ] [ ] [ ]  Ritirare tiro | Assorbire Ferita | Ritirare danno'),
    ('HK417',   '20: [ ][ ][ ][ ] (casella=5)'),
    ('LPO/92FS','LPO: [ ][ ][ ] (1 carica cad.)  92FS: [ ][ ] (7+8)'),
]
for i, (lbl, val) in enumerate(data):
    c0 = tracker.rows[i].cells[0]; c0.text = ''; c0.width = Cm(2)
    r = c0.paragraphs[0].add_run(lbl); r.bold = True; r.font.name = F_TITLE; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    c1 = tracker.rows[i].cells[1]; c1.text = ''
    r = c1.paragraphs[0].add_run(val); r.font.name = F_MONO; r.font.size = Pt(7.5)

# ==========================================
# PAGINA 2 — RETRO (SEGRETO)
# ==========================================
doc.add_page_break()

# Flag bar
flag2 = doc.add_table(rows=1, cols=3)
flag2.alignment = WD_TABLE_ALIGNMENT.CENTER
for c in flag2.rows[0].cells:
    c.height = Cm(0.3); c.text = ''
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
shd(flag2.rows[0].cells[0], '009246')
shd(flag2.rows[0].cells[1], 'FFFFFF')
shd(flag2.rows[0].cells[2], 'CE2B37')
noborder(flag2)

# Stamp grande
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.space_before = Pt(4)
p.add_run().add_picture(STAMP, width=Cm(7))

txt('RISERVATO — SOLO PER TE — NON CONDIVIDERE', bold=True, size=14,
    color=RGBColor(0xBF, 0x0A, 0x30), align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

hdr('Il Tuo Segreto — Parte 1')
txt('', sa=4)
txt('Nel 1986 eri all\'esterno del bunker. Team NATO.', size=12, bold=True, sa=4)
txt('L\'ordine era chiaro: chi esce dal bunker muore e viene bruciato.', size=11, sa=4)

txt('Hai ucciso civili che imploravano di essere lasciati passare.', size=11, sa=2)
txt('Uomini, donne. Terrorizzati. Tu hai eseguito gli ordini.', size=11, sa=6)

txt('Poi e\' arrivato un soldato russo in barella. Giovane, ferito gravemente.', size=11, sa=2)
txt('Hai disobbedito. L\'hai lasciato passare.', size=11, bold=True, sa=4)
txt('Quel soldato era Mikhail "Ghost" Voronov.', size=11, sa=2)
txt('Lui non lo sa. Non si ricorda di te.', size=11, italic=True, sa=8)

hdr('Il Tuo Segreto — Parte 2')
txt('', sa=4)
txt('Sei contagiato. Stadio 1. Dal 1986.', size=12, bold=True, sa=4)
txt('Due settimane fa sono iniziati i sintomi: mal di testa, sogni vividi, assenze.', size=11, sa=4)
txt('I tuoi superiori pensano che tu sia immune.', size=11, sa=2)
txt('Non lo sei.', size=11, bold=True, sa=6)

txt('L\'IA puo\' controllarti. Impulsi, blackout, suggestioni.', size=11, sa=4)
txt('Il GM ti passera\' fogliettini durante la sessione.', size=11, sa=2)
txt('Devi reagire a quello che leggi SENZA rivelare il contenuto.', size=11, bold=True, sa=8)

hdr('Cosa Sai')
txt('Riconosci il russo — Mikhail. L\'hai visto nel 1986 sulla barella. L\'hai lasciato passare.', size=10.5, sa=4)
txt('Lui non sembra ricordarsi di te.', size=10.5, italic=True, sa=8)

hdr('Come Giocare')
txt('Sei stato esposto a qualcosa nel 1986. E\' rimasto dormiente fino a due settimane fa.', size=10, sa=4)
txt('E\' FONDAMENTALE che tu reagisca a quello che leggi nei fogliettini senza rivelare il contenuto.', size=10, bold=True, sa=4)
txt('Mal di testa improvvisi, assenze, momenti in cui "non sei tu" — recitali.', size=10, sa=4)
txt('Se il GM ti dice di fare qualcosa nel fogliettino, FALLO. E\' l\'IA che ti controlla.', size=10, sa=4)
txt('Se ti chiedono se stai bene: menti. Sempre.', size=10, bold=True, sa=10)

txt('CLASSIFIED // OPERAZIONE PROMETHEUS // NATO SECRET', bold=True, size=7,
    color=RGBColor(0x99,0,0), align=WD_ALIGN_PARAGRAPH.CENTER)

# SAVE
out = os.path.join(BASE, 'IT_Marco_Torcia_Scheda.docx')
doc.save(out)
print(f'Salvata: {out}')
