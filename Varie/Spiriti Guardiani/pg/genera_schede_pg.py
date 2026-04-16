"""
Generatore Schede PG — Spiriti Guardiani
Stile: fantasy romano (oro, rosso scuro, pergamena) con accenti kanji.
Obiettivo: singola pagina A4.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.normpath(os.path.join(BASE, "..", "Immagini"))

# Paletta Fantasy Romano
OLD_GOLD   = RGBColor(0xB4, 0x95, 0x4B)
ROMAN_RED  = RGBColor(0x7A, 0x1F, 0x1F)
BRONZE     = RGBColor(0x4A, 0x34, 0x24)
INK        = RGBColor(0x1A, 0x14, 0x10)
SHADOW     = RGBColor(0x6B, 0x5B, 0x45)
CREAM_HEX  = "FBF4E4"
PARCH_HEX  = "F2E6C8"
GOLD_HEX   = "B4954B"
RED_HEX    = "7A1F1F"


def run(p, text, size=10, bold=False, italic=False, color=INK, font=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if font:
        r.font.name = font
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
    return r


def no_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none"/><w:left w:val="none"/>'
        '<w:right w:val="none"/><w:bottom w:val="none"/>'
        '</w:tcBorders>'))


def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))


def set_width(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcW')):
        tcPr.remove(ex)
    tcPr.append(parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(cm*567)}" w:type="dxa"/>'))


def valign(cell, val="center"):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>'))


def border(cell, side, size=6, color=GOLD_HEX):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)
    for ex in tcBorders.findall(qn(f'w:{side}')):
        tcBorders.remove(ex)
    tcBorders.append(parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{size}" w:color="{color}"/>'))


def box(cell, color=GOLD_HEX, size=6):
    for s in ("top", "left", "right", "bottom"):
        border(cell, s, size, color)


def ornament(doc, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run(p, "\u2756   \u22c6   \u2756   \u22c6   \u2756", size=size, color=OLD_GOLD)


def section_title(doc, kanji, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run(p, kanji, size=12, color=OLD_GOLD)
    run(p, "   ", size=10)
    run(p, label.upper(), size=10, bold=True, color=ROMAN_RED)


PG_TITUS = {
    "nome": "TITUS",
    "titolo": "IL BARBONE",
    "eta": 23,
    "accusa": "Rissa in una taverna. Ha rotto la mascella a un nobile.",
    "chi_e": (
        "Ex soldato. Ha combattuto sulle fosse. Ha visto troppo, ha bevuto per "
        "dimenticare. Le mani ricordano ancora come si impugna un'arma, anche "
        "se la testa vorrebbe scordarlo."
    ),
    "kage": (
        "L'alcol. Non riesce a smettere. Ha iniziato a bere dopo una battaglia "
        "in cui il suo plotone e' stato massacrato. E' l'unico sopravvissuto."
    ),
    "arma": "Martello da Guerra",
    "animale": "Salamandra di Fuoco",
    "nome_spirito": "Caesar",
    "condottiero": "Giulio Cesare",
    "caratteristica": "Scintilla (\u706b)",
    "caratteristica_nome": "Scintilla",
    "spirito_personalita": (
        "Fuoco trattenuto. Brucia a fuoco lento, rispetta chi agisce, "
        "disprezza chi esita. Alea iacta est."
    ),
    "spirito_passivo": (
        "SANGUE ALCHEMICO \u2014 Il sangue scalda le vene come brace sotto la cenere. "
        "Non teme il freddo, le fiamme non lo ustionano, la sua arma fuma quando colpisce."
    ),
    "spirito_attivo": (
        "MANI DI BRACE (costa Ki) \u2014 Ogni colpo a segno +2 danno e incendia il "
        "bersaglio (1 danno residuo). Dura 1 round per ogni Ki speso."
    ),
    "spirito_fusione": (
        "CONFLAGRAZIONE (1/sessione, 3 round, scelta dallo spirito) \u2014 Scintilla 9, "
        "Ki al massimo, immune al fuoco e al danno fisico, ogni attacco e' critico automatico."
    ),
    "attributi": [
        ("Radice",    "\u6839", 7, 7),
        ("Eco",       "\u97ff", 5, 5),
        ("Scintilla", "\u706b", 7, 8),
        ("Ombra",     "\u5f71", 5, 5),
        ("Flusso",    "\u6d41", 6, 6),
        ("Battito",   "\u62cd", 6, 6),
        ("Ki",        "\u6c17", 10, 13),
    ],
    "armi_evoluzione": [
        ("IMPROVVISATA",    "Gamba di tavolo con pietre",   "Legno scheggiato, pietre legate in cima"),
        ("FORMA I",   "Mazza chiodata",               "Ferro grezzo, chiodi sporgenti, pesante"),
        ("FORMA II",  "Martello da guerra a una mano","Ferro temprato, impugnatura di cuoio rosso"),
        ("FORMA III", "Malleus Iacta",                "Bronzo incandescente, lascia impronte di fuoco"),
    ],
    "immagine": os.path.join(IMG_DIR, "IL BARBONE - Titus.png"),
}


def crea_scheda(pg, output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.4)
    sec.right_margin = Cm(1.4)

    base_style = doc.styles['Normal']
    base_style.font.name = 'Garamond'
    base_style.font.size = Pt(10)

    # TITOLO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run(p, pg["titolo"], size=12, bold=True, color=OLD_GOLD, font="Trajan Pro")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run(p, pg["nome"], size=28, bold=True, color=ROMAN_RED, font="Trajan Pro")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run(p, f"{pg['animale']}  \u00b7  {pg['nome_spirito']}  \u00b7  {pg['caratteristica']}",
        size=9.5, italic=True, color=SHADOW)

    ornament(doc)

    # RITRATTO + INFO
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    left, right = t.rows[0].cells
    set_width(left, 5.5)
    set_width(right, 12.5)
    no_borders(left)
    no_borders(right)

    if os.path.exists(pg["immagine"]):
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(pg["immagine"], width=Cm(5.0))

    right.paragraphs[0].text = ''
    info_t = right.add_table(rows=4, cols=2)
    info_t.autofit = False
    info_rows = [
        ("ETA'",         str(pg["eta"])),
        ("ACCUSA",       pg["accusa"]),
        ("ARMA",         pg["arma"]),
        ("CONDOTTIERO",  pg["condottiero"]),
    ]
    for i, (label, value) in enumerate(info_rows):
        c_lbl = info_t.rows[i].cells[0]
        c_val = info_t.rows[i].cells[1]
        set_width(c_lbl, 3.0)
        set_width(c_val, 9.0)
        no_borders(c_lbl)
        no_borders(c_val)
        p = c_lbl.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run(p, label, size=8, bold=True, color=OLD_GOLD)
        p = c_val.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run(p, value, size=9.5, color=INK)

    # CHI E'
    section_title(doc, "\u4eba", "Chi e'")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, pg["chi_e"], size=10, italic=True, color=INK)

    # KAGE
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    kcell = t.rows[0].cells[0]
    set_width(kcell, 18.0)
    shade(kcell, PARCH_HEX)
    border(kcell, "left", 24, RED_HEX)
    border(kcell, "top", 4, GOLD_HEX)
    border(kcell, "bottom", 4, GOLD_HEX)
    border(kcell, "right", 4, GOLD_HEX)

    p = kcell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run(p, "\u5f71   KAGE  \u2014  Il peso del passato", size=9, bold=True, color=ROMAN_RED)
    p = kcell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, pg["kage"], size=9.5, italic=True, color=INK)

    # ATTRIBUTI
    section_title(doc, "\u529b", "Attributi")
    t = doc.add_table(rows=3, cols=7)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    for i, (name, kanji, base, after) in enumerate(pg["attributi"]):
        is_star = name == pg["caratteristica_nome"]
        hdr_color = ROMAN_RED if is_star else INK
        val_color = ROMAN_RED if is_star else INK
        shade_hex = "F2DFA0" if is_star else CREAM_HEX

        c = t.rows[0].cells[i]
        set_width(c, 2.55)
        shade(c, shade_hex)
        valign(c)
        box(c, GOLD_HEX, 6)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run(p, kanji, size=14, color=OLD_GOLD, font="MS Mincho")

        c = t.rows[1].cells[i]
        set_width(c, 2.55)
        shade(c, shade_hex)
        valign(c)
        box(c, GOLD_HEX, 6)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run(p, name.upper(), size=8, bold=True, color=hdr_color)

        c = t.rows[2].cells[i]
        set_width(c, 2.55)
        shade(c, "FFFBEE")
        valign(c)
        box(c, GOLD_HEX, 6)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run(p, str(after), size=18, bold=True, color=val_color, font="Trajan Pro")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run(p, "\u2605  bonus dello spirito gia' applicato \u00b7 collegamento nelle Terme",
        size=8, italic=True, color=SHADOW)

    # SPIRITO
    section_title(doc, "\u970a", f"Spirito \u2014 {pg['animale']}")

    t = doc.add_table(rows=4, cols=1)
    t.autofit = False
    for r in t.rows:
        set_width(r.cells[0], 18.0)
        shade(r.cells[0], CREAM_HEX)

    c = t.rows[0].cells[0]
    box(c, GOLD_HEX, 4)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run(p, "CARATTERE  \u00b7  ", size=8, bold=True, color=OLD_GOLD)
    run(p, pg["spirito_personalita"], size=9.5, italic=True, color=INK)

    poteri = [
        ("PASSIVO",  pg["spirito_passivo"]),
        ("ATTIVO",   pg["spirito_attivo"]),
        ("FUSIONE",  pg["spirito_fusione"]),
    ]
    for idx, (label, testo) in enumerate(poteri, start=1):
        c = t.rows[idx].cells[0]
        box(c, GOLD_HEX, 4)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run(p, f"{label}  \u00b7  ", size=8, bold=True, color=ROMAN_RED)
        run(p, testo, size=9.5, color=INK)

    # EVOLUZIONE ARMA
    section_title(doc, "\u6b66", f"Arma \u2014 {pg['arma']}")

    t = doc.add_table(rows=4, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    header_shade = ["DCC792", "E7D7A5", "F1E3B8", "F8EFD2"]
    for i, (fase, nome, aspetto) in enumerate(pg["armi_evoluzione"]):
        c1 = t.rows[i].cells[0]
        c2 = t.rows[i].cells[1]
        c3 = t.rows[i].cells[2]
        set_width(c1, 2.5)
        set_width(c2, 5.5)
        set_width(c3, 10.0)
        shade(c1, header_shade[i])
        shade(c2, CREAM_HEX)
        shade(c3, CREAM_HEX)
        box(c1, GOLD_HEX, 4)
        box(c2, GOLD_HEX, 4)
        box(c3, GOLD_HEX, 4)
        valign(c1); valign(c2); valign(c3)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        is_final = i == 3
        run(p, fase, size=9, bold=True, color=ROMAN_RED if is_final else BRONZE)
        p = c2.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run(p, nome, size=9.5, bold=is_final, color=INK)
        p = c3.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run(p, aspetto, size=9, italic=True, color=SHADOW)

    ornament(doc)

    doc.save(output_path)
    print(f"[OK] Scheda salvata: {output_path}")


if __name__ == "__main__":
    out = os.path.join(BASE, f"ANTEPRIMA_{PG_TITUS['nome']}.docx")
    crea_scheda(PG_TITUS, out)
